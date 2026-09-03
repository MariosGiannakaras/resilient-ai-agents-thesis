#!/usr/bin/env python3
"""T-711 review-ready Word composition hardening adapter.

The base builder remains the auditable composition implementation. This adapter applies
only deterministic presentation/metadata fixes discovered by full rendered-page QA:
- verified citation-ready bibliographic identities instead of noisy catalog display titles;
- readable inline/display math without raw LaTeX residue;
- ordered-list restart at each independent Markdown list block;
- height-aware registered-figure placement and caption adjacency;
- manifest-role-aware main-text vs appendix figure routing;
- removal/normalization of T-710/T-711 drafting and placement instructions;
- meaningful cached TOC/list field results while preserving automatic Word fields;
- stricter post-save residue/reference QA.

No scientific estimand, value, interval, denominator, censoring decision, or asset byte is
computed or changed here.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

import build_review_ready_thesis_docx as builder


ORDER_STYLE = "T711 Ordered List"
_current_render_mode = "main"
_caption_counters = {"Σχήμα": 0, "Πίνακας": 0}

_original_render_markdown = builder.render_markdown
_original_add_asset_table = builder.add_asset_table
_original_add_markdown_table = builder.add_markdown_table
_original_build = builder.build


def _normalize_math_text(value: str) -> str:
    replacements = {
        r"\mathbb{R}": "ℝ",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\le": "≤",
        r"\ge": "≥",
        r"\leftarrow": "←",
        r"\rightarrow": "→",
        r"\gamma": "γ",
        r"\alpha": "α",
        r"\beta": "β",
        r"\pi": "π",
        r"\mu": "μ",
        r"\sigma": "σ",
        r"\Delta": "Δ",
        r"\delta": "δ",
        r"\epsilon": "ε",
        r"\varepsilon": "ε",
        r"\times": "×",
        r"\pm": "±",
        r"\in": "∈",
        r"\sum": "Σ",
        r"\max": "max",
        r"\min": "min",
        r"\mid": "|",
        r"\left": "",
        r"\right": "",
        r"\;": " ",
        r"\,": " ",
        r"\!": "",
    }
    out = value.strip()
    out = re.sub(r"\\mathcal\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\\text\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\\operatorname\{([^{}]+)\}", r"\1", out)
    for source, target in replacements.items():
        out = out.replace(source, target)
    out = out.replace(r"\{", "{").replace(r"\}", "}")
    out = out.replace("\\", "")
    return out


def _math_atom(value: str) -> str:
    value = _normalize_math_text(value)
    return value.replace("{", "").replace("}", "")


def _add_math_runs(paragraph, expression: str, size: float = 11):
    expression = _normalize_math_text(expression)
    token = re.compile(r"([_^])(?:\{([^{}]*)\}|([A-Za-z0-9+\-πγαβμσΔδεℝ]+))")
    pos = 0
    for match in token.finditer(expression):
        if match.start() > pos:
            plain = expression[pos:match.start()].replace("{", "").replace("}", "")
            if plain:
                run = paragraph.add_run(plain)
                builder.set_run_font(run, name="Cambria Math", size=size)
        atom = _math_atom(match.group(2) or match.group(3) or "")
        run = paragraph.add_run(atom)
        builder.set_run_font(run, name="Cambria Math", size=size)
        if match.group(1) == "_":
            run.font.subscript = True
        else:
            run.font.superscript = True
        pos = match.end()
    if pos < len(expression):
        tail = expression[pos:].replace("{", "").replace("}", "")
        if tail:
            run = paragraph.add_run(tail)
            builder.set_run_font(run, name="Cambria Math", size=size)


def _add_inline(paragraph, value: str, cmap: dict[str, int]):
    value = builder.replace_citations(value, cmap)
    token_re = re.compile(r"(\\\(.+?\\\)|\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    pos = 0
    for match in token_re.finditer(value):
        if match.start() > pos:
            run = paragraph.add_run(value[pos:match.start()])
            builder.set_run_font(run, size=11)
        token = match.group(0)
        if token.startswith(r"\("):
            _add_math_runs(paragraph, token[2:-2], 11)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            builder.set_run_font(run, size=11, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            builder.set_run_font(run, name="Consolas", size=9.5)
        else:
            run = paragraph.add_run(token[1:-1])
            builder.set_run_font(run, size=11, italic=True)
        pos = match.end()
    if pos < len(value):
        run = paragraph.add_run(value[pos:])
        builder.set_run_font(run, size=11)


def _add_equation(doc: Document, latex: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    _add_math_runs(paragraph, latex, 11)
    return paragraph


def _ordered_style(doc: Document):
    try:
        return doc.styles[ORDER_STYLE]
    except KeyError:
        style = doc.styles.add_style(ORDER_STYLE, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        builder.set_style_font(style, "Times New Roman", 11)
        return style


def _add_list_paragraph(doc: Document, value: str, level: int, ordered: bool, cmap: dict[str, int]):
    if not ordered:
        return builder.Document.add_paragraph if False else _add_bullet(doc, value, level, cmap)
    style = _ordered_style(doc)
    last = doc.paragraphs[-1] if doc.paragraphs else None
    previous_is_same_block = last is not None and last.style.name == ORDER_STYLE and getattr(doc, "_t711_order_level", None) == level
    number = getattr(doc, "_t711_order_number", 0) + 1 if previous_is_same_block else 1
    doc._t711_order_number = number
    doc._t711_order_level = level
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Cm(0.65 * level)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(f"{number}. ")
    builder.set_run_font(run, size=11)
    _add_inline(paragraph, value, cmap)
    return paragraph


def _add_bullet(doc: Document, value: str, level: int, cmap: dict[str, int]):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.65 * level)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.35
    _add_inline(paragraph, value, cmap)
    return paragraph


def _add_seq_caption(doc: Document, label: str, caption: str, prefix: str = ""):
    _caption_counters[label] = _caption_counters.get(label, 0) + 1
    number = _caption_counters[label]
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    if prefix:
        run = paragraph.add_run(prefix)
        builder.set_run_font(run, size=10, bold=True)
    run = paragraph.add_run(f"{label} ")
    builder.set_run_font(run, size=10, bold=True)
    builder.add_field(paragraph, f" SEQ {label} \\* ARABIC ", str(number))
    run = paragraph.add_run(f" — {caption}")
    builder.set_run_font(run, size=10)
    return paragraph


def _add_figure(doc: Document, asset: dict, inserted: list[str]):
    png = builder.choose_output(asset, "png")
    if png is None:
        return
    with Image.open(png) as image:
        width_px, height_px = image.size
    ratio = height_px / max(width_px, 1)
    max_width = 6.1
    max_height = 7.2
    width_inches = min(max_width, max_height / max(ratio, 0.01))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    if ratio > 1.12:
        paragraph.paragraph_format.page_break_before = True
    run = paragraph.add_run()
    run.add_picture(str(png), width=Inches(width_inches))
    _add_seq_caption(doc, "Σχήμα", asset.get("caption") or asset["asset_id"])
    inserted.append(asset["asset_id"])


def _mark_table_for_pagination(table):
    if not table.rows:
        return
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_asset_table(doc: Document, asset: dict, inserted: list[str]):
    normalized = dict(asset)
    normalized.setdefault("caption", normalized.get("title") or normalized["asset_id"])
    before = len(doc.tables)
    result = _original_add_asset_table(doc, normalized, inserted)
    if len(doc.tables) > before:
        _mark_table_for_pagination(doc.tables[-1])
    return result


def _add_markdown_table(doc: Document, lines: list[str], section_title: str, cmap: dict[str, int], table_counter: list[int]):
    before = len(doc.tables)
    result = _original_add_markdown_table(doc, lines, section_title, cmap, table_counter)
    if len(doc.tables) > before:
        _mark_table_for_pagination(doc.tables[-1])
    return result


def _insert_asset_ids(doc: Document, ids, assets: dict[str, dict], inserted: list[str], inserted_rq_tables: set[str]):
    for asset_id in ids:
        asset = assets.get(asset_id)
        if asset is None or asset_id in inserted:
            continue
        intended = set(asset.get("intended_use", []))
        if _current_render_mode == "main" and "main-thesis" not in intended:
            continue
        if _current_render_mode == "appendix" and "appendix" not in intended:
            continue
        if _current_render_mode not in {"main", "appendix"}:
            continue
        if asset.get("kind") == "figure":
            _add_figure(doc, asset, inserted)
        # Registered CSV tables remain available as evidence, but T-711 no longer
        # auto-inserts every RQ table merely because a figure from that RQ appears.
        # The accepted manuscript's compact result tables remain the main Word tables.


def _appendix_transform(line: str) -> str | None:
    stripped = line.strip()
    transforms = {
        "Να καταγραφούν διακριτά:": "Η αλυσίδα εκτέλεσης διακρίνει τα ακόλουθα στοιχεία:",
        "Να συνοψιστούν:": "Το αναπαραγώγιμο λογισμικό περιβάλλον συνοψίζεται ως εξής:",
        "Να περιγραφούν πλήρως οι τέσσερις frozen Phase-B conditions:": "Οι τέσσερις frozen Phase-B conditions είναι:",
    }
    if stripped in transforms:
        return transforms[stripped]
    if stripped.startswith("Να καταγραφούν οι δύο held-out final 7×7 layouts"):
        return "Οι δύο held-out final 7×7 layouts, το κοινό shortest-path length, το episode limit και τα 12 independent root IDs τεκμηριώνονται στο frozen protocol και στο Κεφάλαιο 3. Οι root seeds παραμένουν reproducibility metadata και δεν ερμηνεύονται ως method hyperparameters."
    if stripped.startswith("Να συμπεριληφθούν ως reproducibility metadata"):
        return "Τα ακόλουθα canonical hashes αποτελούν reproducibility metadata των accepted outputs:"
    if stripped.startswith("Κάθε conditional recovery-time table/figure πρέπει"):
        return "Στα conditional recovery-time summaries εμφανίζεται το recovered n. Οι right-censored roots παραμένουν censored και δεν εμφανίζονται ως observed recovery time 256."
    return line


def _preprocess_markdown(md: str, mode: str) -> str:
    lines = md.splitlines()
    out: list[str] = []
    i = 0
    skip_config_bullets = False
    skip_app_section = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if mode == "main" and "Προτεινόμενη T-711" in stripped:
            ids = builder.ASSET_RE.findall(stripped)
            if ids:
                out.append("Προτεινόμενη T-711 τοποθέτηση: " + ", ".join(ids))
            i += 1
            continue

        if mode == "glossary" and ("Draft authority" in stripped or "T-711" in stripped):
            i += 1
            continue

        if mode == "appendix":
            if stripped.startswith("# Παραρτήματα"):
                out.append("# Παραρτήματα")
                i += 1
                continue
            if stripped.startswith("## T-711 appendix assembly rules"):
                break
            if stripped.startswith(">"):
                i += 1
                continue
            if stripped.startswith("## Παράρτημα Δ"):
                out.extend([
                    "## Παράρτημα Δ — Ερευνητική εφαρμογή και όριο παρουσίασης",
                    "",
                    "Δεν ενσωματώνονται screenshots της εφαρμογής στην παρούσα review-ready έκδοση, επειδή δεν έχει παραχθεί provenance-registered application capture. Η απουσία screenshot δεν επηρεάζει την επιστημονική τεκμηρίωση: όλα τα quantitative αποτελέσματα προέρχονται από registered T-613 figures/tables και όχι από το UI.",
                    "",
                ])
                skip_app_section = True
                i += 1
                continue
            if skip_app_section:
                if stripped.startswith("## Παράρτημα Ε"):
                    skip_app_section = False
                else:
                    i += 1
                    continue
            if "Στο τελικό Word έγγραφο μπορεί να ενταχθεί appendix table" in stripped:
                skip_config_bullets = True
                i += 1
                continue
            if skip_config_bullets:
                if re.match(r"^[-*]\s+", stripped):
                    i += 1
                    continue
                skip_config_bullets = False
            if stripped == "Προτεινόμενα assets:":
                ids: list[str] = []
                j = i + 1
                while j < len(lines) and re.match(r"^\s*[-*]\s+", lines[j]):
                    ids.extend(builder.ASSET_RE.findall(lines[j]))
                    j += 1
                if ids:
                    out.append("Προτεινόμενη T-711 τοποθέτηση: " + ", ".join(ids))
                i = j
                continue
            if "Κύριο οπτικό asset:" in stripped:
                ids = builder.ASSET_RE.findall(stripped)
                if ids:
                    out.append("Προτεινόμενη T-711 τοποθέτηση: " + ", ".join(ids))
                i += 1
                continue
            if "T-711" in stripped or "Προτεινόμενα assets" in stripped:
                i += 1
                continue
            if stripped.startswith("Ο πίνακας πρέπει να παράγεται") or stripped.startswith("Το appendix πρέπει να συνεχίσει"):
                i += 1
                continue
            transformed = _appendix_transform(line)
            if transformed is None:
                i += 1
                continue
            line = transformed

        if stripped.startswith("> **Σημείωση drafting:**") or stripped.startswith("> Το T-710"):
            i += 1
            continue
        out.append(line)
        i += 1
    return "\n".join(out)


def _render_markdown(doc: Document, md: str, cmap: dict[str, int], assets: dict[str, dict], inserted: list[str], inserted_rq_tables: set[str], start_new_page: bool = True):
    global _current_render_mode
    if re.search(r"^#\s+Παραρτήματα", md, re.MULTILINE):
        mode = "appendix"
    elif "GLOSSARY" in md.upper() or "Γλωσσ" in md or "Draft authority for the alphabetical" in md:
        mode = "glossary"
    elif re.search(r"^#\s+Κεφάλαιο", md, re.MULTILINE):
        mode = "main"
    else:
        mode = "front"
    previous = _current_render_mode
    _current_render_mode = mode
    cleaned = _preprocess_markdown(md, mode)
    # Heading 1 already has page_break_before; avoid a second explicit break that can
    # materialize as a blank page in headless Word/LibreOffice rendering.
    effective_new_page = start_new_page and not cleaned.lstrip().startswith("#")
    try:
        return _original_render_markdown(doc, cleaned, cmap, assets, inserted, inserted_rq_tables, effective_new_page)
    finally:
        _current_render_mode = previous


def _analysis_identity(source_id: str) -> str:
    path = builder.BIB / "analyses" / f"{source_id}.md"
    if not path.exists():
        raise FileNotFoundError(f"citation-ready analysis missing for {source_id}")
    md = path.read_text(encoding="utf-8")
    marker = "## Bibliographic identity"
    if marker not in md:
        raise ValueError(f"verified bibliographic identity missing for {source_id}")
    tail = md.split(marker, 1)[1].lstrip()
    paragraph_lines: list[str] = []
    for line in tail.splitlines():
        stripped = line.strip()
        if not stripped:
            if paragraph_lines:
                break
            continue
        if stripped.startswith("#") or stripped.startswith("-"):
            if paragraph_lines:
                break
            continue
        paragraph_lines.append(stripped)
    identity = " ".join(paragraph_lines)
    identity = re.sub(r"\*\*([^*]+)\*\*", r"\1", identity)
    identity = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", identity)
    identity = identity.replace("`", "").strip()
    if len(identity) < 20 or identity.lower().startswith("http") or "https---" in identity or "http---" in identity:
        raise ValueError(f"unsafe bibliographic identity for {source_id}: {identity!r}")
    if not identity.endswith("."):
        identity += "."
    return identity


def _add_references(doc: Document, cmap: dict[str, int], bib: dict[str, dict[str, str]]):
    doc.add_page_break()
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Βιβλιογραφία")
    for source_id, number in sorted(cmap.items(), key=lambda item: item[1]):
        identity = _analysis_identity(source_id)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.first_line_indent = Cm(-0.75)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(f"[{number}] {identity}")
        builder.set_run_font(run, size=10)


def _multiline_field(paragraph, instruction: str, lines: list[str]):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend([begin, instr, separate])
    for index, line in enumerate(lines or [""]):
        text_node = OxmlElement("w:t")
        text_node.set(qn("xml:space"), "preserve")
        text_node.text = line
        run._r.append(text_node)
        if index < len(lines) - 1:
            run._r.append(OxmlElement("w:br"))
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)
    builder.set_run_font(run, size=10)


def _toc_cache_lines() -> list[str]:
    lines: list[str] = []
    for path in builder.MANUSCRIPT_FILES + [builder.GLOSSARY_FILE, builder.APPENDIX_FILE]:
        for raw in path.read_text(encoding="utf-8").splitlines():
            match = re.match(r"^(#{1,3})\s+(.+)$", raw.strip())
            if not match:
                continue
            level = len(match.group(1))
            title = builder.strip_markdown_inline(match.group(2))
            title = re.sub(r"\s+—\s+Draft.*$", "", title)
            if "T-711 appendix assembly rules" in title:
                continue
            lines.append("    " * (level - 1) + title)
    lines.append("Βιβλιογραφία")
    return lines


def _mentioned_figure_cache_lines() -> list[str]:
    assets = builder.load_assets()
    seen: set[str] = set()
    result: list[str] = []
    source = "\n".join(path.read_text(encoding="utf-8") for path in builder.MANUSCRIPT_FILES + [builder.APPENDIX_FILE])
    for asset_id in builder.ASSET_RE.findall(source):
        if asset_id in seen:
            continue
        seen.add(asset_id)
        asset = assets.get(asset_id)
        if not asset or asset.get("kind") != "figure":
            continue
        if not set(asset.get("intended_use", [])) & {"main-thesis", "appendix"}:
            continue
        result.append(asset.get("caption") or asset_id)
    return result


def _markdown_table_cache_lines() -> list[str]:
    result: list[str] = []
    for path in builder.MANUSCRIPT_FILES:
        current = path.stem
        lines = path.read_text(encoding="utf-8").splitlines()
        for index, raw in enumerate(lines):
            stripped = raw.strip()
            if stripped.startswith("#"):
                current = builder.strip_markdown_inline(stripped.lstrip("#").strip())
            if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
                label = f"Σύνοψη για την ενότητα «{current}»"
                if label not in result:
                    result.append(label)
    return result


def _clean_title_page(doc: Document, title: str, subtitle: str, english: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Cm(4)
    run = paragraph.add_run(title)
    builder.set_run_font(run, size=18, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    builder.set_run_font(run, size=14, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Cm(4)
    label = "Student: [to be completed from the official form]" if english else "Ονοματεπώνυμο φοιτητή: [να συμπληρωθεί από το επίσημο έντυπο]"
    run = paragraph.add_run(label)
    builder.set_run_font(run, size=11)
    doc.add_page_break()


def _add_front_matter(doc: Document, cmap: dict[str, int]):
    _clean_title_page(doc, "Σύγκριση και Αξιολόγηση Ανθεκτικών Πρακτόρων Τεχνητής Νοημοσύνης σε Περιβάλλοντα με Αβεβαιότητα", "Διπλωματική Εργασία")
    _clean_title_page(doc, "Comparison and Evaluation of Resilient AI Agents in Uncertain Environments", "Diploma Thesis", english=True)

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Δηλώσεις πνευματικών δικαιωμάτων και λογοκλοπής")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("[Θέση για την ακριβή επίσημη δήλωση πριν από την τελική υποβολή]")
    builder.set_run_font(run, size=11, italic=True)
    doc.add_page_break()

    greek, greek_keywords, english, english_keywords = builder.extract_front_sections()
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Περίληψη")
    _render_markdown(doc, greek, cmap, {}, [], set(), start_new_page=False)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Λέξεις-κλειδιά: ")
    builder.set_run_font(run, size=11, bold=True)
    run = paragraph.add_run(greek_keywords)
    builder.set_run_font(run, size=11)
    doc.add_page_break()

    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Abstract")
    _render_markdown(doc, english, cmap, {}, [], set(), start_new_page=False)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Keywords: ")
    builder.set_run_font(run, size=11, bold=True)
    run = paragraph.add_run(english_keywords)
    builder.set_run_font(run, size=11)
    doc.add_page_break()

    cached_fields = [
        ("Πίνακας Περιεχομένων", ' TOC \\o "1-3" \\h \\z \\u ', _toc_cache_lines()),
        ("Κατάλογος Σχημάτων", ' TOC \\h \\z \\c "Σχήμα" ', _mentioned_figure_cache_lines()),
        ("Κατάλογος Πινάκων", ' TOC \\h \\z \\c "Πίνακας" ', _markdown_table_cache_lines()),
    ]
    for title, instruction, cache in cached_fields:
        heading = doc.add_paragraph(style="Heading 1")
        heading.add_run(title)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.15
        _multiline_field(paragraph, instruction, cache)
        doc.add_page_break()


def _remove_release_note(output: Path):
    doc = Document(output)
    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Σημείωση έκδοσης review-ready":
            target = paragraph
            break
    if target is None:
        return
    body = doc._element.body
    children = list(body)
    target_index = children.index(target._p)
    if target_index > 0:
        previous = children[target_index - 1]
        if previous.tag == qn("w:p") and not "".join(previous.itertext()).strip():
            body.remove(previous)
            target_index -= 1
            children = list(body)
    for child in list(body)[target_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    doc.save(output)


def _enhanced_qa(output: Path, qa_output: Path):
    existing = json.loads(qa_output.read_text(encoding="utf-8"))
    doc = Document(output)
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    unresolved_src = sorted(set(builder.SRC_RE.findall(full_text)))
    internal_patterns = [
        r"Σημείωση drafting",
        r"Προτεινόμενη T-711",
        r"Προτεινόμενα assets",
        r"Draft authority",
        r"Draft περιεχομένου",
        r"T-711 appendix",
        r"authoritative metadata",
        r"Ενημερώστε το πεδίο",
        r"Σημείωση έκδοσης review-ready",
        r"\[ΕΚΚΡΕΜΕΙ",
    ]
    internal_hits = [pattern for pattern in internal_patterns if re.search(pattern, full_text, re.IGNORECASE)]
    raw_latex_patterns = [r"\\\(", r"\\\)", r"\\\[", r"\\\]", r"\\mathcal", r"\\gamma", r"\\pi", r"\\leftarrow", r"[_^]\{"]
    latex_hits = [pattern for pattern in raw_latex_patterns if re.search(pattern, full_text)]
    malformed_reference = bool(re.search(r"\[\d+\].*(?:https---|http---|“https?[:\-])", full_text))
    references = [paragraph.text for paragraph in doc.paragraphs if re.match(r"^\[\d+\]\s", paragraph.text.strip())]
    expected_reference_count = existing.get("citation_count", 0)
    status = "pass" if not unresolved_src and not internal_hits and not latex_hits and not malformed_reference and len(references) == expected_reference_count else "fail"
    existing.update({
        "schema_version": 2,
        "status": status,
        "output_sha256": builder.sha256(output),
        "unresolved_src_ids": unresolved_src,
        "contains_drafting_notes": bool(internal_hits),
        "internal_handoff_patterns": internal_hits,
        "contains_raw_latex": bool(latex_hits),
        "raw_latex_patterns": latex_hits,
        "malformed_reference_identity": malformed_reference,
        "verified_reference_identity_count": len(references),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "visual_qa_required": True,
    })
    qa_output.write_text(json.dumps(existing, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if status != "pass":
        raise RuntimeError(f"T-711 enhanced structural QA failed: {existing}")


def _build(output: Path, qa_output: Path):
    _caption_counters["Σχήμα"] = 0
    _caption_counters["Πίνακας"] = 0
    _original_build(output, qa_output)
    _remove_release_note(output)
    _enhanced_qa(output, qa_output)


# Activate the deterministic review-ready overrides.
builder.add_inline = _add_inline
builder.add_equation = _add_equation
builder.add_list_paragraph = _add_list_paragraph
builder.add_seq_caption = _add_seq_caption
builder.add_figure = _add_figure
builder.add_asset_table = _add_asset_table
builder.add_markdown_table = _add_markdown_table
builder.insert_asset_ids = _insert_asset_ids
builder.render_markdown = _render_markdown
builder.add_references = _add_references
builder.add_front_matter = _add_front_matter
builder.build = _build


if __name__ == "__main__":
    builder.main()
