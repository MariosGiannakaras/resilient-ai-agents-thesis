from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import hashlib, zipfile, json, re

import os

SRC = Path(os.environ.get('T716_SOURCE_DOCX', 'thesis/archive/T714_run66_full_review_ready.docx'))
OUT = Path(os.environ.get('T716_OUTPUT_DOCX', 'artifacts/t716/T716_stage3_full_content_review_ready.docx'))
QA = Path(os.environ.get('T716_QA_JSON', 'artifacts/t716/T716_stage3_qa-report.json'))
OUT.parent.mkdir(parents=True, exist_ok=True)
QA.parent.mkdir(parents=True, exist_ok=True)


def sha256(path: Path) -> str:
    h=hashlib.sha256()
    with path.open('rb') as f:
        for b in iter(lambda:f.read(1024*1024), b''):
            h.update(b)
    return h.hexdigest()


def media_hashes(path: Path):
    with zipfile.ZipFile(path) as z:
        return {n:hashlib.sha256(z.read(n)).hexdigest() for n in sorted(z.namelist()) if n.startswith('word/media/')}


def find_exact(doc, text):
    ms=[p for p in doc.paragraphs if p.text.strip()==text]
    if len(ms)!=1:
        raise RuntimeError(f'anchor {text!r}: {len(ms)} matches')
    return ms[0]


def insert_before(doc, anchor_text, paragraphs):
    anchor=find_exact(doc, anchor_text)
    for style,text in paragraphs:
        p=anchor.insert_paragraph_before(text, style=style)
        p.paragraph_format.keep_together=False


def add_ref_before(doc, anchor_text, text):
    p=find_exact(doc, anchor_text).insert_paragraph_before(text, style='Normal')
    p.paragraph_format.keep_together=False
    return p


def words(doc, upto=None):
    ps=doc.paragraphs if upto is None else doc.paragraphs[:upto]
    return sum(len(re.findall(r"\b[\wΆ-ώ]+\b",p.text,flags=re.UNICODE)) for p in ps)

before_media=media_hashes(SRC)
doc=Document(SRC)
base_paragraphs=len(doc.paragraphs)
base_words=words(doc)

# --- Chapter 1: sharper problem framing and controlled-testbed rationale ---
insert_before(doc, '1.2 Πρόβλημα και ερευνητικό πλαίσιο', [
('Normal', 'Η διάκριση ανάμεσα σε ονομαστική μάθηση και προσαρμογή μετά από μεταβολή είναι και θεωρητική. Η κλασική Q-Learning διαθέτει ισχυρό αποτέλεσμα σύγκλισης υπό τις δηλωμένες υποθέσεις του σταθερού, διακριτού προβλήματος και επαρκούς επαναληπτικής δειγματοληψίας, όμως το αποτέλεσμα αυτό δεν αποτελεί εγγύηση πεπερασμένου χρόνου ανάκαμψης όταν αλλάξει χωρίς προειδοποίηση η σχέση κατάστασης–ενέργειας–συνέπειας [18]. Η βιβλιογραφία για dynamically varying environments χρησιμοποιεί ακριβώς αυτή τη διαφορά ως όριο: όταν παραβιάζεται η stationarity assumption, η ικανότητα tracking ή adaptation πρέπει να μετρηθεί στο νέο καθεστώς και δεν μπορεί να συναχθεί από stationary convergence theory [21].'),
('Normal', 'Επιπλέον, robustness, zero-shot generalisation, online adaptation και recovery δεν είναι εναλλάξιμοι όροι. Σε αυστηρό zero-shot regime δεν επιτρέπονται νέες learning updates στο test instance, ενώ στην online adaptation ο agent συνεχίζει να ενημερώνει policy, value function ή model μετά τη μεταβολή [22]. Η temporal resilience literature προσθέτει μία ακόμη διάσταση: η άμεση πτώση, η πορεία αποκατάστασης και η τελική επίδοση μπορεί να περιγράφουν διαφορετικές ιδιότητες της ίδιας trajectory [23], [24]. Η παρούσα εργασία διατηρεί αυτές τις διακρίσεις μέσω των τριών χωριστών RQ αντί να τις συμπιέζει σε έναν ενιαίο δείκτη.'),
])
insert_before(doc, '1.3 Σκοπός και στόχοι', [
('Normal', 'Η επιλογή ενός μικρού GridWorld υπηρετεί αυτόν τον διαχωρισμό και όχι την απλοποίηση για την ίδια την απλοποίηση. Η βιβλιογραφία generalisation τονίζει ότι ένα benchmark ορίζεται από τον συνδυασμό environment, παραγόντων μεταβολής, train/test protocol και διαθέσιμου interaction budget, όχι μόνο από το όνομα του simulator [22]. Παράλληλα, το NovGrid δείχνει ότι ένα GridWorld μπορεί να μεταβάλει ελεγχόμενα μηχανισμούς ή action effects διατηρώντας συμβατή διεπαφή και να αξιολογεί χωριστά degradation, adaptation και recovery [23]. Στην παρούσα μελέτη η ίδια λογική εφαρμόζεται με project-owned disturbance semantics, frozen seeds και matched branches, ώστε οι αιτίες των διαφορών να παραμένουν όσο γίνεται επιθεωρήσιμες.'),
])

# --- Chapter 2: foundational, non-stationarity, uncertainty, resilience, safety ---
insert_before(doc, '2.4 Deep Q-Network', [
('Normal', 'Η πρωτογενής εργασία των Watkins και Dayan παρέχει το θεμελιώδες όριο για την ερμηνεία της Q-Learning. Ο one-step bootstrap target χρησιμοποιεί τη μέγιστη εκτιμώμενη αξία της επόμενης κατάστασης και επομένως η μέθοδος είναι off-policy ως προς την exploratory behavior policy [18]. Η SARSA, αντίθετα, χρησιμοποιεί την αξία της επόμενης ενέργειας που πράγματι επιλέγεται από την behavior policy. Η μηχανιστική αυτή διαφορά είναι σημαντική για το πείραμα, αλλά δεν συνεπάγεται εκ των προτέρων ότι μία από τις δύο μεθόδους είναι «πιο ανθεκτική». Αυτό αποτελεί εμπειρικό ερώτημα υπό το frozen protocol.'),
('Normal', 'Το κλασικό convergence result της Q-Learning πρέπει επίσης να διαβάζεται μέσα στις παραδοχές του. Η πιθανότητα-ένα σύγκλιση αφορά επαναλαμβανόμενη δειγματοληψία σε σταθερό Markovian domain και δεν προβλέπει πόσες πραγματικές αλληλεπιδράσεις απαιτούνται για να εντοπιστεί και να απορροφηθεί μία κρυφή persistent action remap [18]. Για αυτό η εργασία δεν χρησιμοποιεί θεωρητική σύγκλιση ως proxy για resilience· μετρά άμεσα την post-change loss και, χωριστά, το stable recovery.'),
])
insert_before(doc, '2.7 Μη στασιμότητα και δυναμικές μεταβολές', [
('Normal', 'Η πρωτογενής διατύπωση της Dyna από τον Sutton ενοποιεί direct reinforcement learning, εκμάθηση forward model και planning από model-generated experience σε μία κοινή incremental αρχιτεκτονική [19]. Οι planning updates δεν αποτελούν νέες πραγματικές αλληλεπιδράσεις· είναι εσωτερικές ενημερώσεις που αξιοποιούν το μαθημένο μοντέλο. Η διάκριση αυτή αιτιολογεί γιατί η παρούσα μελέτη χρησιμοποιεί actual environment interactions ως κοινό fairness currency, επιτρέποντας παράλληλα στη Dyna-Q+ το method-native planning που αποτελεί μέρος του αλγορίθμου.'),
('Normal', 'Στα changing-world experiments της ίδιας εργασίας, η recency-driven διερεύνηση ενθαρρύνει την επαναδοκιμή ενεργειών των οποίων οι συνέπειες μπορεί να έχουν αλλάξει, ενώ ένα μοντέλο μπορεί να παραμένει προσωρινά stale μέχρι να επαναπαρατηρηθούν τα επηρεασμένα state–action pairs [19]. Το ιστορικό αυτό evidence στηρίζει τη μηχανιστική υπόθεση πίσω από τη Dyna-Q+, όχι μια πρόβλεψη ότι θα ανακάμψει ταχύτερα σε κάθε αλλαγή. Η μόνιμη ανααντιστοίχιση ενεργειών της παρούσας εργασίας διαφέρει από τα συγκεκριμένα maze changes και αξιολογείται ανεξάρτητα.'),
])
insert_before(doc, '2.8 Αβεβαιότητα ενεργειών, παρατηρήσεων και κρυφή μεταβολή', [
('Normal', 'Η μη στασιμότητα είναι ευρύτερη από μία απλή αλλαγή μετάβασης. Η peer-reviewed επισκόπηση των Khetarpal et al. οργανώνει το continual RL με βάση το scope και τον driver της non-stationarity και επισημαίνει ότι μπορεί να μεταβάλλονται dynamics, rewards, observations ή άλλες συνιστώσες της αλληλεπίδρασης, με ή χωρίς καθαρά task boundaries [20]. Η εστιασμένη survey του Padakandla συμπληρώνει αυτή την εικόνα τονίζοντας ότι οι stationary convergence formulations δεν μεταφέρονται αυτομάτως σε dynamically varying environments και ότι διαφορετικές οικογένειες αλλαγής απαιτούν διαφορετικούς μηχανισμούς αντίδρασης [21].'),
('Normal', 'Αυτό έχει άμεση συνέπεια στη σχεδίαση της παρούσας μελέτης. Μία abrupt persistent action remap, ένα στοχαστικό no-op actuator failure και μία corruption της delivered observation δεν θεωρούνται τρεις βαθμοί της ίδιας μεταβλητής, αλλά τρεις διαφορετικοί information/dynamics mechanisms. Η ανάλυση τα διατηρεί χωριστά ώστε ένα θετικό αποτέλεσμα σε μία οικογένεια να μην μετατραπεί σε γενικό ισχυρισμό για «non-stationary RL». Η πρακτική αυτή ακολουθεί την αρχή ότι η φύση και ο βαθμός της non-stationarity πρέπει να ελέγχονται και να δηλώνονται ρητά [20], [21].'),
('Normal', 'Η continual-learning βιβλιογραφία επιβάλλει επίσης προσοχή στη διάκριση μεταξύ retention και adaptation. Η διατήρηση παλαιότερης γνώσης και η επαρκής plasticity για νέα εμπειρία δημιουργούν tension που δεν λύνεται απλώς με το να παραμένει ενεργός ο optimizer [20]. Πρόσφατη model-based εργασία στο ICML 2025 εξετάζει διαφορετικό αλλά συναφές setting: διαδοχικά tasks με κοινή world dynamics, online world model και MPC planning, με κύριο πρόβλημα το catastrophic forgetting [30]. Το setting αυτό δεν ισοδυναμεί με Dyna-Q+ ή με hidden action remap, αλλά αποτελεί σύγχρονο παράδειγμα του πόσο καθοριστικές είναι οι assumptions για το τι ακριβώς «συνεχίζεται» και τι αλλάζει.'),
])
insert_before(doc, '2.9 Ανθεκτικότητα, προσαρμογή και ανάκαμψη', [
('Normal', 'Η αβεβαιότητα ενεργειών πρέπει να διαχωρίζει τουλάχιστον την τυχαία ή adversarial αντικατάσταση action από τη συνεχή perturbation και από τη μόνιμη αλλαγή της σημασιολογίας των actions. Οι Tessler et al. διατυπώνουν probabilistic action replacement και noisy-action models και δείχνουν ότι το robustness–performance trade-off εξαρτάται από το disturbance model και τη severity [26]. Η thesis condition action-failure-0.15 είναι στενότερη: πρόκειται για ανεξάρτητο stochastic no-op και όχι για adversary που επιλέγει worst-case action. Οι persistent remaps είναι ακόμη διαφορετικό πρόβλημα, επειδή αλλάζουν συστηματικά την αντιστοίχιση intended προς executed action.'),
('Normal', 'Αντίστοιχα, observation corruption δεν είναι συνώνυμη με αλλαγή dynamics. Εργασία σε POMDPs με incomplete/noisy observations χρησιμοποιεί belief inference και learned transition components για να αντιμετωπίσει information loss [27], ενώ η bounded-robustness προσέγγιση των Jarne Ornia et al. μοντελοποιεί άγνωστο observation kernel και θέτει ρητό trade-off μεταξύ nominal utility και robustness [28]. Η παρούσα εργασία δεν υλοποιεί belief state, filtering ή robust-policy training. Η συνθήκη 0,05 λειτουργεί ως controlled corruption της delivered observation ώστε να εξεταστεί η συμπεριφορά των ίδιων frozen configurations υπό διαφορετικό information quality.'),
('Normal', 'Η διάκριση από zero-shot generalisation είναι επίσης ουσιώδης. Στη ZSG βιβλιογραφία η test policy αξιολογείται χωρίς πρόσθετο learning από τα test instances, ενώ online updates αποτελούν διαφορετικό regime [22]. Οι Frozen branches της παρούσας μελέτης προσεγγίζουν αυτή τη no-update λογική ως προς το deployment, ενώ οι Adaptive branches επιτρέπουν ordinary method-native learning. Το matched FN/FD/AN/AD design δεν πρέπει συνεπώς να περιγράφεται ως standard zero-shot benchmark, αλλά ως ελεγχόμενη σύγκριση no-update robustness και online adaptation.'),
])
insert_before(doc, '2.10 Συνεχής προσαρμογή, replay και οργάνωση μοντέλου', [
('Normal', 'Η ανθεκτικότητα αντιμετωπίζεται εδώ ως χρονική διαδικασία και όχι ως ετικέτα για μία υψηλή τελική απόδοση. Το NovGrid διαχωρίζει την immediate response to novelty, την asymptotic adaptive performance και την adaptive efficiency, δείχνοντας ότι διαφορετικά σημεία της post-change curve απαντούν σε διαφορετικά ερωτήματα [23]. Η εργασία για cooperative resilience των Chacon-Chamorro et al. χρησιμοποιεί reference και performance curves για να διαχωρίσει failure και recovery profiles [24]. Παρότι το domain της είναι multi-agent, η εννοιολογική διάκριση πτώσης και αποκατάστασης υποστηρίζει τη χρήση ξεχωριστών RQ2 και RQ3.'),
('Normal', 'Πρόσφατη peer-reviewed εργασία για autonomous cyber defense ενισχύει το ίδιο methodological point από διαφορετικό domain: πολύ coarse χρονική aggregation μπορεί να κρύψει πληροφορία για το recovery process, ενώ υπερβολικά fine resolution μπορεί να δυσκολεύει την ερμηνεία [25]. Η δική μας επιλογή fixed 32-interaction windows δεν προέρχεται από τη cyber metric και δεν υιοθετεί τα domain-specific weights ή smoothing της. Η πηγή χρησιμοποιείται μόνο για να αιτιολογήσει γιατί η temporal granularity πρέπει να δηλώνεται ρητά και γιατί το whole-horizon score δεν αρκεί για recovery claim.'),
('Normal', 'Τέλος, resilience και safety παραμένουν διακριτές έννοιες. Η safe-RL βιβλιογραφία συνήθως εισάγει explicit costs, constraints ή risk semantics πέρα από το task reward και τονίζει ότι safety κατά την εξερεύνηση μπορεί να διαφέρει από την τελική policy utility [29]. Η παρούσα μελέτη δεν έχει frozen safety-cost objective ούτε αποδεικνύει constraint satisfaction. Συνεπώς, η διατήρηση ή ανάκτηση task performance υπό disturbance δεν παρουσιάζεται ως safety guarantee· το safe continual adaptation αποτελεί ξεχωριστή κατεύθυνση μελλοντικής εργασίας.'),
])
insert_before(doc, '2.11 Εμπειρικός σχεδιασμός και δίκαιη σύγκριση RL', [
('Normal', 'Η αποθηκευμένη εμπειρία και το learned model δημιουργούν διαφορετικές μορφές ιστορικής εξάρτησης. Στην DQN, το replay buffer επαναχρησιμοποιεί παλαιότερες μεταβάσεις και η replay design μπορεί να αλλάξει αισθητά τη learning dynamics [11]. Σε non-stationary deployment, αυτό δημιουργεί εύλογο stale-history concern, αλλά η κατεύθυνση της επίδρασης δεν είναι δεδομένη: η διατήρηση παλιών samples μπορεί να παρεμβαίνει στην προσαρμογή, ενώ η επιθετική απόρριψή τους μπορεί να αυξήσει forgetting. Χωρίς ειδική replay ablation, η παρούσα εργασία αντιμετωπίζει αυτό το σημείο ως mechanism hypothesis και όχι ως αιτιώδη εξήγηση των DQN outcomes.'),
('Normal', 'Στα model-based systems το αντίστοιχο ζήτημα είναι τι γνώση αποθηκεύει το model και πώς χρησιμοποιείται. Η Dyna κάνει learning updates από model-generated experience [19], ενώ η πρόσφατη εργασία των Liu et al. σχεδιάζει actions με MPC πάνω σε online world model που επαναχρησιμοποιείται μεταξύ tasks με κοινή dynamics [30]. Πρόκειται για διαφορετικούς planning mechanisms. Η σύγκρισή τους είναι χρήσιμη επειδή δείχνει ότι ο όρος «model-based adaptation» δεν ορίζει μία ενιαία μέθοδο και ότι assumptions για task boundaries, reward changes, uncertainty και exploration πρέπει να διατηρούνται ορατές.'),
])
insert_before(doc, '2.13 Ερευνητικό κενό και θέση της παρούσας εργασίας', [
('Normal', 'Συνολικά, η σχετική βιβλιογραφία περιλαμβάνει εξειδικευμένες λύσεις για action robustness [26], observation robustness [27], [28], zero-shot generalisation [22], continual-learning retention [20], [30] και temporal resilience [23]–[25]. Η παρούσα εργασία δεν επιχειρεί να αναπαραγάγει αυτές τις specialized methods μέσα σε ένα ενιαίο leaderboard. Χρησιμοποιεί τις πηγές για να ορίσει τα σύνορα των εννοιών και να επιλέξει ελεγχόμενες disturbances, ενώ το empirical question παραμένει η συμπεριφορά πέντε συγκεκριμένων baseline mechanisms υπό κοινό protocol.'),
('Normal', 'Αυτή η τοποθέτηση αποφεύγει δύο αντίθετα σφάλματα. Το πρώτο θα ήταν να παρουσιαστούν standard RL algorithms ως state of the art για κάθε μορφή non-stationarity. Το δεύτερο θα ήταν να θεωρηθεί ότι ένα specialized robust ή continual method πρέπει εξ ορισμού να υπερέχει σε οποιοδήποτε disturbance. Οι υπάρχουσες εργασίες δείχνουν ισχυρή εξάρτηση από domain, information regime, perturbation family, architecture και objective [22], [26], [28]. Επομένως, η συνεισφορά της διπλωματικής είναι controlled evidence για έναν περιορισμένο μηχανιστικό χώρο και όχι καθολική κατάταξη.'),
])

# --- Chapter 3: T-715 scientific corrections integrated into fuller T-714 structure ---
insert_before(doc, '3.4 Φάση A: ανεξάρτητη ονομαστική μάθηση και σημεία ελέγχου', [
('Normal', 'Οι τελικές ρυθμίσεις δεν επιλέχθηκαν από τα final outcomes. Πριν ανοίξει το final reserve εκτελέστηκε προκαθορισμένο, ίσης ευκαιρίας tuning μόνο σε DEVELOPMENT evidence: έξι method-specific υποψήφιες ρυθμίσεις για καθεμία από τις πέντε μεθόδους, στις ίδιες τρεις tuning-only roots και στις ίδιες δύο development layouts, με κοινό budget 8.192 πραγματικών αλληλεπιδράσεων και κοινό πλέγμα probes. Το design παρήγαγε 6×3×2×5=180 tuning units. Οι random seeds δεν αποτέλεσαν tuning parameter.'),
('Normal', 'Η επιλογή έγινε χωριστά ανά μέθοδο με προκαθορισμένο μηχανικό κανόνα: πρώτα μεγιστοποιήθηκε η ισοβαρής ως προς roots και layouts τραπεζοειδής time-average success στα no-learning probes, έπειτα χρησιμοποιήθηκαν ως tie-breakers η τελική success, η time-average return και, μόνο σε πλήρη ισοβαθμία, το λεξικογραφικά μικρότερο config ID. Οι επιλεγμένες ρυθμίσεις ήταν q-c06, sarsa-c06, dqn-c05, ppo-c06 και dyna-c03. Το bounded search space ήταν ίσης έκτασης αλλά όχι εξαντλητικό, άρα η σύγκριση αφορά αυτές τις frozen επιλογές και όχι τη θεωρητικά μέγιστη δυνατή επίδοση κάθε algorithm family.'),
])
insert_before(doc, '3.7 RQ1: μεγέθη εκτίμησης ονομαστικής μάθησης', [
('Normal', 'Η operational σημασιολογία των disturbances είναι μέρος του frozen protocol. Στο swap-right-down ισχύει up→up, right→down, down→right, left→left, ενώ στο cycle-clockwise ισχύει up→right, right→down, down→left, left→up. Η αλλαγή εφαρμόζεται στο actually executed action mapping και παραμένει ενεργή στα επόμενα episodes. Ο agent δεν λαμβάνει change indicator ούτε executed-action truth.'),
('Normal', 'Στο action-failure-0.15 πραγματοποιείται ανεξάρτητη Bernoulli δοκιμή σε κάθε πραγματική αλληλεπίδραση. Σε αποτυχία, η intended action μετατρέπεται σε no-op: η ground-truth θέση παραμένει αμετάβλητη, collision=false και αποδίδεται η κανονική step reward −0,1, εκτός αν εφαρμόζεται άλλη terminal/reward rule. Η μετάβαση παραμένει κανονική εμπειρία και στους Adaptive branches τροφοδοτεί τον method-native learning mechanism.'),
('Normal', 'Στο observation-corruption-0.05 η δοκιμή γίνεται μετά την πραγματική μετάβαση. Όταν ενεργοποιηθεί, η delivered observation δειγματοληπτείται ομοιόμορφα από όλα τα μη-obstacle κελιά που δεν είναι η τρέχουσα ground-truth θέση. Goal και start δεν αποκλείονται ως κατηγορίες εφόσον είναι έγκυρα και δεν συμπίπτουν με την τρέχουσα πραγματική θέση. Η corruption δεν αλλάζει transition ή reward· αλλάζει μόνο την observation που λαμβάνει ο agent και, όταν η μάθηση είναι ενεργή, αυτή η observation χρησιμοποιείται στην κανονική ενημέρωση.'),
])
insert_before(doc, '3.10 Ανεξάρτητες επαναλήψεις, διατάξεις, προϋπολογισμοί και πειραματικός πίνακας', [
('Normal', 'Για πληρότητα, ο directed recovery gap γράφεται gₖ=Nₖ−Dₖ, όπου Nₖ είναι η equal-layout Adaptive-Nominal μέση reward ανά interaction στο παράθυρο k και Dₖ η αντίστοιχη Adaptive-Disturbed τιμή. Το window είναι in-tolerance όταν gₖ≤0,10· το κριτήριο δεν είναι |gₖ|≤0,10. Stable recovery απαιτεί δύο διαδοχικά in-tolerance windows. recovery_time είναι το endpoint του πρώτου window του πρώτου stable pair και confirmation_time το endpoint του δεύτερου. Αν το pair δεν εμφανιστεί μέχρι το 256, το observation είναι right-censored και recovery_time=null.'),
])
insert_before(doc, '3.11 Στατιστική ανάλυση και άμεσες συγκρίσεις', [
('Normal', 'Το πλήθος των 12 roots επιλέχθηκε πριν από το final experiment μέσω DEVELOPMENT precision sizing. Εξετάστηκαν candidate counts 12, 16, 20 και 24 και επιλέχθηκε το μικρότερο n που ικανοποιούσε το προδηλωμένο Student-t 95% half-width <0,20 τόσο για Phase-A AUC όσο και για Phase-B adaptation benefit. Το κριτήριο ικανοποιήθηκε στα 12 roots, όπου το μέγιστο observed sizing half-width ήταν 0,1428. Οι layouts δεν αυξάνουν το n επειδή μειώνονται ισοβαρώς μέσα σε κάθε root.'),
('Normal', 'Οι δύο final layouts έχουν generation seeds 57001 και 57002. Για root r_i, i=01,…,12, χρησιμοποιούνται χωριστά streams initialization=71000+i, exploration=72000+i, scenario=73000+i, environment=74000+i, action disturbance=75000+i και observation disturbance=76000+i. Η διάκριση αυτή αποτρέπει την ανεπιθύμητη επαναχρησιμοποίηση ενός ενιαίου RNG stream μεταξύ μηχανισμών και διατηρεί το pairing όπου αυτό απαιτείται από το protocol.'),
])

# --- Chapter 4: architecture explanation without changing project authority ---
insert_before(doc, '4.5 Checkpoints και ακριβής συνέχεια της μάθησης', [
('Normal', 'Κρίσιμη αρχιτεκτονική αρχή είναι ο information firewall μεταξύ agent και evaluator. Η εφαρμογή και το evidence layer μπορούν να γνωρίζουν ground-truth θέση, branch identity, disturbance configuration και telemetry για σκοπούς validation ή visualization, όμως αυτά τα δεδομένα δεν μετατρέπονται σε learning input. Ο agent ενημερώνεται μόνο από την προβλεπόμενη observation/reward/action interface. Έτσι η hidden-change υπόθεση δεν παραβιάζεται από το ίδιο το σύστημα παρακολούθησης.'),
])
insert_before(doc, '4.7 Πακέτα εκτέλεσης, manifests και ακεραιότητα', [
('Normal', 'Η αναπαραγωγιμότητα εξαρτάται επίσης από την απομόνωση των πηγών τυχαιότητας. Αν initialization, exploration, environment sampling και disturbance sampling κατανάλωναν μία κοινή mutable RNG ακολουθία, μία αθώα αλλαγή στον αριθμό random draws ενός component θα μπορούσε να μεταβάλει άσχετο component και να σπάσει το intended pairing. Για αυτό τα seed streams είναι χωριστά και προκαθορισμένα ανά root. Η αρχιτεκτονική καταγράφει τις ταυτότητες αυτές στα run bundles αντί να βασίζεται σε implicit global seed.'),
])
insert_before(doc, '4.16 Συνολική αρχιτεκτονική ροή', [
('Normal', 'Ο διαχωρισμός scientific core, Study orchestration, evidence validation, analysis και UI περιορίζει δύο κατηγορίες σφάλματος. Πρώτον, η παρουσίαση δεν έχει εξουσία να επαναϋπολογίζει estimands ή να αλλάζει thresholds. Δεύτερον, μία αποτυχία ή αλλαγή στο presentation layer δεν απαιτεί επανάληψη της επιστημονικής εκτέλεσης εφόσον τα frozen evidence artifacts παραμένουν έγκυρα. Η ίδια αρχή επιτρέπει το Word manuscript να αναφέρεται σε registered figures/tables ως immutable scientific media, ενώ η συγγραφή αλλάζει μόνο exposition και contextual interpretation.'),
('Normal', 'Η Study-first σχεδίαση λειτουργεί επομένως ως provenance boundary: immutable recipe → deterministic plan → Phase-A state → exact matched Phase-B branches → validated evidence → frozen analysis → registered thesis assets. Η PySide6 εφαρμογή είναι thin inspection/control surface πάνω από αυτή την αλυσίδα. Η διάκριση είναι ουσιώδης για τη διπλωματική, επειδή η αναπαραγωγιμότητα δεν στηρίζεται σε screenshots ή χειροκίνητες σημειώσεις αλλά σε versioned machine-readable identities.'),
])

# --- Chapter 6: deeper evidence-aware interpretation and validity boundaries ---
insert_before(doc, '6.3 RQ2: η αποτελεσματικότητα της προσαρμογής εξαρτάται από τη συνθήκη', [
('Normal', 'Η διάκριση τελικής επίδοσης και learning efficiency συμφωνεί με τη γενικότερη αρχή ότι διαφορετικά σημεία ή summaries μιας learning curve απαντούν σε διαφορετικές ερωτήσεις. Σε non-stationary evaluation η ίδια λογική γίνεται ακόμη ισχυρότερη: η performance κατά την adaptation phase μπορεί να διαφέρει από την τελική post-change policy [21], [23]. Για αυτό η υψηλή nominal time-average απόδοση της Dyna-Q+ δεν χρησιμοποιείται ως ένδειξη ότι θα έχει υποχρεωτικά μικρότερο adaptation loss ή ταχύτερο recovery.'),
])
insert_before(doc, '6.5 RQ3: συχνότητα, χρόνος και λογοκρισία της ανάκαμψης', [
('Normal', 'Η πρωτογενής Dyna εργασία προσφέρει μία εύλογη αλλά περιορισμένη μηχανιστική ερμηνεία: planning πάνω σε learned model μπορεί να πολλαπλασιάζει την αξιοποίηση κάθε πραγματικής transition, αλλά μετά από change τμήματα του model μπορούν να παραμένουν stale έως ότου συλλεχθεί νέα πραγματική εμπειρία [19]. Η Dyna-Q+ προσθέτει recency-driven re-exploration ακριβώς για να αυξήσει την πιθανότητα επανεξέτασης παλιών state–action pairs. Το αποτέλεσμα της παρούσας μελέτης δείχνει ότι αυτοί οι μηχανισμοί δεν μεταφράστηκαν σε καθολικά ταχύτερη recovery στο συγκεκριμένο protocol· δεν αποδεικνύει ότι το model-based planning είναι γενικά επιζήμιο.'),
('Normal', 'Η σύγκριση με σύγχρονο model-based continual RL απαιτεί ακόμη μεγαλύτερη προσοχή. Οι Liu et al. διατηρούν reusable online world dynamics και σχεδιάζουν με MPC σε sequence tasks με κοινή dynamics και μεταβαλλόμενα rewards [30]. Η παρούσα action-remap αλλάζει τη σχέση intended προς executed action και το Dyna-Q+ κάνει value/planning backups αντί για direct MPC. Επομένως, το recent result είναι χρήσιμο ως απόδειξη ότι η οργάνωση και επαναχρησιμοποίηση model knowledge είναι ενεργό ερευνητικό θέμα, όχι ως cross-domain predictor των thesis outcomes.'),
])
insert_before(doc, '6.8 Απειλές προς την εσωτερική εγκυρότητα', [
('Normal', 'Η RQ3 operationalization είναι σκόπιμα αυστηρή ως προς το τι ονομάζεται observed recovery. Η resilience literature δείχνει ότι degradation, recovery trajectory και τελική κατάσταση μπορεί να αποκλίνουν [24], ενώ το NovGrid διαχωρίζει επίσης immediate resilience, adaptive efficiency και asymptotic performance [23]. Επιπλέον, το temporal-granularity analysis των Cadet et al. δείχνει ότι whole-horizon aggregation μπορεί να κρύψει recovery information [25]. Αυτά δεν καθορίζουν το δικό μας threshold ή window size, αλλά στηρίζουν την επιλογή να διατηρούνται χωριστά recovery incidence, observed time και restricted delay.'),
('Normal', 'Η λογοκρισία είναι ουσιαστικό μέρος αυτής της διάκρισης. Ένα root που δεν ικανοποιεί το stable criterion έως το 256 δεν παρέχει observed event time· παρέχει μόνο την πληροφορία ότι το event δεν παρατηρήθηκε μέσα στον διαθέσιμο horizon. Η χρήση του 256 ως restricted-horizon quantity επιτρέπει bounded comparison χωρίς να μετατρέπει τη μη ανάκαμψη σε ψευδή παρατήρηση. Παράλληλα, η sensitivity στα 0,05/0,10/0,20 δείχνει πόσο εξαρτάται η incidence από τον operational ορισμό, χωρίς post-hoc επιλογή threshold.'),
])
insert_before(doc, '6.12 Εγκυρότητα αναπαραγωγιμότητας', [
('Normal', 'Η εξωτερική εγκυρότητα πρέπει να διαβάζεται μαζί με το είδος benchmark. Η ZSG survey επισημαίνει ότι environment και evaluation protocol είναι αδιαχώριστα και ότι held-out seeds από μόνα τους δεν αρκούν για ισχυρούς ισχυρισμούς generalisation όταν οι παράγοντες μεταβολής δεν είναι ελεγχόμενοι [22]. Το NovGrid αντίστοιχα χρησιμοποιεί controllable novelty σε μικρό GridWorld για να απομονώνει post-change behavior [23]. Η παρούσα μελέτη ακολουθεί αυτή τη controlled-factor φιλοσοφία, αλλά παραμένει μία συγκεκριμένη 7×7 task family με δύο final layouts και δεν τεκμηριώνει transfer σε continuous control ή robotics.'),
('Normal', 'Επίσης, specialized robustness methods δεν προσφέρουν καθολικό comparator χωρίς αλλαγή του research question. Η action-robust εργασία των Tessler et al. χρησιμοποιεί adversarial action models και δείχνει sensitivity στη severity και στο domain [26], ενώ η bounded observation robustness των Jarne Ornia et al. βελτιστοποιεί explicit lexicographic robustness objective με nominal-utility tolerance [28]. Η απουσία αυτών των interventions από το thesis δεν είναι ένδειξη ότι είναι ασήμαντες· σημαίνει ότι το πείραμα εστιάζει σε standard method-native adaptation υπό κοινό information contract.'),
])
insert_before(doc, '6.14 Κεντρικό συμπέρασμα της Συζήτησης', [
('Normal', 'Ένα ακόμη όριο είναι η ασφάλεια. Η safe-RL βιβλιογραφία διαχωρίζει task return από costs, constraints και risk semantics και επισημαίνει ότι μία policy μπορεί να έχει καλή utility χωρίς να ικανοποιεί safety requirement [29]. Η παρούσα μελέτη δεν ορίζει safety cost ούτε safe-exploration constraint· επομένως τα resilience findings δεν πρέπει να διατυπωθούν ως αποδείξεις ασφαλούς deployment. Η συμβολή περιορίζεται στην εμπειρική συμπεριφορά task performance και recovery υπό συγκεκριμένες disturbances.'),
])

# --- Chapter 7: evidence-bounded future directions ---
insert_before(doc, '7.8 Τελικό συμπέρασμα', [
('Normal', 'Μία φυσική επέκταση είναι να διαχωριστούν πειραματικά τα regimes no-update robustness, hidden online adaptation και explicit context/change detection. Η ZSG βιβλιογραφία παρέχει καθαρό protocol boundary για το πρώτο [22], ενώ το continual RL καλύπτει πλουσιότερες περιπτώσεις retention και μεταβολής χωρίς καθαρά task boundaries [20]. Ένα επόμενο protocol θα μπορούσε να προσθέσει detector ή context inference ως ξεχωριστό treatment, χωρίς να αλλοιώσει αναδρομικά τα αποτελέσματα της παρούσας μελέτης.'),
('Normal', 'Για model-based agents, χρήσιμη κατεύθυνση είναι η ελεγχόμενη μελέτη model freshness, planning strategy και history management. Η Dyna foundation δείχνει τον ρόλο του recency-driven exploration σε changing worlds [19], ενώ το recent online-world-model work δείχνει διαφορετικό design point όπου reusable dynamics χρησιμοποιείται άμεσα για MPC planning [30]. Συγκριτικές ablations θα μπορούσαν να ξεχωρίσουν planning quantity, model reset/recency, uncertainty handling και replay/model retention αντί να αποδίδεται κάθε post-change διαφορά γενικά στο «model-based RL».'),
('Normal', 'Τέλος, η μετάβαση από resilience σε safe continual adaptation απαιτεί νέο επιστημονικό συμβόλαιο. Θα χρειάζονταν explicit cost/constraint variables, safety evaluation τόσο κατά την adaptation transient όσο και μετά από αυτή, και πιθανώς methods σχεδιασμένες για constrained learning [29]. Αυτή η επέκταση είναι σημαντική ακριβώς επειδή η σημερινή εργασία δεν μετρά αυτά τα μεγέθη και δεν πρέπει να τα υπονοεί μέσω του task return.'),
])

# --- Bibliography additions [18]-[30] ---
refs = [
'[18] Christopher J. C. H. Watkins and Peter Dayan, “Q-learning,” Machine Learning, vol. 8, pp. 279–292, 1992.',
'[19] Richard S. Sutton, “Integrated Modeling and Control Based on Reinforcement Learning and Dynamic Programming,” in Advances in Neural Information Processing Systems 3, pp. 471–478, 1990.',
'[20] Khimya Khetarpal, Matthew Riemer, Irina Rish, and Doina Precup, “Towards Continual Reinforcement Learning: A Review and Perspectives,” Journal of Artificial Intelligence Research, vol. 75, pp. 1401–1476, 2022, doi: 10.1613/JAIR.1.13673.',
'[21] Sindhu Padakandla, “A Survey of Reinforcement Learning Algorithms for Dynamically Varying Environments,” ACM Computing Surveys, vol. 54, no. 6, 2021, doi: 10.1145/3459991.',
'[22] Robert Kirk, Amy Zhang, Edward Grefenstette, and Tim Rocktäschel, “A Survey of Zero-shot Generalisation in Deep Reinforcement Learning,” Journal of Artificial Intelligence Research, vol. 76, pp. 201–264, 2023.',
'[23] Jonathan Balloch, Zhiyu Lin, Mustafa Hussain, Aarun Srinivas, Robert Wright, Xiangyu Peng, Julia Kim, and Mark Riedl, “NovGrid: A Flexible Grid World for Evaluating Agent Response to Novelty,” arXiv:2203.12117, 2022.',
'[24] Manuela Chacon-Chamorro et al., “Cooperative Resilience in Artificial Intelligence Multiagent Systems,” arXiv:2409.13187, 2024.',
'[25] Xavier Cadet, Simona Boboila, Edward Koh, Peter Chin, and Alina Oprea, “Quantitative Resilience Modeling for Autonomous Cyber Defense,” Reinforcement Learning Journal, vol. 6, pp. 894–908, 2025.',
'[26] Chen Tessler, Yonathan Efroni, and Shie Mannor, “Action Robust Reinforcement Learning and Applications in Continuous Control,” in Proceedings of the 36th International Conference on Machine Learning, PMLR, vol. 97, 2019.',
'[27] Yuhui Wang, Hao He, and Xiaoyang Tan, “Robust Reinforcement Learning in POMDPs with Incomplete and Noisy Observations,” arXiv:1902.05795, 2019.',
'[28] Daniel Jarne Ornia, Licio Romao, Lewis Hammond, Manuel Mazo Jr., and Alessandro Abate, “Bounded Robustness in Reinforcement Learning via Lexicographic Objectives,” in Proceedings of Learning for Dynamics and Control, PMLR, vol. 242, pp. 954–967, 2024.',
'[29] Shangding Gu, Long Yang, Yali Du, Guang Chen, Florian Walter, Jun Wang, and Alois Knoll, “A Review of Safe Reinforcement Learning: Methods, Theories and Applications,” arXiv:2205.10330, 2022.',
'[30] Zichen Liu, Guoji Fu, Chao Du, Wee Sun Lee, and Min Lin, “Continual Reinforcement Learning by Planning with Online World Models,” in Proceedings of the 42nd International Conference on Machine Learning, PMLR, vol. 267, pp. 38397–38423, 2025.',
]
for r in refs:
    add_ref_before(doc, 'Παραρτήματα', r)

# Append audit corrections to appendix without duplicating existing exact facts too heavily.
insert_before(doc, 'Α.2 Τελικές διατάξεις και ανεξάρτητες επαναλήψεις', [
('Normal', 'Το DEVELOPMENT tuning χρησιμοποίησε έξι προκαθορισμένες υποψήφιες ρυθμίσεις ανά μέθοδο στις ίδιες 3 tuning-only roots × 2 development layouts, δηλαδή 180 tuning units συνολικά. Οι επιλογές πριν από το final reserve ήταν q-c06, sarsa-c06, dqn-c05, ppo-c06 και dyna-c03. Το search space ήταν bounded και ίσης έκτασης ανά method, όχι εξαντλητικό.'),
])
insert_before(doc, 'Α.3 Ορισμοί των διαταραχών', [
('Normal', 'Generation seeds των final layouts: gw-l1-final-a=57001 και gw-l1-final-b=57002. Για root r_i, i=01,…,12, τα seed streams είναι initialization=71000+i, exploration=72000+i, scenario=73000+i, environment=74000+i, action disturbance=75000+i και observation disturbance=76000+i. Η επιλογή n=12 προέκυψε από predeclared DEVELOPMENT sizing μεταξύ {12,16,20,24}, με μέγιστο 95% Student-t half-width 0,1428 στο πρώτο qualifying candidate.'),
])
insert_before(doc, 'Α.4 Προϋπολογισμοί και συμβόλαιο ανάκαμψης', [
('Normal', 'Action failure: με πιθανότητα 0,15 η intended action εκτελείται ως no-op, η ground-truth θέση δεν αλλάζει, collision=false και η reward είναι η ordinary step reward −0,1 εκτός άλλης terminal/reward rule. Observation corruption: με πιθανότητα 0,05 η delivered observation δειγματοληπτείται ομοιόμορφα από τα μη-obstacle κελιά εξαιρώντας την τρέχουσα ground-truth θέση· goal και start δεν αποκλείονται ειδικά από το support. Η πραγματική μετάβαση και reward δεν αλλάζουν.'),
])

# Synchronize the bibliography provenance appendix with the immutable T-716 consumer snapshot.
find_exact(doc, 'ada0d1aec7511098fd12610ae9e5abe7aea875cd').text = '27674a566ab55e4491b74243fe077a31ef81ae73'
find_exact(doc, 'Το formal citation layer περιλαμβάνει 123 citation-ready sources. Η πλήρης research corpus περιλαμβάνει 599 canonical sources και 19 research materials, αλλά formal thesis citations επιτρέπονται μόνο από το citation-ready manifest.').text = (
    'Το formal citation layer περιλαμβάνει 129 citation-ready sources. Η πλήρης research corpus περιλαμβάνει 601 canonical sources, '
    '19 research materials και 281 indexed originals, αλλά formal thesis citations επιτρέπονται μόνο από το citation-ready manifest.'
)

# Save and verify.
doc.save(OUT)
after=Document(OUT)
after_media=media_hashes(OUT)
# Locate bibliography/main-body boundaries after edits.
bib_idx=next(i for i,p in enumerate(after.paragraphs) if p.text.strip()=='Βιβλιογραφία')
app_idx=next(i for i,p in enumerate(after.paragraphs) if p.text.strip()=='Παραρτήματα')
main_words=words(after, bib_idx)
whole_words=words(after)
ref_nums=[]
for p in after.paragraphs[bib_idx+1:app_idx]:
    m=re.match(r'^\[(\d+)\]',p.text.strip())
    if m: ref_nums.append(int(m.group(1)))
text='\n'.join(p.text for p in after.paragraphs[:bib_idx])
used=sorted({int(x) for x in re.findall(r'\[(\d+)\]',text)})
unresolved_src=re.findall(r'SRC-[A-F0-9]{10}',text)
report={
    'status':'pass',
    'source_sha256':sha256(SRC),
    'output_sha256':sha256(OUT),
    'base_paragraph_count':base_paragraphs,
    'paragraph_count':len(after.paragraphs),
    'base_word_count_local':base_words,
    'whole_document_word_count_local':whole_words,
    'main_body_word_count_to_bibliography_local':main_words,
    'bibliography_reference_numbers':ref_nums,
    'bibliography_reference_count':len(ref_nums),
    'used_reference_numbers':used,
    'unused_reference_numbers':sorted(set(ref_nums)-set(used)),
    'missing_reference_numbers':sorted(set(used)-set(ref_nums)),
    'inline_shape_count':len(after.inline_shapes),
    'table_count':len(after.tables),
    'media_preserved':before_media==after_media,
    'media_count':len(after_media),
    'unresolved_src_ids':unresolved_src,
    'contains_required_tuning_count':'180 tuning units' in text,
    'contains_required_winners':all(x in text for x in ['q-c06','sarsa-c06','dqn-c05','ppo-c06','dyna-c03']),
    'contains_root_sizing_half_width':'0,1428' in text,
    'contains_action_failure_semantics':all(x in text for x in ['no-op','collision=false','−0,1']),
    'contains_observation_support_goal':'Goal και start δεν αποκλείονται' in text or 'goal και start δεν αποκλείονται' in text,
    'contains_directed_recovery_gap':'gₖ=Nₖ−Dₖ' in text and 'gₖ≤0,10' in text,
    'contains_two_window_recovery':'δύο διαδοχικά in-tolerance windows' in text,
    'contains_right_censoring':'recovery_time=null' in text,
    'bibliography_consumer_snapshot':'27674a566ab55e4491b74243fe077a31ef81ae73',
    'bibliography_canonical_source_count':601,
    'bibliography_citation_ready_count':129,
    'bibliography_research_material_count':19,
    'bibliography_indexed_original_count':281,
    'stale_bibliography_provenance_absent':all(x not in '\n'.join(p.text for p in after.paragraphs) for x in [
        'ada0d1aec7511098fd12610ae9e5abe7aea875cd',
        '123 citation-ready sources',
        '599 canonical sources',
    ]),
}
if (not report['media_preserved'] or report['missing_reference_numbers'] or report['unused_reference_numbers'] or unresolved_src or not report['stale_bibliography_provenance_absent']):
    report['status']='fail'
QA.write_text(json.dumps(report,ensure_ascii=False,indent=2)+'\n',encoding='utf-8')
print(json.dumps(report,ensure_ascii=False,indent=2))
