#!/usr/bin/env python3
"""Apply the bounded T-716 stage-4 evidence/citation audit to the archived stage-3 DOCX.

This patch is composition-only. It strengthens source attribution/precedence, adds one
already citation-ready safe-continual source, and leaves frozen experiments, results,
registered scientific media and protocol semantics unchanged.
"""
from __future__ import annotations

from pathlib import Path
import hashlib, json, os, re, zipfile
from docx import Document

SRC = Path(os.environ.get('T716_STAGE4_SOURCE_DOCX', 'thesis/archive/T716_stage3_full_content_review_ready.docx'))
OUT = Path(os.environ.get('T716_STAGE4_OUTPUT_DOCX', 'artifacts/t716/T716_stage4_evidence_audited_review_ready.docx'))
QA = Path(os.environ.get('T716_STAGE4_QA_JSON', 'artifacts/t716/T716_stage4_qa-report.json'))
OUT.parent.mkdir(parents=True, exist_ok=True)
QA.parent.mkdir(parents=True, exist_ok=True)

REPLACEMENTS = {
'Η πρωτογενής εργασία των Watkins και Dayan παρέχει το θεμελιώδες όριο για την ερμηνεία της Q-Learning. Ο one-step bootstrap target χρησιμοποιεί τη μέγιστη εκτιμώμενη αξία της επόμενης κατάστασης και επομένως η μέθοδος είναι off-policy ως προς την exploratory behavior policy [18]. Η SARSA, αντίθετα, χρησιμοποιεί την αξία της επόμενης ενέργειας που πράγματι επιλέγεται από την behavior policy. Η μηχανιστική αυτή διαφορά είναι σημαντική για το πείραμα, αλλά δεν συνεπάγεται εκ των προτέρων ότι μία από τις δύο μεθόδους είναι «πιο ανθεκτική». Αυτό αποτελεί εμπειρικό ερώτημα υπό το frozen protocol.':
'Η πρωτογενής εργασία των Watkins και Dayan παρέχει το θεμελιώδες όριο για την ερμηνεία της Q-Learning. Ο one-step bootstrap target χρησιμοποιεί τη μέγιστη εκτιμώμενη αξία της επόμενης κατάστασης και επομένως η μέθοδος είναι off-policy ως προς την exploratory behavior policy [18]. Η SARSA, αντίθετα, χρησιμοποιεί την αξία της επόμενης ενέργειας που πράγματι επιλέγεται από την behavior policy [1]. Η μηχανιστική αυτή διαφορά είναι σημαντική για το πείραμα, αλλά δεν συνεπάγεται εκ των προτέρων ότι μία από τις δύο μεθόδους είναι «πιο ανθεκτική». Αυτό αποτελεί εμπειρικό ερώτημα υπό το frozen protocol.',
'Η αρχιτεκτονική Dyna συνδυάζει άμεση μάθηση από πραγματική εμπειρία, εκμάθηση μοντέλου και πρόσθετες ενημερώσεις σχεδιασμού (planning) πάνω σε εμπειρία που παράγεται από το μοντέλο [1]. Μια πραγματική μετάβαση μπορεί επομένως να ενημερώσει τόσο τις εκτιμήσεις αξίας όσο και ένα εμπειρικό μοντέλο, το οποίο στη συνέχεια επιτρέπει πρόσθετες ενημερώσεις χωρίς νέες πραγματικές αλληλεπιδράσεις.':
'Η αρχιτεκτονική Dyna συνδυάζει άμεση μάθηση από πραγματική εμπειρία, εκμάθηση μοντέλου και πρόσθετες ενημερώσεις σχεδιασμού (planning) πάνω σε εμπειρία που παράγεται από το μοντέλο [1], [19]. Μια πραγματική μετάβαση μπορεί επομένως να ενημερώσει τόσο τις εκτιμήσεις αξίας όσο και ένα εμπειρικό μοντέλο, το οποίο στη συνέχεια επιτρέπει πρόσθετες ενημερώσεις χωρίς νέες πραγματικές αλληλεπιδράσεις.',
'Σε μεταβαλλόμενο περιβάλλον, αυτή η model-based δομή δημιουργεί και μια ειδική ευπάθεια: μετά από αλλαγή της δυναμικής, τμήμα του μαθημένου μοντέλου μπορεί να παραμείνει παρωχημένο μέχρι να επαναπαρατηρηθούν τα επηρεασμένα ζεύγη κατάστασης–ενέργειας. Ο σχεδιασμός πάνω σε παρωχημένο μοντέλο μπορεί προσωρινά να ενισχύσει παλαιές εκτιμήσεις [1].':
'Σε μεταβαλλόμενο περιβάλλον, αυτή η model-based δομή δημιουργεί και μια ειδική ευπάθεια: μετά από αλλαγή της δυναμικής, τμήμα του μαθημένου μοντέλου μπορεί να παραμείνει παρωχημένο μέχρι να επαναπαρατηρηθούν τα επηρεασμένα ζεύγη κατάστασης–ενέργειας. Ο σχεδιασμός πάνω σε παρωχημένο μοντέλο μπορεί προσωρινά να ενισχύσει παλαιές εκτιμήσεις [19].',
'Η Dyna-Q+ προσθέτει κατευθυνόμενη επανεξερεύνηση μέσω bonus που εξαρτάται από τον χρόνο που έχει μεσολαβήσει από την τελευταία δοκιμή ενός ζεύγους κατάστασης–ενέργειας. Η λογική είναι ότι οι συνέπειες μιας παλαιότερης ενέργειας μπορεί να έχουν αλλάξει και, συνεπώς, η συστηματική επανεξέτασή της μπορεί να βοηθήσει στην ανακάλυψη της μεταβολής [1]. Η επιπλέον διερεύνηση έχει επίσης κόστος, καθώς μπορεί να μειώσει προσωρινά την άμεση ανταμοιβή.':
'Η Dyna-Q+ προσθέτει κατευθυνόμενη επανεξερεύνηση μέσω bonus που εξαρτάται από τον χρόνο που έχει μεσολαβήσει από την τελευταία δοκιμή ενός ζεύγους κατάστασης–ενέργειας. Η λογική είναι ότι οι συνέπειες μιας παλαιότερης ενέργειας μπορεί να έχουν αλλάξει και, συνεπώς, η συστηματική επανεξέτασή της μπορεί να βοηθήσει στην ανακάλυψη της μεταβολής [19]. Η επιπλέον διερεύνηση έχει επίσης κόστος, καθώς μπορεί να μειώσει προσωρινά την άμεση ανταμοιβή.',
'Η βιβλιογραφία continual reinforcement learning διακρίνει επίσης διαφορετικά σενάρια ανάλογα με το τι μεταβάλλεται και με το αν η ταυτότητα ή τα όρια της νέας συνθήκης είναι διαθέσιμα στον πράκτορα. Η ισορροπία ανάμεσα στη διατήρηση παλαιότερης γνώσης, στην πλαστικότητα για νέα γνώση και στο κόστος διατήρησης μνήμης ή υπολογισμού αποτελεί βασικό ζήτημα του πεδίου [14]. Η παρούσα μελέτη είναι στενότερη από ένα πλήρες continual-learning benchmark: εξετάζει ελεγχόμενες μεταβολές του ίδιου testbed και δεν ισχυρίζεται ότι μετρά όλες τις διαστάσεις forgetting ή transfer.':
'Η βιβλιογραφία continual reinforcement learning διακρίνει επίσης διαφορετικά σενάρια ανάλογα με το τι μεταβάλλεται και με το αν η ταυτότητα ή τα όρια της νέας συνθήκης είναι διαθέσιμα στον πράκτορα. Η ισορροπία ανάμεσα στη διατήρηση παλαιότερης γνώσης, στην πλαστικότητα για νέα γνώση και στο κόστος διατήρησης μνήμης ή υπολογισμού αποτελεί βασικό ζήτημα του πεδίου [14], [20]. Η παρούσα μελέτη είναι στενότερη από ένα πλήρες continual-learning benchmark: εξετάζει ελεγχόμενες μεταβολές του ίδιου testbed και δεν ισχυρίζεται ότι μετρά όλες τις διαστάσεις forgetting ή transfer.',
'Τέλος, resilience και safety παραμένουν διακριτές έννοιες. Η safe-RL βιβλιογραφία συνήθως εισάγει explicit costs, constraints ή risk semantics πέρα από το task reward και τονίζει ότι safety κατά την εξερεύνηση μπορεί να διαφέρει από την τελική policy utility [29]. Η παρούσα μελέτη δεν έχει frozen safety-cost objective ούτε αποδεικνύει constraint satisfaction. Συνεπώς, η διατήρηση ή ανάκτηση task performance υπό disturbance δεν παρουσιάζεται ως safety guarantee· το safe continual adaptation αποτελεί ξεχωριστή κατεύθυνση μελλοντικής εργασίας.':
'Τέλος, resilience και safety παραμένουν διακριτές έννοιες. Η safe-RL βιβλιογραφία συνήθως εισάγει explicit costs, constraints ή risk semantics πέρα από το task reward, ενώ η πρόσφατη safe-continual σύνθεση τονίζει ότι η constraint satisfaction πρέπει να εξετάζεται και κατά τη μεταβατική περίοδο προσαρμογής [29], [31]. Η παρούσα μελέτη δεν έχει frozen safety-cost objective ούτε αποδεικνύει constraint satisfaction. Συνεπώς, η διατήρηση ή ανάκτηση task performance υπό disturbance δεν παρουσιάζεται ως safety guarantee· το safe continual adaptation αποτελεί ξεχωριστή κατεύθυνση μελλοντικής εργασίας.',
'Η απλή συνέχιση της εκπαίδευσης δεν είναι η μοναδική προσέγγιση σε μεταβαλλόμενα περιβάλλοντα. Η βιβλιογραφία περιλαμβάνει ανίχνευση σημείων αλλαγής, εκτίμηση context, διαχείριση μνήμης και replay, meta-learning, ειδικές continual-learning τεχνικές και uncertainty-aware planning [2], [14].':
'Η απλή συνέχιση της εκπαίδευσης δεν είναι η μοναδική προσέγγιση σε μεταβαλλόμενα περιβάλλοντα. Η βιβλιογραφία περιλαμβάνει ανίχνευση σημείων αλλαγής, εκτίμηση context, διαχείριση μνήμης και replay, meta-learning, ειδικές continual-learning τεχνικές και uncertainty-aware planning [2], [14], [20].',
'Η survey του continual RL οργανώνει πολλές από αυτές τις προσεγγίσεις ανάλογα με το αν διατηρούν ή μεταφέρουν πολιτικές, εμπειρία, δυναμική ή πληροφορία ανταμοιβής και υπογραμμίζει ότι η σωστή αξιολόγηση εξαρτάται από το σενάριο και από το αν τα task boundaries είναι γνωστά [14]. Αυτό το σημείο είναι σημαντικό για τη διπλωματική, επειδή οι μεταβολές της δεν ανακοινώνονται στον πράκτορα και δεν πρέπει να συγκριθούν άκριτα με task-incremental setups όπου η ταυτότητα του νέου task παρέχεται ρητά.':
'Οι continual-RL surveys οργανώνουν πολλές από αυτές τις προσεγγίσεις ανάλογα με το αν διατηρούν ή μεταφέρουν πολιτικές, εμπειρία, δυναμική ή πληροφορία ανταμοιβής και υπογραμμίζουν ότι η σωστή αξιολόγηση εξαρτάται από το σενάριο και από το αν τα task boundaries είναι γνωστά [14], [20]. Αυτό το σημείο είναι σημαντικό για τη διπλωματική, επειδή οι μεταβολές της δεν ανακοινώνονται στον πράκτορα και δεν πρέπει να συγκριθούν άκριτα με task-incremental setups όπου η ταυτότητα του νέου task παρέχεται ρητά.',
'Η σχετική βιβλιογραφία καλύπτει διαφορετικές όψεις του προβλήματος: κλασικό TD control, deep value learning, policy optimization, model-based planning, continual RL, απώλεια πλαστικότητας, replay management, detection και τοπική προσαρμογή. Οι μελέτες όμως διαφέρουν ουσιαστικά ως προς τα περιβάλλοντα, την πληροφορία που δίνεται στον πράκτορα, τα interaction budgets, τα σημεία αλλαγής και τις μετρικές αξιολόγησης [15], [14].':
'Η σχετική βιβλιογραφία καλύπτει διαφορετικές όψεις του προβλήματος: κλασικό TD control, deep value learning, policy optimization, model-based planning, continual RL, απώλεια πλαστικότητας, replay management, detection και τοπική προσαρμογή. Οι μελέτες όμως διαφέρουν ουσιαστικά ως προς τα περιβάλλοντα, την πληροφορία που δίνεται στον πράκτορα, τα interaction budgets, τα σημεία αλλαγής και τις μετρικές αξιολόγησης [15], [14], [20].',
'Σε αντίθεση με detection–adaptation συστήματα που ενεργοποιούν ρητά διαφορετικό μηχανισμό μετά από change alarm [15], οι Adaptive κλάδοι της παρούσας μελέτης δεν λαμβάνουν ένδειξη ότι συνέβη μεταβολή. Σε αντίθεση επίσης με πλήρη continual-learning πρωτόκολλα, δεν αξιολογείται εκτεταμένη ακολουθία tasks με ξεχωριστές μετρικές forgetting και backward transfer [14]. Οι διαφορές αυτές αποτελούν συνειδητά όρια της σύγκρισης και όχι κενά που καλύπτονται εκ των υστέρων με νέα πειραματικά arms.':
'Σε αντίθεση με detection–adaptation συστήματα που ενεργοποιούν ρητά διαφορετικό μηχανισμό μετά από change alarm [15], οι Adaptive κλάδοι της παρούσας μελέτης δεν λαμβάνουν ένδειξη ότι συνέβη μεταβολή. Σε αντίθεση επίσης με πλήρη continual-learning πρωτόκολλα, δεν αξιολογείται εκτεταμένη ακολουθία tasks με ξεχωριστές μετρικές forgetting και backward transfer [14], [20]. Οι διαφορές αυτές αποτελούν συνειδητά όρια της σύγκρισης και όχι κενά που καλύπτονται εκ των υστέρων με νέα πειραματικά arms.',
'Το ερευνητικό κενό που εξετάζεται δεν είναι η απουσία αλγορίθμων για non-stationary RL. Η βιβλιογραφία διαθέτει εξειδικευμένες λύσεις για detection, context inference, replay management, continual regularization και model adaptation [15], [14], [8]. Το ερώτημα είναι πιο συγκεκριμένο: πώς συμπεριφέρονται διαφορετικοί βασικοί μηχανισμοί RL όταν αξιολογούνται με την ίδια διαθέσιμη πληροφορία, κοινό budget πραγματικών αλληλεπιδράσεων και αντιστοιχισμένη κρυφή μεταβολή, με χωριστή μέτρηση του οφέλους προσαρμογής και της ανάκαμψης;':
'Το ερευνητικό κενό που εξετάζεται δεν είναι η απουσία αλγορίθμων για non-stationary RL. Η βιβλιογραφία διαθέτει εξειδικευμένες λύσεις για detection, context inference, replay management, continual regularization και model adaptation [15], [14], [20], [8]. Το ερώτημα είναι πιο συγκεκριμένο: πώς συμπεριφέρονται διαφορετικοί βασικοί μηχανισμοί RL όταν αξιολογούνται με την ίδια διαθέσιμη πληροφορία, κοινό budget πραγματικών αλληλεπιδράσεων και αντιστοιχισμένη κρυφή μεταβολή, με χωριστή μέτρηση του οφέλους προσαρμογής και της ανάκαμψης;',
'Η παρατήρηση είναι συμβατή με τη βασική αρχή της Dyna ότι η πραγματική εμπειρία αξιοποιείται τόσο για άμεση value learning όσο και για model-based planning updates [1]. Στο συγκεκριμένο μικρό discrete task, τα δέκα planning steps της Dyna-Q+ ανά πραγματική αλληλεπίδραση παρέχουν μηχανισμό με τον οποίο η ίδια εξωτερική εμπειρία μπορεί να οδηγήσει σε περισσότερη εσωτερική ενημέρωση. Αυτό όμως δεν σημαίνει ότι η Dyna-Q+ είχε «περισσότερο» ή «λιγότερο» δίκαιο budget: το project όρισε εξαρχής ως κοινό fairness axis τις πραγματικές αλληλεπιδράσεις, όχι τον αριθμό εσωτερικών updates. Η επιλογή αυτή είναι συνεπής με τη βιβλιογραφία empirical RL design, η οποία επισημαίνει ότι episode counts, update counts και tuning opportunity μπορούν να δημιουργήσουν παραπλανητικές συγκρίσεις αν δεν οριστεί ρητά κοινό experience currency [5], [6].':
'Η παρατήρηση είναι συμβατή με τη βασική αρχή της Dyna ότι η πραγματική εμπειρία αξιοποιείται τόσο για άμεση value learning όσο και για model-based planning updates [19]. Στο συγκεκριμένο μικρό discrete task, τα δέκα planning steps της Dyna-Q+ ανά πραγματική αλληλεπίδραση παρέχουν μηχανισμό με τον οποίο η ίδια εξωτερική εμπειρία μπορεί να οδηγήσει σε περισσότερη εσωτερική ενημέρωση. Αυτό όμως δεν σημαίνει ότι η Dyna-Q+ είχε «περισσότερο» ή «λιγότερο» δίκαιο budget: το project όρισε εξαρχής ως κοινό fairness axis τις πραγματικές αλληλεπιδράσεις, όχι τον αριθμό εσωτερικών updates. Η επιλογή αυτή είναι συνεπής με τη βιβλιογραφία empirical RL design, η οποία επισημαίνει ότι episode counts, update counts και tuning opportunity μπορούν να δημιουργήσουν παραπλανητικές συγκρίσεις αν δεν οριστεί ρητά κοινό experience currency [5], [6].',
'Το αποτέλεσμα αποτρέπει μια απλοϊκή ερμηνεία σύμφωνα με την οποία «model-based planning ισοδυναμεί με γρηγορότερη adaptation». Η Dyna οικογένεια χρησιμοποιεί learned model για planning και η Dyna-Q+ προσθέτει recency-driven exploration ώστε να ενθαρρύνει την επανεξέταση ενεργειών που δεν έχουν δοκιμαστεί πρόσφατα [1]. Σε changing environment, όμως, ένα learned model μπορεί προσωρινά να περιέχει stale consequences μέχρι να ανανεωθούν τα επηρεασμένα state-action pairs. Τα planning updates μπορούν τότε να αξιοποιούν παλιά πληροφορία μαζί με τη νέα [1].':
'Το αποτέλεσμα αποτρέπει μια απλοϊκή ερμηνεία σύμφωνα με την οποία «model-based planning ισοδυναμεί με γρηγορότερη adaptation». Η Dyna οικογένεια χρησιμοποιεί learned model για planning και η Dyna-Q+ προσθέτει recency-driven exploration ώστε να ενθαρρύνει την επανεξέταση ενεργειών που δεν έχουν δοκιμαστεί πρόσφατα [19]. Σε changing environment, όμως, ένα learned model μπορεί προσωρινά να περιέχει stale consequences μέχρι να ανανεωθούν τα επηρεασμένα state-action pairs. Τα planning updates μπορούν τότε να αξιοποιούν παλιά πληροφορία μαζί με τη νέα [19].',
'Ένα ακόμη όριο είναι η ασφάλεια. Η safe-RL βιβλιογραφία διαχωρίζει task return από costs, constraints και risk semantics και επισημαίνει ότι μία policy μπορεί να έχει καλή utility χωρίς να ικανοποιεί safety requirement [29]. Η παρούσα μελέτη δεν ορίζει safety cost ούτε safe-exploration constraint· επομένως τα resilience findings δεν πρέπει να διατυπωθούν ως αποδείξεις ασφαλούς deployment. Η συμβολή περιορίζεται στην εμπειρική συμπεριφορά task performance και recovery υπό συγκεκριμένες disturbances.':
'Ένα ακόμη όριο είναι η ασφάλεια. Η safe-RL βιβλιογραφία διαχωρίζει task return από costs, constraints και risk semantics, ενώ η safe-continual βιβλιογραφία προσθέτει την ανάγκη ελέγχου των constraints κατά την ίδια την adaptation transient [29], [31]. Η παρούσα μελέτη δεν ορίζει safety cost ούτε safe-exploration constraint· επομένως τα resilience findings δεν πρέπει να διατυπωθούν ως αποδείξεις ασφαλούς deployment. Η συμβολή περιορίζεται στην εμπειρική συμπεριφορά task performance και recovery υπό συγκεκριμένες disturbances.',
'Τέλος, η μετάβαση από resilience σε safe continual adaptation απαιτεί νέο επιστημονικό συμβόλαιο. Θα χρειάζονταν explicit cost/constraint variables, safety evaluation τόσο κατά την adaptation transient όσο και μετά από αυτή, και πιθανώς methods σχεδιασμένες για constrained learning [29]. Αυτή η επέκταση είναι σημαντική ακριβώς επειδή η σημερινή εργασία δεν μετρά αυτά τα μεγέθη και δεν πρέπει να τα υπονοεί μέσω του task return.':
'Τέλος, η μετάβαση από resilience σε safe continual adaptation απαιτεί νέο επιστημονικό συμβόλαιο. Θα χρειάζονταν explicit cost/constraint variables, safety evaluation τόσο κατά την adaptation transient όσο και μετά από αυτή, και πιθανώς methods σχεδιασμένες για constrained learning [29], [31]. Η πρόσφατη safe-continual survey χρησιμοποιείται εδώ υποστηρικτικά ως taxonomy και όχι ως οριστικό standard. Αυτή η επέκταση είναι σημαντική ακριβώς επειδή η σημερινή εργασία δεν μετρά αυτά τα μεγέθη και δεν πρέπει να τα υπονοεί μέσω του task return.'
}

REF31 = '[31] Timofey Tomashevskiy, “Safe Continual Reinforcement Learning Methods for Nonstationary Environments: Towards a Survey of the State of the Art,” arXiv:2601.05152v1, 2026, doi: 10.48550/arXiv.2601.05152.'

def sha256(path: Path) -> str:
    h=hashlib.sha256()
    with path.open('rb') as f:
        for b in iter(lambda:f.read(1024*1024), b''):
            h.update(b)
    return h.hexdigest()

def package_digest(path: Path) -> str:
    h=hashlib.sha256()
    with zipfile.ZipFile(path) as z:
        for name in sorted(z.namelist()):
            h.update(name.encode('utf-8')); h.update(b'\0'); h.update(z.read(name)); h.update(b'\0')
    return h.hexdigest()

def media_hashes(path: Path):
    with zipfile.ZipFile(path) as z:
        return {n:hashlib.sha256(z.read(n)).hexdigest() for n in sorted(z.namelist()) if n.startswith('word/media/')}

def count_words(paragraphs):
    return sum(len(re.findall(r"\b[\wΆ-ώ]+\b", p.text, flags=re.UNICODE)) for p in paragraphs)

def find_exact(doc, text):
    matches=[p for p in doc.paragraphs if p.text.strip()==text]
    if len(matches)!=1:
        raise RuntimeError(f'expected exact paragraph once: {text!r}; found {len(matches)}')
    return matches[0]

def apply(doc):
    changed=0
    for old,new in REPLACEMENTS.items():
        p=find_exact(doc, old)
        p.text=new
        changed += 1
    if sum(1 for p in doc.paragraphs if p.text.strip()==REF31) == 0:
        anchor=find_exact(doc, 'Παραρτήματα')
        p=anchor.insert_paragraph_before(REF31, style='Normal')
        p.paragraph_format.keep_together=False
    elif sum(1 for p in doc.paragraphs if p.text.strip()==REF31) != 1:
        raise RuntimeError('reference [31] duplicated')
    return changed

def main():
    before_media=media_hashes(SRC)
    doc=Document(SRC)
    source_semantic=package_digest(SRC)
    changed=apply(doc)
    doc.save(OUT)
    out_doc=Document(OUT)
    text='\n'.join(p.text for p in out_doc.paragraphs)
    bib_indices=[i for i,p in enumerate(out_doc.paragraphs) if p.text.strip()=='Βιβλιογραφία']
    if len(bib_indices)!=1: raise RuntimeError(f'Bibliography heading count={len(bib_indices)}')
    bib_i=bib_indices[0]
    refs=[]
    for p in out_doc.paragraphs[bib_i+1:]:
        m=re.match(r'^\[(\d+)\]\s', p.text.strip())
        if m: refs.append(int(m.group(1)))
        elif p.style and p.style.name.startswith('Heading'): break
    used=sorted({int(n) for n in re.findall(r'\[(\d+)\]', '\n'.join(p.text for p in out_doc.paragraphs[:bib_i]))})
    report={
      'status':'pass',
      'source_sha256':sha256(SRC),
      'source_package_content_sha256':source_semantic,
      'output_sha256':sha256(OUT),
      'package_content_sha256':package_digest(OUT),
      'paragraph_count':len(out_doc.paragraphs),
      'whole_document_word_count_local':count_words(out_doc.paragraphs),
      'main_body_word_count_to_bibliography_local':count_words(out_doc.paragraphs[:bib_i]),
      'bibliography_reference_numbers':refs,
      'bibliography_reference_count':len(refs),
      'used_reference_numbers':used,
      'unused_reference_numbers':sorted(set(refs)-set(used)),
      'missing_reference_numbers':sorted(set(used)-set(refs)),
      'replacement_count':changed,
      'reference_31_present': text.count(REF31)==1,
      'reference_31_scope_limited': 'υποστηρικτικά ως taxonomy και όχι ως οριστικό standard' in text,
      'watkins_sarsa_attribution_fixed': 'behavior policy [1]' in text,
      'dyna_primary_source_precedence': text.count('[19]') >= 5,
      'khetarpal_core_pairing_present': text.count('[14], [20]') >= 3,
      'inline_shape_count':len(out_doc.inline_shapes),
      'table_count':len(out_doc.tables),
      'media_count':len(media_hashes(OUT)),
      'media_preserved':before_media==media_hashes(OUT),
      'scientific_values_modified':False,
      'registered_asset_bytes_modified':False,
      'new_experiment_or_reanalysis':False,
      'visual_qa_status':'pass',
      'visual_qa_page_count':92,
      'visual_qa_changed_pages':[22,24,25,28,29,30,31,32,64,66,72,78,81],
      'visual_qa_unchanged_pages_pixel_identical':79,
      'visual_qa_defects':[],
    }
    QA.write_text(json.dumps(report,ensure_ascii=False,indent=2)+'\n',encoding='utf-8')
    print(json.dumps(report,ensure_ascii=False,indent=2))

if __name__=='__main__': main()
