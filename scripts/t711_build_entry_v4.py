#!/usr/bin/env python3
"""T-711 final cached-table accuracy layer.

Extends v3 after rendered-page QA found that the glossary acronym Markdown table is a
real Word table but was omitted from the cached List of Tables. The cache now scans the
same main manuscript, glossary, and cleaned appendix surfaces that can actually render
Word tables. Structural QA records the exact cache/table count equality.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

import t711_build_entry_v3 as v3


t711 = v3.t711


def _table_cache_lines() -> list[str]:
    result: list[str] = []
    sources = [
        *((path, "main") for path in t711.builder.MANUSCRIPT_FILES),
        (t711.builder.GLOSSARY_FILE, "glossary"),
        (t711.builder.APPENDIX_FILE, "appendix"),
    ]
    for path, mode in sources:
        md = t711._preprocess_markdown(path.read_text(encoding="utf-8"), mode)
        lines = md.splitlines()
        current = path.stem
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()
            if stripped.startswith("#"):
                current = t711.builder.strip_markdown_inline(stripped.lstrip("#").strip())
            if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                result.append(f"Πίνακας {len(result) + 1} — Σύνοψη για την ενότητα «{current}»")
                while i < len(lines) and lines[i].strip().startswith("|"):
                    i += 1
                continue
            i += 1
    return result


_original_v3_qa = t711._enhanced_qa


def _enhanced_qa(output: Path, qa_output: Path):
    _original_v3_qa(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    doc = Document(output)
    cached_tables = _table_cache_lines()
    report["cached_table_entry_count"] = len(cached_tables)
    report["rendered_table_count"] = len(doc.tables)
    if len(cached_tables) != len(doc.tables):
        report["status"] = "fail"
    qa_output.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if report["status"] != "pass":
        raise RuntimeError(f"T-711 v4 structural QA failed: {report}")


t711._markdown_table_cache_lines = _table_cache_lines
t711._enhanced_qa = _enhanced_qa

if __name__ == "__main__":
    t711.builder.main()
