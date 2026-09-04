#!/usr/bin/env python3
"""T-711A source-aware academic/compliance hardening.

This layer bypasses the v12 build wrapper while reusing its presentation helpers. It is
compatible with canonical manuscript text that has already received the verified
citation/wording corrections. Scientific results and registered T-613 media remain
immutable.
"""

from __future__ import annotations

import json
import os
from pathlib import Path

from docx import Document

import t711_build_entry_v12 as v12


t711 = v12.t711
_previous_build = v12._previous_build  # accepted v11 build; bypass v12's old-text assumptions

EXPECTED_ADDED_CITATIONS = {
    "SRC-0A4AFAC8E9",  # Agarwal et al. — RL statistical reporting
    "SRC-95C9DAEE68",  # Liu — non-stationary detection/adaptation
    "SRC-F909CABDEB",  # continual-RL survey / scope taxonomy
}
EXPECTED_REFERENCE_COUNT = 17


def _build(output: Path, qa_output: Path) -> None:
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    media_before = v12._media_hashes(output)
    doc = Document(output)

    # Keep compatibility with the accepted pre-hardening Results source while allowing
    # canonical source-level corrections to remove the need for post-processing later.
    applied_replacements = 0
    for paragraph in doc.paragraphs:
        replacement = v12.REPLACEMENTS.get(paragraph.text)
        if replacement is not None:
            paragraph.text = replacement
            applied_replacements += 1

    caption_count = 0
    rq3_caption = None
    for paragraph in doc.paragraphs:
        number = v12._caption_number(paragraph)
        if number not in v12.CAPTIONS:
            continue
        if len(paragraph.runs) < 3:
            raise RuntimeError(
                f"Unexpected caption run structure for figure {number}: {paragraph.text!r}"
            )
        paragraph.runs[2].text = " — " + v12.CAPTIONS[number]
        for extra_run in paragraph.runs[3:]:
            extra_run.text = ""
        caption_count += 1
        if number == 11:
            rq3_caption = paragraph

    if rq3_caption is None:
        raise RuntimeError("Figure 11 conditional-recovery caption was not found")
    v12._insert_after(
        rq3_caption,
        "Σημείωση ερμηνείας: ο χρόνος αυτός είναι υπό συνθήκη ανάκαμψης "
        "(conditional on recovery). Όταν το recovered n είναι πολύ μικρό, ιδίως "
        "n=1 ή n=2, η εκτίμηση είναι ασταθής και πρέπει να διαβάζεται μαζί με "
        "το recovery proportion και το restricted delay.",
    )

    cached_figure_count = v12._localize_cached_figure_list(doc)
    alt_text_count = v12._assign_alt_text(doc)

    props = doc.core_properties
    props.title = (
        "Σύγκριση και Αξιολόγηση Ανθεκτικών Πρακτόρων Τεχνητής Νοημοσύνης "
        "σε Περιβάλλοντα με Αβεβαιότητα"
    )
    props.subject = (
        "Διπλωματική εργασία — ενισχυτική μάθηση, ανθεκτικότητα και "
        "προσαρμογή υπό αβεβαιότητα"
    )
    props.author = ""  # official student identity remains externally gated
    props.last_modified_by = ""
    props.keywords = (
        "reinforcement learning; resilient AI agents; non-stationarity; adaptation; recovery"
    )
    props.comments = ""
    props.category = "Diploma Thesis"
    props.created = v12.datetime(2026, 9, 3, 0, 0, tzinfo=v12.timezone.utc)
    props.modified = v12.datetime(2026, 9, 3, 0, 0, tzinfo=v12.timezone.utc)
    doc.save(output)

    final_doc = Document(output)
    media_after = v12._media_hashes(output)
    placeholder_hits = v12._placeholder_hits(final_doc)
    placeholder_set = set(placeholder_hits)
    final_mode = os.environ.get("T711_FINAL_MODE", "0").strip().lower() in {
        "1", "true", "yes"
    }

    superiority_residue = [p.text for p in final_doc.paragraphs if "υπερείχε" in p.text]
    metadata_residue = v12._generator_metadata_residue(output)
    dqn_paragraphs = [p.text for p in final_doc.paragraphs if p.text.startswith("Η Deep Q-Network (DQN)")]
    dqn_foundation_ok = (
        len(dqn_paragraphs) == 1
        and "experience replay" in dqn_paragraphs[0]
        and "target network" not in dqn_paragraphs[0]
    )
    source_text = t711.builder.manuscript_order_text()
    added_citations_present = EXPECTED_ADDED_CITATIONS.issubset(set(t711.builder.SRC_RE.findall(source_text)))
    final_sha = v12._sha256(output)
    final_paragraph_count = len(final_doc.paragraphs)

    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": final_paragraph_count,
            "post_synthesis_paragraph_count": final_paragraph_count,
            "t711a_hardening_version": 13,
            "t711a_reader_text_rescanned_after_save": True,
            "dqn_foundation_citation_wording_corrected": dqn_foundation_ok,
            "postprocess_replacement_count": applied_replacements,
            "superiority_wording_residue": superiority_residue,
            "localized_figure_caption_count": caption_count,
            "localized_cached_figure_entry_count": cached_figure_count,
            "rq3_conditional_recovery_note_added": True,
            "image_alt_text_count": alt_text_count,
            "embedded_media_bytes_preserved": media_before == media_after,
            "docx_core_metadata_scrubbed": not metadata_residue,
            "generator_metadata_residue": metadata_residue,
            "placeholder_hits": placeholder_hits,
            "placeholder_count": len(placeholder_hits),
            "review_placeholders_expected": placeholder_set == v12.EXPECTED_REVIEW_PLACEHOLDERS,
            "final_mode": final_mode,
            "final_submission_ready": final_mode and not placeholder_hits,
            "targeted_related_work_citations_present": added_citations_present,
            "expected_reference_count": EXPECTED_REFERENCE_COUNT,
            "scientific_values_modified": False,
        }
    )

    hardening_ok = (
        report.get("status") == "pass"
        and report.get("citation_count") == EXPECTED_REFERENCE_COUNT
        and report.get("verified_reference_identity_count") == EXPECTED_REFERENCE_COUNT
        and added_citations_present
        and dqn_foundation_ok
        and caption_count == cached_figure_count == 24
        and alt_text_count == 25
        and media_before == media_after
        and not superiority_residue
        and not metadata_residue
        and final_paragraph_count >= 708
        and (
            (not final_mode and placeholder_set == v12.EXPECTED_REVIEW_PLACEHOLDERS)
            or (final_mode and not placeholder_hits)
        )
        and report.get("registered_asset_bytes_modified") is False
    )
    if not hardening_ok:
        report["status"] = "fail"

    qa_output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if report["status"] != "pass":
        raise RuntimeError(f"T-711A v13 hardening failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
