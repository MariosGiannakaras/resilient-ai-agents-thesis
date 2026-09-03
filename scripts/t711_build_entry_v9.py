#!/usr/bin/env python3
"""T-711 final reader-facing appendix cleanup.

Full rendered-page QA of v8 found three groups of visible appendix handoff bullets whose
paragraph text consisted only of registered FIG-* identifiers. The figures themselves
were already inserted correctly. This final layer removes only those identifier-only
paragraphs after composition, then re-opens the saved DOCX and fails closed if any
reader-visible asset-ID residue remains.

No manuscript claim, citation, scientific value, T-613 asset byte, manifest, caption or
figure placement is changed.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

import t711_build_entry_v8 as v8


t711 = v8.t711
_previous_build = t711.builder.build

_ASSET_ONLY_RE = re.compile(r"^FIG-(?:RQ\d+|METHOD)-[A-Z0-9-]+,?$")
_ASSET_ANY_RE = re.compile(r"\bFIG-(?:RQ\d+|METHOD)-[A-Z0-9-]+\b")


def _all_reader_text(doc: Document) -> list[str]:
    texts: list[str] = []
    texts.extend(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return texts


def _remove_identifier_only_paragraphs(output: Path) -> list[str]:
    doc = Document(output)
    removed: list[str] = []
    for paragraph in list(doc.paragraphs):
        value = paragraph.text.strip()
        if _ASSET_ONLY_RE.fullmatch(value):
            removed.append(value.rstrip(","))
            paragraph._element.getparent().remove(paragraph._element)
    doc.save(output)
    return removed


def _build(output: Path, qa_output: Path):
    _previous_build(output, qa_output)
    removed = _remove_identifier_only_paragraphs(output)

    report = json.loads(qa_output.read_text(encoding="utf-8"))
    doc = Document(output)
    reader_text = "\n".join(_all_reader_text(doc))
    residue = sorted(set(_ASSET_ANY_RE.findall(reader_text)))

    expected_removed = {
        "FIG-RQ1-004-FINAL-ROOTS",
        "FIG-RQ1-005-TIME-ROOTS",
        "FIG-RQ1-007-CONTRASTS",
        "FIG-RQ2-010-CONDITIONS",
        "FIG-RQ2-012-BENEFIT-ROOTS",
        "FIG-RQ2-013-HEATMAP",
        "FIG-RQ2-014-CONTRASTS",
        "FIG-RQ2-015-PAIRED-ROOTS",
        "FIG-RQ3-019-CONDITIONAL",
        "FIG-RQ3-021-ROOT-TRAJECTORIES",
        "FIG-RQ3-022-CENSORING",
        "FIG-RQ3-023-SENSITIVITY",
        "FIG-RQ3-024-CONTRASTS",
        "FIG-RQ3-025-TIMELINE",
    }
    removed_set = set(removed)

    report.update(
        {
            "reader_visible_asset_id_residue": residue,
            "removed_appendix_handoff_ids": sorted(removed_set),
            "removed_appendix_handoff_count": len(removed_set),
            "reader_text_rescanned_after_save": True,
            "registered_asset_bytes_modified": False,
            "scientific_values_modified": False,
        }
    )

    if residue or removed_set != expected_removed:
        report["status"] = "fail"

    qa_output.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if report["status"] != "pass":
        raise RuntimeError(f"T-711 v9 reader-facing QA failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
