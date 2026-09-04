#!/usr/bin/env python3
"""T-714 Chapter 4 paragraph pagination hardening.

Keeps the single bold architecture-flow summary paragraph together so the visual flow
is not split across a page boundary. Scientific content, values, citations and media
are untouched.
"""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document

import t711_build_entry_v17 as v17


t711 = v17.t711
_previous_build = v17._build

FLOW_PREFIX = "Frozen protocol/configuration → immutable Study recipe → deterministic plan"


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _build(output: Path, qa_output: Path) -> None:
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))

    doc = Document(output)
    matches = [p for p in doc.paragraphs if p.text.startswith(FLOW_PREFIX)]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected exactly one Chapter 4 architecture-flow paragraph, found {len(matches)}"
        )

    flow = matches[0]
    flow.paragraph_format.keep_together = True
    doc.save(output)

    final_doc = Document(output)
    final_matches = [p for p in final_doc.paragraphs if p.text.startswith(FLOW_PREFIX)]
    applied = (
        len(final_matches) == 1
        and final_matches[0].paragraph_format.keep_together is True
    )

    final_sha = _sha256(output)
    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": len(final_doc.paragraphs),
            "post_synthesis_paragraph_count": len(final_doc.paragraphs),
            "t711a_hardening_version": 18,
            "chapter4_flow_keep_together_applied": applied,
            "chapter4_flow_match_count": len(final_matches),
            "scientific_values_modified": False,
            "registered_asset_bytes_modified": False,
            "final_visual_qa_required": True,
        }
    )
    if not applied:
        report["status"] = "fail"

    qa_output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if report.get("status") != "pass":
        raise RuntimeError(f"T-714 v18 Chapter 4 pagination hardening failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
