#!/usr/bin/env python3
"""T-711 final presentation repair for the registered RQ-to-evidence map.

The frozen T-613 FIG-METHOD-027-RQ-MAP asset contains text overlap inside the
registered SVG/PNG geometry, confirmed by 100% rendered-page QA. This wrapper
preserves the frozen asset bytes and manifest unchanged, but replaces only its
Word presentation with a deterministic composition-native rendering carrying the
same non-quantitative semantic labels. All scientific figures and values remain
registered T-613 assets without alteration.
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

import t711_build_entry_v7 as v7


t711 = v7.t711

REPAIRED_ASSET_ID = "FIG-METHOD-027-RQ-MAP"
_previous_add_figure = t711._add_figure
_previous_enhanced_qa = t711._enhanced_qa


def _font(size: int):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _centered_multiline(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str):
    title_font = _font(34)
    body_font = _font(25)
    x0, y0, x1, y1 = box
    cx = (x0 + x1) // 2
    title_box = draw.textbbox((0, 0), title, font=title_font)
    title_w = title_box[2] - title_box[0]
    draw.text((cx - title_w / 2, y0 + 42), title, font=title_font, fill=(20, 20, 20))

    lines = body.split("\n")
    line_height = 34
    start_y = y0 + 103
    for index, line in enumerate(lines):
        bounds = draw.textbbox((0, 0), line, font=body_font)
        width = bounds[2] - bounds[0]
        draw.text((cx - width / 2, start_y + index * line_height), line, font=body_font, fill=(35, 35, 35))


def _render_native_rq_map() -> Path:
    width, height = 2100, 620
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = _font(34)
    title = "Research questions to evidence"
    bounds = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (bounds[2] - bounds[0])) / 2, 42), title, font=title_font, fill=(20, 20, 20))

    boxes = [
        (60, 170, 390, 430),
        (465, 170, 795, 430),
        (870, 170, 1200, 430),
        (1275, 170, 1605, 430),
        (1680, 170, 2010, 430),
    ]
    content = [
        ("RQ1", "final return\n+ time-average return"),
        ("RQ2", "Frozen loss\nAdaptive loss\nadaptation benefit"),
        ("RQ3", "recovery status\nrestricted delay\nconditional time"),
        ("Contrasts", "root-paired\ndirect comparisons"),
        ("Outputs", "registered\nfigures + tables"),
    ]

    for box, (title_text, body_text) in zip(boxes, content):
        draw.rounded_rectangle(box, radius=24, fill=(233, 242, 247), outline=(0, 114, 178), width=4)
        _centered_multiline(draw, box, title_text, body_text)

    arrow_y = 300
    for left, right in zip(boxes[:-1], boxes[1:]):
        x_start = left[2] + 22
        x_end = right[0] - 22
        draw.line((x_start, arrow_y, x_end, arrow_y), fill=(85, 85, 85), width=5)
        draw.polygon(
            [(x_end, arrow_y), (x_end - 20, arrow_y - 12), (x_end - 20, arrow_y + 12)],
            fill=(85, 85, 85),
        )

    path = Path(tempfile.gettempdir()) / "t711-rq-map-presentation-repair.png"
    image.save(path, format="PNG", optimize=False)
    return path


def _add_figure(doc, asset: dict, inserted: list[str]):
    if asset.get("asset_id") != REPAIRED_ASSET_ID:
        return _previous_add_figure(doc, asset, inserted)

    png = _render_native_rq_map()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(png), width=Inches(6.1))
    t711._add_seq_caption(doc, "Σχήμα", asset.get("caption") or REPAIRED_ASSET_ID)
    inserted.append(REPAIRED_ASSET_ID)


def _enhanced_qa(output: Path, qa_output: Path):
    _previous_enhanced_qa(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    doc = Document(output)
    report.update(
        {
            "presentation_repaired_assets": [REPAIRED_ASSET_ID],
            "presentation_repair_reason": "registered T-613 rendering contains overlapping labels; frozen bytes and manifest preserved",
            "registered_asset_bytes_modified": False,
            "scientific_values_modified": False,
            "final_visual_qa_required": True,
            "post_repair_paragraph_count": len(doc.paragraphs),
        }
    )
    if report.get("inserted_asset_count") != 24 or REPAIRED_ASSET_ID not in report.get("inserted_asset_ids", []):
        report["status"] = "fail"
    qa_output.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if report["status"] != "pass":
        raise RuntimeError(f"T-711 v8 structural QA failed: {report}")


t711._add_figure = _add_figure
t711._enhanced_qa = _enhanced_qa

if __name__ == "__main__":
    t711.builder.main()
