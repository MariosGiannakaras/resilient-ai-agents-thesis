#!/usr/bin/env python3
"""Composition-only T-715 audit hardening for the reader-facing thesis DOCX.

The helper inserts bounded methodology/reproducibility clarifications that reconcile the
reader-facing thesis with the already-frozen protocol and analysis authorities. It does not
recompute scientific evidence, alter numerical outcomes, or touch embedded media bytes.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as _Document


@dataclass(frozen=True)
class Block:
    anchor: str
    paragraphs: tuple[tuple[str, str], ...]


BLOCKS = (
    Block(
        "3.4 Αρχική εκπαίδευση και αξιολόγηση μετά τη μεταβολή",
        (
            ("Heading 3", "3.3.1 Επιλογή υπερπαραμέτρων και όριο της σύγκρισης"),
            ("Normal", "Οι τελικές ρυθμίσεις δεν προήλθαν από defaults ούτε επιλέχθηκαν από τα τελικά αποτελέσματα. Πριν ανοίξει το τελικό σύνολο αξιολόγησης εκτελέστηκε προκαθορισμένο, ίσης ευκαιρίας tuning σε DEVELOPMENT δεδομένα: έξι method-specific υποψήφιες ρυθμίσεις για καθεμία από τις πέντε μεθόδους, στις ίδιες τρεις tuning-only ρίζες, στις ίδιες δύο development διατάξεις, με κοινό budget 8.192 πραγματικών αλληλεπιδράσεων και κοινό πλέγμα probes. Έτσι προέκυψαν 6×3×2×5=180 tuning units, τα οποία ολοκληρώθηκαν χωρίς αποτυχία. Οι σπόροι τυχαιότητας δεν αποτέλεσαν παράμετρο tuning και το final reserve δεν χρησιμοποιήθηκε στην επιλογή."),
            ("Normal", "Η επιλογή έγινε ξεχωριστά για κάθε μέθοδο με προκαθορισμένο μηχανικό κανόνα: πρώτα μεγιστοποιήθηκε η ισοβαρής ως προς roots/layouts τραπεζοειδής χρονικά σταθμισμένη μέση καμπύλη επιτυχίας στα no-learning probes, έπειτα χρησιμοποιήθηκαν ως διαδοχικά tie-breakers η τελική επιτυχία, η χρονικά σταθμισμένη μέση απόδοση και, μόνο σε πλήρη ισοβαθμία, το λεξικογραφικά μικρότερο config ID. Οι επιλεγμένες ρυθμίσεις ήταν q-c06, sarsa-c06, dqn-c05, ppo-c06 και dyna-c03."),
            ("Normal", "Το search space ήταν σκόπιμα περιορισμένο και ίσης έκτασης, όχι εξαντλητικό: για Q-Learning και SARSA εξετάστηκαν learning rate {0,1, 0,2, 0,4} × ε {0,05, 0,10}, για Dyna-Q+ planning steps {5, 10, 20} × κ {0,0005, 0,001}, για DQN learning rate {0,0003, 0,001, 0,003} × target-update interval {128, 256} και για PPO learning rate {0,0001, 0,0003, 0,001} × epochs {5, 10}. Οι υπόλοιπες παράμετροι παρέμειναν σταθερές όπως καταγράφονται στο Παράρτημα Α. Επομένως, ιδιαίτερα για DQN και PPO, τα αποτελέσματα αφορούν αυτό το GridWorld, αυτό το interaction budget και αυτό το προκαθορισμένο bounded tuning· δεν αποτελούν γενική αξιολόγηση της μέγιστης δυνατής επίδοσης των αλγορίθμων."),
        ),
    ),
    Block(
        "3.6 Μετρήσεις για τα τρία ερευνητικά ερωτήματα",
        (
            ("Heading 3", "3.5.1 Ακριβής σημασιολογία των διαταραχών"),
            ("Normal", "Στις δύο μόνιμες αλλαγές η μεταβολή εφαρμόζεται στο πραγματικά εκτελούμενο action mapping και παραμένει ενεργή στα επόμενα επεισόδια. Στο swap-right-down ισχύει up→up, right→down, down→right, left→left. Στο cycle-clockwise ισχύει up→right, right→down, down→left, left→up. Ο πράκτορας δεν λαμβάνει εξωτερικό change indicator ούτε το πραγματικά εκτελεσμένο action, άρα η αλλαγή πρέπει να γίνει αντιληπτή μόνο μέσα από την εμπειρία."),
            ("Normal", "Στην αποτυχία ενέργειας, σε κάθε πραγματική αλληλεπίδραση πραγματοποιείται ανεξάρτητη δοκιμή με πιθανότητα 0,15. Όταν συμβεί αποτυχία, η intended action μετατρέπεται σε no-op: η πραγματική θέση δεν αλλάζει, δεν δηλώνεται collision και αποδίδεται η κανονική step reward −0,1. Η μετάβαση καταγράφεται κανονικά και, στους Adaptive κλάδους, χρησιμοποιείται από τον συνηθισμένο μηχανισμό μάθησης της μεθόδου. Η παρατήρηση που επιστρέφεται μετά το βήμα είναι η παρατήρηση της προκύπτουσας κατάστασης, εκτός αν ενεργοποιηθεί ξεχωριστά μηχανισμός αλλοίωσης παρατήρησης στη σχετική συνθήκη."),
            ("Normal", "Στην αλλοίωση παρατήρησης, μετά την πραγματική μετάβαση του περιβάλλοντος πραγματοποιείται ανεξάρτητη δοκιμή με πιθανότητα 0,05. Αν ενεργοποιηθεί, η delivered observation επιλέγεται ομοιόμορφα από όλα τα κελιά του 7×7 που δεν είναι εμπόδια και δεν συμπίπτουν με την τρέχουσα πραγματική θέση. Ο στόχος ή η αρχική θέση δεν αποκλείονται ως κατηγορίες, εφόσον είναι έγκυρα κελιά και δεν αποτελούν την τρέχουσα πραγματική θέση. Η αλλοίωση δεν μετακινεί τον πράκτορα και δεν αλλάζει reward ή ground-truth transition· αλλάζει μόνο την παρατήρηση που λαμβάνει ο agent και, όταν η μάθηση είναι ενεργή, αυτή η delivered observation τροφοδοτεί την κανονική διαδικασία ενημέρωσης."),
        ),
    ),
    Block(
        "3.7 Επαναλήψεις και στατιστική σύνοψη",
        (
            ("Heading 3", "3.6.1 Μαθηματικοί ορισμοί των βασικών μετρικών"),
            ("Normal", "Για το RQ1, έστω (tₖ, yₖ) τα προκαθορισμένα Phase-A probe σημεία μιας μονάδας method/root/layout, όπου tₖ είναι ο αριθμός πραγματικών αλληλεπιδράσεων και yₖ το no-learning probe return_mean. Η κύρια τελική ονομαστική επίδοση είναι P_final = y_K. Η δευτερεύουσα μέση επίδοση κατά τη μάθηση είναι η τραπεζοειδής χρονικά σταθμισμένη μέση τιμή πάνω στον πραγματικό άξονα αλληλεπιδράσεων:"),
            ("Normal", "P_avg = [1 / (t_K − t_0)] · Σₖ₌₀ᴷ⁻¹ (tₖ₊₁ − tₖ) · (yₖ + yₖ₊₁) / 2"),
            ("Normal", "Για το RQ2 χρησιμοποιείται το Phase-B return_sum. Για τις matched τιμές Frozen-Nominal (FN), Frozen-Disturbed (FD), Adaptive-Nominal (AN) και Adaptive-Disturbed (AD), επειδή μεγαλύτερη απόδοση είναι καλύτερη, οι απώλειες και το όφελος προσαρμογής ορίζονται ως:"),
            ("Normal", "L_F = FN − FD"),
            ("Normal", "L_A = AN − AD"),
            ("Normal", "AB = L_F − L_A = (FN − FD) − (AN − AD)"),
            ("Normal", "Θετικό AB σημαίνει ότι η συνέχιση της μάθησης μείωσε την απώλεια που συνδέεται με τη διαταραχή σε σχέση με το matched nominal reference· δεν σημαίνει από μόνο του ότι η Adaptive-Disturbed απόδοση υπερέβη την ονομαστική."),
            ("Normal", "Για το RQ3, μετά από ισοβαρή συνένωση των δύο layouts μέσα σε κάθε root, έστω Nₖ η Adaptive-Nominal και Dₖ η Adaptive-Disturbed μέση reward ανά πραγματική αλληλεπίδραση στο παράθυρο k. Ο κατευθυνόμενος gap είναι gₖ = Nₖ − Dₖ. Ένα παράθυρο βρίσκεται εντός της κύριας ανοχής όταν gₖ ≤ 0,10 — όχι όταν |gₖ| ≤ 0,10. Σταθερή ανάκαμψη απαιτεί δύο διαδοχικά παράθυρα εντός ανοχής. recovery_time είναι το τέλος του πρώτου παραθύρου του πρώτου τέτοιου ζεύγους και confirmation_time το τέλος του δεύτερου. Αν δεν εμφανιστεί τέτοιο ζεύγος έως τις 256 αλληλεπιδράσεις, το root είναι right-censored και recovery_time = null."),
        ),
    ),
    Block(
        "3.8 Αναπαραγωγιμότητα",
        (
            ("Normal", "Ο αριθμός των 12 roots δεν επιλέχθηκε εκ των υστέρων από τα τελικά αποτελέσματα. Στο προγενέστερο DEVELOPMENT sizing εξετάστηκαν υποψήφια πλήθη 12, 16, 20 και 24 και επιλέχθηκε μηχανικά το μικρότερο πλήθος που έδινε Student-t 95% half-width μικρότερο από 0,20 τόσο για το Phase-A AUC όσο και για το Phase-B adaptation benefit. Το κριτήριο ικανοποιήθηκε ήδη στα 12 roots, όπου το μέγιστο half-width ήταν 0,1428."),
            ("Normal", "Για οποιαδήποτε συνεχή root-level ποσότητα q_r, αφού πρώτα ληφθεί ο ισοβαρής μέσος των δύο layouts μέσα στο ίδιο root, το αναφερόμενο pointwise διάστημα 95% έχει τη μορφή q̄ ± t_(0,975;n−1)·s/√n, με n τον πραγματικό αριθμό διαθέσιμων ανεξάρτητων roots. Για n=12 χρησιμοποιείται η προδηλωμένη κρίσιμη τιμή t=2,201. Δεν ορίστηκε οικογένεια p-values ή διόρθωση πολλαπλών συγκρίσεων."),
            ("Normal", "Οι αναλογίες ανάκαμψης x/12 είναι περιγραφικές root-level αναλογίες με πάντοτε εμφανή παρονομαστή. Δεν προστίθεται εκ των υστέρων ξεχωριστό binomial confidence interval ή significance test που δεν ανήκε στο παγωμένο statistical plan· η αβεβαιότητα της ανάκαμψης τεκμηριώνεται με τα root-level trajectories, τις προκαθορισμένες sensitivity ανοχές και τις δηλωμένες root-paired συγκρίσεις."),
        ),
    ),
    Block(
        "Κεφάλαιο 4 — Αρχιτεκτονική και Υλοποίηση του Συστήματος",
        (("Normal", "Οι δύο τελικές layouts δημιουργήθηκαν με generation seeds 57001 και 57002. Οι 12 τελικές roots χρησιμοποιούν έξι ανεξάρτητα, προκαθορισμένα seed streams: initialization 71001–71012, exploration 72001–72012, scenario 73001–73012, environment 74001–74012, action-disturbance 75001–75012 και observation-disturbance 76001–76012. Η αντιστοίχιση είναι κατά root index, π.χ. r01 χρησιμοποιεί τις τιμές που λήγουν σε 001 και r12 αυτές που λήγουν σε 012. Οι ίδιες ταυτότητες χρησιμοποιούνται στις matched συγκρίσεις όπου το πρωτόκολλο απαιτεί pairing."),),
    ),
    Block(
        "6.6 Πρακτική αποτίμηση",
        (
            ("Normal", "Το bounded tuning ήταν ίσης ευκαιρίας αλλά όχι εξαντλητικό. Επομένως οι συγκρίσεις — ιδίως για DQN και PPO — αφορούν τις προδηλωμένες έξι υποψήφιες ρυθμίσεις ανά μέθοδο, το συγκεκριμένο 7×7 GridWorld και το budget των 8.192 αλληλεπιδράσεων. Δεν τεκμηριώνουν ότι διαφορετικό search space, μεγαλύτερο budget ή διαφορετική νευρωνική αρχιτεκτονική δεν θα άλλαζε τη σχετική εικόνα."),
            ("Normal", "Η ανάκαμψη είναι επίσης συγκεκριμένη operationalization της ανθεκτικότητας: μετράται σε ορίζοντα 256 αλληλεπιδράσεων, με παράθυρα 32 αλληλεπιδράσεων, κατευθυνόμενη ανοχή 0,10 και σταθερότητα δύο παραθύρων. Ένα root που δεν ικανοποιεί το κριτήριο έως το 256 χαρακτηρίζεται μόνο ως μη ανακτημένο μέσα σε αυτόν τον πεπερασμένο ορίζοντα· δεν αποδεικνύεται ότι δεν θα ανακάμψει αργότερα."),
        ),
    ),
    Block(
        "Α.2 Τελικές διατάξεις και ανεξάρτητες επαναλήψεις",
        (("Normal", "Το DEVELOPMENT tuning είχε έξι προκαθορισμένες υποψήφιες ρυθμίσεις ανά μέθοδο και χρησιμοποιούσε τις ίδιες 3 tuning-only roots × 2 development layouts. Τα ζεύγη παραμέτρων ήταν: Q-Learning/SARSA learning rate {0,1, 0,2, 0,4} × ε {0,05, 0,10}; Dyna-Q+ planning steps {5,10,20} × κ {0,0005,0,001}; DQN learning rate {0,0003,0,001,0,003} × target-update interval {128,256}; PPO learning rate {0,0001,0,0003,0,001} × n_epochs {5,10}. Η τελική επιλογή έγινε πριν από το final reserve και έδωσε q-c06, sarsa-c06, dqn-c05, ppo-c06 και dyna-c03."),),
    ),
    Block(
        "Α.3 Ορισμοί των διαταραχών",
        (("Normal", "Generation seeds των τελικών layouts: gw-l1-final-a = 57001 και gw-l1-final-b = 57002. Για root r_i, i=01,…,12, τα seed streams είναι initialization = 71000+i, exploration = 72000+i, scenario = 73000+i, environment = 74000+i, action disturbance = 75000+i και observation disturbance = 76000+i. Η 12-root επιλογή προέκυψε από το προδηλωμένο sizing rule: μικρότερο candidate count με Student-t 95% half-width <0,20 για Phase-A AUC και Phase-B adaptation benefit· στα 12 roots το μέγιστο half-width ήταν 0,1428."),),
    ),
    Block(
        "Α.4 Προϋπολογισμοί και συμβόλαιο ανάκαμψης",
        (("Normal", "Action failure: με πιθανότητα 0,15 η intended action εκτελείται ως no-op. Η ground-truth θέση παραμένει ίδια, collision=false και η reward είναι η κανονική step reward −0,1. Η μετάβαση παραμένει έγκυρη εμπειρία για τους Adaptive κλάδους. Observation corruption: με πιθανότητα 0,05 η delivered observation δειγματοληπτείται ομοιόμορφα από τα μη εμποδισμένα κελιά, εξαιρώντας πάντοτε την τρέχουσα ground-truth θέση. Η πραγματική μετάβαση και η reward δεν αλλάζουν· ο agent λαμβάνει τη delivered observation. Το goal δεν αποκλείεται ειδικά από το support."),),
    ),
)

SENTINELS = tuple(block.paragraphs[0][1] for block in BLOCKS)
EXPECTED_INSERTED_PARAGRAPHS = sum(len(block.paragraphs) for block in BLOCKS)
EXPECTED_INSERTED_HEADINGS = sum(1 for block in BLOCKS for style, _ in block.paragraphs if style.startswith("Heading"))


def _find_exact(doc: _Document, text: str):
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"T-715 audit hardening anchor expected exactly once: {text!r}; found {len(matches)}")
    return matches[0]


def apply(doc: _Document) -> dict[str, object]:
    sentinel_presence = {text: sum(1 for p in doc.paragraphs if p.text.strip() == text) for text in SENTINELS}
    present = [text for text, count in sentinel_presence.items() if count]
    if present:
        if all(count == 1 for count in sentinel_presence.values()):
            return {
                "applied": False,
                "already_applied": True,
                "inserted_paragraph_count": 0,
                "inserted_heading_count": 0,
                "sentinel_count": len(SENTINELS),
            }
        raise RuntimeError(f"T-715 audit hardening found a partial/duplicate prior insertion: {sentinel_presence}")

    anchors = [(block, _find_exact(doc, block.anchor)) for block in BLOCKS]
    for block, anchor in anchors:
        for style, text in block.paragraphs:
            paragraph = anchor.insert_paragraph_before(text, style=style)
            paragraph.paragraph_format.keep_together = False

    final_presence = {text: sum(1 for p in doc.paragraphs if p.text.strip() == text) for text in SENTINELS}
    if not all(count == 1 for count in final_presence.values()):
        raise RuntimeError(f"T-715 audit hardening sentinel verification failed: {final_presence}")

    return {
        "applied": True,
        "already_applied": False,
        "inserted_paragraph_count": EXPECTED_INSERTED_PARAGRAPHS,
        "inserted_heading_count": EXPECTED_INSERTED_HEADINGS,
        "sentinel_count": len(SENTINELS),
    }


def apply_to_path(path: Path) -> dict[str, object]:
    doc = Document(path)
    before = len(doc.paragraphs)
    result = apply(doc)
    doc.save(path)
    result = dict(result)
    result["paragraph_count_before"] = before
    result["paragraph_count_after"] = len(Document(path).paragraphs)
    return result
