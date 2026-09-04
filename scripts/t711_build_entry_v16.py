#!/usr/bin/env python3
"""T-711A front-matter pagination hardening.

Normalizes explicit trailing page-break paragraphs before long generated front-matter
lists. A page break attached to the following Heading 1 is stable when the preceding
list grows to the bottom of a page; a standalone break paragraph can itself flow to the
next page and create a blank page. Scientific content and media are untouched.
"""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import t711_build_entry_v15 as v15


t711 = v15.t711
_previous_build = v15._build

TARGET_HEADINGS = {
    "Κατάλογος Σχημάτων",
    "Κατάλογος Πινάκων",
    "Γλωσσάριο και Ακρωνύμια",
}


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _is_explicit_page_break(paragraph) -> bool:
    if paragraph.text:
        return False
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def _build(output: Path, qa_output: Path) -> None:
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))

    doc = Document(output)
    removed = []
    for heading_text in TARGET_HEADINGS:
        paragraphs = doc.paragraphs
        matches = [i for i, p in enumerate(paragraphs) if p.text == heading_text]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one front-matter heading {heading_text!r}, found {len(matches)}")
        index = matches[0]
        heading = paragraphs[index]
        heading.paragraph_format.page_break_before = True
        if index > 0 and _is_explicit_page_break(paragraphs[index - 1]):
            prior = paragraphs[index - 1]
            prior._element.getparent().remove(prior._element)
            removed.append(heading_text)

    doc.save(output)
    final_doc = Document(output)
    remaining_unstable = []
    paragraphs = final_doc.paragraphs
    for index, paragraph in enumerate(paragraphs[:-1]):
        if not _is_explicit_page_break(paragraph):
            continue
        if paragraphs[index + 1].text in TARGET_HEADINGS:
            remaining_unstable.append(paragraphs[index + 1].text)

    final_sha = _sha256(output)
    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": len(final_doc.paragraphs),
            "post_synthesis_paragraph_count": len(final_doc.paragraphs),
            "t711a_hardening_version": 16,
            "front_matter_pagebreak_targets": sorted(TARGET_HEADINGS),
            "front_matter_explicit_breaks_removed": sorted(removed),
            "unstable_front_matter_pagebreak_residue": remaining_unstable,
            "blank_page_pagination_repair": (
                set(removed) == TARGET_HEADINGS and not remaining_unstable
            ),
            "scientific_values_modified": False,
            "registered_asset_bytes_modified": False,
            "final_visual_qa_required": True,
        }
    )
    if set(removed) != TARGET_HEADINGS or remaining_unstable:
        report["status"] = "fail"

    qa_output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if report.get("status") != "pass":
        raise RuntimeError(f"T-711A v16 pagination hardening failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
