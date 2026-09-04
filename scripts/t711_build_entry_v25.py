#!/usr/bin/env python3
"""T-715 final application-illustration composition layer.

Adds exactly two deterministic DEVELOPMENT-only PySide6 screenshots to Chapter 4:
Phase A GridWorld monitoring and the matched Frozen/Adaptive Phase-B presentation.
They are implementation illustrations, not scientific evidence or quantitative sources.
The pre-existing 24 registered T-613 scientific figures and the unnumbered results
synthesis remain byte-for-byte unchanged and keep their numbering/list semantics.
"""

from __future__ import annotations

import hashlib
import json
import os
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import t711_build_entry_v24 as v24


t711 = v24.t711
_previous_build = t711.builder.build
ROOT = Path(__file__).resolve().parents[1]

PHASE_A_NAME = "03-run-phase-a-1440x900.png"
PHASE_B_NAME = "04-run-phase-b-frozen-adaptive-1440x900.png"
EXPECTED_CAPTURE_NAMES = (PHASE_A_NAME, PHASE_B_NAME)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        result: dict[str, str] = {}
        for name in sorted(n for n in archive.namelist() if n.startswith("word/media/")):
            result[name] = hashlib.sha256(archive.read(name)).hexdigest()
        return result


def _find_paragraph(doc: Document, prefix: str):
    matches = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"T-715 application screenshot anchor expected once for {prefix!r}, found {len(matches)}")
    return matches[0]


def _insert_screenshot_after(doc: Document, anchor, image_path: Path, caption: str, alt_text: str) -> None:
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(6)
    picture_paragraph.paragraph_format.space_after = Pt(3)
    picture_paragraph.paragraph_format.keep_with_next = True
    run = picture_paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(5.95))
    shape._inline.docPr.set("descr", alt_text)

    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_paragraph.paragraph_format.keep_together = True
    caption_run = caption_paragraph.add_run(caption)
    t711.builder.set_run_font(caption_run, size=9.5)

    anchor._p.addnext(picture_paragraph._p)
    picture_paragraph._p.addnext(caption_paragraph._p)


def _validated_screenshot_inputs() -> tuple[Path, Path, dict]:
    directory = Path(os.environ.get("T715_UI_SCREENSHOT_DIR", ROOT / "artifacts" / "t715-ui")).resolve()
    manifest_path = directory / "manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(f"T-715 deterministic UI manifest missing: {manifest_path}")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    required_manifest_state = (
        manifest.get("development_fixture_created_only") is True
        and manifest.get("scientific_jobs_executed") == 0
        and manifest.get("environment_steps_executed") == 0
        and manifest.get("final_reserve_accessed") is False
        and manifest.get("final_experiment_authorized") is False
    )
    if not required_manifest_state:
        raise RuntimeError(f"T-715 refuses non-DEVELOPMENT UI capture manifest: {manifest}")

    records = {record.get("file"): record for record in manifest.get("screenshots", [])}
    paths: list[Path] = []
    for name in EXPECTED_CAPTURE_NAMES:
        path = directory / name
        record = records.get(name)
        if not path.exists() or record is None:
            raise RuntimeError(f"T-715 required deterministic UI screenshot missing: {name}")
        if int(record.get("width", 0)) != 1440 or int(record.get("height", 0)) != 900:
            raise RuntimeError(f"T-715 unexpected screenshot dimensions for {name}: {record}")
        actual = _sha256(path)
        if actual != record.get("sha256"):
            raise RuntimeError(f"T-715 screenshot hash mismatch for {name}")
        paths.append(path)
    return paths[0], paths[1], manifest


def _build(output: Path, qa_output: Path) -> None:
    phase_a, phase_b, capture_manifest = _validated_screenshot_inputs()
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    if report.get("status") != "pass":
        raise RuntimeError(f"T-715 base reader composition failed before application screenshots: {report}")

    media_before = _media_hashes(output)
    doc = Document(output)
    inline_before = len(doc.inline_shapes)

    phase_a_anchor = _find_paragraph(
        doc,
        "Η εφαρμογή οργανώνεται στις κύριες ενότητες",
    )
    _insert_screenshot_after(
        doc,
        phase_a_anchor,
        phase_a,
        "Στιγμιότυπο εφαρμογής 1 — Ενδεικτική προβολή της Φάσης A και του GridWorld κατά την ονομαστική μάθηση. Η εικόνα είναι deterministic DEVELOPMENT capture και δεν αποτελεί πηγή ποσοτικών αποτελεσμάτων.",
        "Desktop εφαρμογή PySide6 στη Φάση A με προβολή του GridWorld και της τρέχουσας κατάστασης εκτέλεσης. DEVELOPMENT capture μόνο για επεξήγηση της υλοποίησης.",
    )

    phase_b_anchor = _find_paragraph(
        doc,
        "Η ενότητα Run της εφαρμογής είναι η πιο άμεση σύνδεση",
    )
    _insert_screenshot_after(
        doc,
        phase_b_anchor,
        phase_b,
        "Στιγμιότυπο εφαρμογής 2 — Ενδεικτική προβολή της Φάσης B με παράλληλη παρουσίαση Frozen και Adaptive κλάδων στο GridWorld. Η εικόνα εξηγεί τη λειτουργία της εφαρμογής και δεν χρησιμοποιείται ως επιστημονικό τεκμήριο.",
        "Desktop εφαρμογή PySide6 στη Φάση B με παράλληλη απεικόνιση Frozen και Adaptive GridWorld. DEVELOPMENT capture μόνο για επεξήγηση της υλοποίησης.",
    )

    doc.save(output)
    final_doc = Document(output)
    media_after = _media_hashes(output)
    inline_after = len(final_doc.inline_shapes)

    # Every pre-existing DOCX media part must survive byte-for-byte. The only allowed
    # additions are the two application screenshots.
    preserved = all(media_after.get(name) == digest for name, digest in media_before.items())
    new_media = sorted(name for name in media_after if name not in media_before)
    capture_hashes = {_sha256(phase_a), _sha256(phase_b)}
    new_media_hashes = {media_after[name] for name in new_media}

    app_captions = [
        paragraph.text
        for paragraph in final_doc.paragraphs
        if paragraph.text.startswith("Στιγμιότυπο εφαρμογής ")
    ]
    final_sha = _sha256(output)
    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": len(final_doc.paragraphs),
            "post_synthesis_paragraph_count": len(final_doc.paragraphs),
            "t715_application_screenshots_added": True,
            "t715_application_screenshot_count": len(app_captions),
            "t715_application_screenshot_names": list(EXPECTED_CAPTURE_NAMES),
            "t715_application_screenshot_sha256": {
                PHASE_A_NAME: _sha256(phase_a),
                PHASE_B_NAME: _sha256(phase_b),
            },
            "t715_application_capture_purpose": capture_manifest.get("purpose"),
            "t715_application_capture_development_only": True,
            "t715_application_capture_scientific_jobs_executed": 0,
            "t715_application_capture_environment_steps_executed": 0,
            "t715_application_capture_final_reserve_accessed": False,
            "t715_application_screenshots_scientific_evidence": False,
            "t715_preexisting_docx_media_preserved": preserved,
            "t715_new_docx_media_count": len(new_media),
            "t715_new_docx_media_hashes_match_screenshots": new_media_hashes == capture_hashes,
            "t715_inline_shape_count_before_app_screenshots": inline_before,
            "t715_final_inline_shape_count": inline_after,
            "t715_final_image_alt_text_count": report.get("image_alt_text_count", 0) + 2,
            "scientific_values_modified": False,
            "registered_asset_bytes_modified": False,
            "final_visual_qa_required": True,
        }
    )

    hardening_ok = (
        preserved
        and len(new_media) == 2
        and new_media_hashes == capture_hashes
        and inline_before == 25
        and inline_after == 27
        and len(app_captions) == 2
        and report.get("inserted_asset_count") == 24
        and report.get("registered_figure_count_preserved") is True
        and report.get("scientific_values_modified") is False
        and report.get("registered_asset_bytes_modified") is False
    )
    if not hardening_ok:
        report["status"] = "fail"

    qa_output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if report.get("status") != "pass":
        raise RuntimeError(f"T-715 application screenshot composition failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
