#!/usr/bin/env python3
"""T-711A academic structure, terminology and appendix hardening.

This layer is presentation/editorial only. It preserves protocol-v2.1, T-611/T-612
scientific evidence, all registered T-613 media bytes, numerical results, estimands,
intervals and censoring decisions. It improves the reader-facing thesis structure and
self-containment after the source-aware v13 build.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

import t711_build_entry_v13 as v13


t711 = v13.t711
_previous_build = t711.builder.build

HEADING_REPLACEMENTS = {
    "3.2 GridWorld ως ελεγχόμενο πειραματικό testbed": "3.2 GridWorld ως ελεγχόμενο πειραματικό περιβάλλον δοκιμών",
    "3.3 Τελικές μέθοδοι και fairness boundaries": "3.3 Τελικές μέθοδοι και όρια δίκαιης σύγκρισης",
    "3.4 Phase A: ανεξάρτητη ονομαστική μάθηση και checkpoints": "3.4 Φάση A: ανεξάρτητη ονομαστική μάθηση και σημεία ελέγχου",
    "3.5 Phase B: matched FN/FD/AN/AD design": "3.5 Φάση B: αντιστοιχισμένος σχεδιασμός FN/FD/AN/AD",
    "3.6 Μηχανισμοί αβεβαιότητας και disturbance conditions": "3.6 Μηχανισμοί αβεβαιότητας και συνθήκες διαταραχής",
    "3.7 RQ1: estimands ονομαστικής μάθησης": "3.7 RQ1: μεγέθη εκτίμησης ονομαστικής μάθησης",
    "3.8 RQ2: resilience και adaptation-benefit estimand": "3.8 RQ2: ανθεκτικότητα και μέγεθος εκτίμησης οφέλους προσαρμογής",
    "3.9 RQ3: recovery, passive windows και right-censoring": "3.9 RQ3: ανάκαμψη, παθητικά χρονικά παράθυρα και δεξιά λογοκρισία",
    "3.10 Roots, layouts, budgets και πειραματικός πίνακας": "3.10 Ανεξάρτητες επαναλήψεις, διατάξεις, προϋπολογισμοί και πειραματικός πίνακας",
    "3.12 Reproducibility, provenance και scientific firewall": "3.12 Αναπαραγωγιμότητα, ιχνηλασιμότητα και διαχωρισμός της επιστημονικής ανάλυσης",
    "4.3 Study ως μονάδα εκτέλεσης και provenance": "4.3 Η Study ως μονάδα εκτέλεσης και ιχνηλασιμότητας",
    "4.4 Deterministic planning και stage barriers": "4.4 Ντετερμινιστικός σχεδιασμός και φραγμοί σταδίων",
    "4.6 RNG isolation και reproducibility": "4.6 Απομόνωση γεννητριών τυχαιότητας και αναπαραγωγιμότητα",
    "4.7 Run bundles, manifests και ακεραιότητα": "4.7 Πακέτα εκτέλεσης, manifests και ακεραιότητα",
    "4.8 Evidence-v2 validation και analysis layer": "4.8 Επικύρωση Evidence-v2 και στρώμα ανάλυσης",
    "4.9 Χειρισμός αποτυχίας και το T-610 recovery path": "4.9 Χειρισμός αποτυχίας και διαδρομή ανάκαμψης του T-610",
    "4.11 Live visualization και επιστημονική παθητικότητα": "4.11 Ζωντανή οπτικοποίηση και επιστημονική παθητικότητα",
    "4.12 Execution supervision και restart safety": "4.12 Εποπτεία εκτέλεσης και ασφαλής επανεκκίνηση",
    "4.13 Final-reserve firewall και authorization": "4.13 Προστασία του τελικού αποθεματικού και εξουσιοδότηση",
    "4.14 Παραγωγή αποτελεσμάτων και thesis assets": "4.14 Παραγωγή αποτελεσμάτων και υλικού της διπλωματικής",
    "5.2.2 Αποτελεσματικότητα μάθησης κατά μήκος του budget": "5.2.2 Αποτελεσματικότητα μάθησης κατά μήκος του διαθέσιμου προϋπολογισμού",
    "5.3.1 Συνολική εικόνα ανά condition": "5.3.1 Συνολική εικόνα ανά συνθήκη",
    "5.3.2 Persistent action-remap cycle": "5.3.2 Μόνιμη κυκλική ανααντιστοίχιση ενεργειών",
    "5.3.3 Persistent action-remap swap": "5.3.3 Μόνιμη εναλλαγή δεξιάς–κάτω ενέργειας",
    "5.3.4 Action failure 15%": "5.3.4 Αποτυχία ενέργειας 15%",
    "5.3.5 Observation corruption 5%": "5.3.5 Αλλοίωση παρατήρησης 5%",
    "5.4 RQ3 — Ανάκαμψη μετά από persistent action remapping": "5.4 RQ3 — Ανάκαμψη μετά από μόνιμη ανααντιστοίχιση ενεργειών",
    "5.4.1 Cycle remap": "5.4.1 Κυκλική ανααντιστοίχιση",
    "5.4.2 Swap remap": "5.4.2 Εναλλαγή δεξιάς–κάτω ενέργειας",
    "6.2 RQ1: τελική επίδοση και sample efficiency δεν είναι το ίδιο": "6.2 RQ1: τελική επίδοση και αποδοτικότητα δείγματος δεν είναι το ίδιο",
    "6.3 RQ2: η προσαρμογή είναι condition-dependent": "6.3 RQ2: η αποτελεσματικότητα της προσαρμογής εξαρτάται από τη συνθήκη",
    "6.5 RQ3: recovery incidence, timing και censoring": "6.5 RQ3: συχνότητα, χρόνος και λογοκρισία της ανάκαμψης",
    "6.6 Ευαισθησία στον operational ορισμό της ανάκαμψης": "6.6 Ευαισθησία στον λειτουργικό ορισμό της ανάκαμψης",
    "6.7 Σχέση με continual RL και hidden change": "6.7 Σχέση με συνεχή ενισχυτική μάθηση και κρυφή μεταβολή",
    "6.8.1 Fairness και tuning": "6.8.1 Δίκαιη σύγκριση και ρύθμιση υπερπαραμέτρων",
    "6.8.2 Checkpoint equivalence": "6.8.2 Ισοδυναμία σημείων ελέγχου",
    "6.8.3 Information leakage": "6.8.3 Διαρροή πληροφορίας",
    "6.9 Construct validity": "6.9 Εγκυρότητα εννοιολογικής κατασκευής",
    "6.10 Statistical conclusion validity": "6.10 Εγκυρότητα στατιστικών συμπερασμάτων",
    "6.11 External validity": "6.11 Εξωτερική εγκυρότητα",
    "6.12 Reproducibility validity": "6.12 Εγκυρότητα αναπαραγωγιμότητας",
    "7.7.2 Διαφορετικά budgets και learning timescales": "7.7.2 Διαφορετικοί προϋπολογισμοί και χρονικές κλίμακες μάθησης",
    "7.7.3 Ρητή ανίχνευση αλλαγής και context inference": "7.7.3 Ρητή ανίχνευση αλλαγής και εκτίμηση πλαισίου",
    "7.7.4 Specialized continual-learning mechanisms": "7.7.4 Εξειδικευμένοι μηχανισμοί συνεχούς μάθησης",
    "7.7.5 Modular model-based adaptation": "7.7.5 Αρθρωτή προσαρμογή με μοντέλο",
    "7.7.6 Διαχείριση replay υπό non-stationarity": "7.7.6 Διαχείριση replay σε μη στάσιμα περιβάλλοντα",
    "7.7.7 Πλουσιότερη μελέτη recovery": "7.7.7 Πλουσιότερη μελέτη ανάκαμψης",
    "7.7.8 Περισσότερες ανεξάρτητες roots και ιεραρχική ανάλυση": "7.7.8 Περισσότερες ανεξάρτητες επαναλήψεις και ιεραρχική ανάλυση",
    "7.7.9 Safety και constraint-aware adaptation": "7.7.9 Ασφάλεια και προσαρμογή υπό περιορισμούς",
    "Α.1 Methods και βασικές frozen ρυθμίσεις": "Α.1 Μέθοδοι και τελικές ρυθμίσεις",
    "Α.2 Final layouts και roots": "Α.2 Τελικές διατάξεις και ανεξάρτητες επαναλήψεις",
    "Α.3 Disturbance definitions": "Α.3 Ορισμοί των διαταραχών",
    "Παράρτημα Β — Πλήρη RQ αποτελέσματα και diagnostics": "Παράρτημα Β — Πλήρη αποτελέσματα και διαγνωστικά ανά ερευνητικό ερώτημα",
    "Παράρτημα Γ — Evidence, provenance και reproducibility": "Παράρτημα Γ — Τεκμήρια, ιχνηλασιμότητα και αναπαραγωγιμότητα",
    "Γ.1 Execution lineage": "Γ.1 Γραμμή εκτέλεσης",
    "Γ.2 Frozen identities": "Γ.2 Παγωμένες ταυτότητες τεκμηρίων",
    "Γ.3 Bibliography provenance": "Γ.3 Ιχνηλασιμότητα βιβλιογραφίας",
}

TEXT_REPLACEMENTS = {
    "Στο RQ1, Q-Learning, SARSA και Dyna-Q+ κατέληξαν στην ίδια τελική ονομαστική επίδοση, αλλά η Dyna-Q+ είχε σημαντικά υψηλότερη time-average επίδοση κατά μήκος του training budget. Η DQN και η PPO είχαν χαμηλότερη τελική και time-average επίδοση και μεγαλύτερη μεταξύ-root διακύμανση στο συγκεκριμένο controlled task και budget.":
        "Στο RQ1, Q-Learning, SARSA και Dyna-Q+ κατέληξαν στην ίδια τελική ονομαστική επίδοση, ενώ η εκτιμώμενη time-average επίδοση της Dyna-Q+ ήταν υψηλότερη κατά μήκος του διαθέσιμου προϋπολογισμού αλληλεπιδράσεων. Η DQN και η PPO είχαν χαμηλότερη τελική και time-average επίδοση και μεγαλύτερη διακύμανση μεταξύ ανεξάρτητων επαναλήψεων στο συγκεκριμένο ελεγχόμενο περιβάλλον και budget.",
    "Η σύνθεση αυτή δεν δημιουργεί ενιαία κατάταξη. Η Dyna-Q+ είναι η ισχυρότερη ως προς time-average nominal learning, ενώ Q-Learning και SARSA έχουν την πιο συνεπή recovery incidence στις persistent remap conditions. Η ίδια η online adaptation μπορεί να είναι ωφέλιμη, ουδέτερη ή επιβλαβής ανάλογα με τον disturbance mechanism. Αυτές οι διαφοροποιήσεις αποτελούν το αντικείμενο της ερμηνείας στο επόμενο κεφάλαιο.":
        "Η σύνθεση αυτή δεν δημιουργεί ενιαία κατάταξη. Η Dyna-Q+ είχε την υψηλότερη καταγεγραμμένη time-average ονομαστική επίδοση, ενώ οι Q-Learning και SARSA εμφάνισαν τις υψηλότερες αναλογίες σταθερής ανάκαμψης στις δύο μόνιμες συνθήκες ανααντιστοίχισης. Η online προσαρμογή μπορεί να είναι ωφέλιμη, ουδέτερη ή επιβαρυντική ανάλογα με τον μηχανισμό της διαταραχής. Αυτές οι διαφοροποιήσεις αποτελούν το αντικείμενο της ερμηνείας στο επόμενο κεφάλαιο.",
    "Η απάντηση στο RQ1 είναι επομένως διπλή. Ως προς το τελικό nominal level, οι τρεις tabular/planning μέθοδοι δεν διαχωρίστηκαν. Ως προς τη learning efficiency, η Dyna-Q+ είχε το ισχυρότερο αποτέλεσμα. Το εύρημα παραμένει περιορισμένο στο συγκεκριμένο controlled task, budget και frozen configuration set.":
        "Η απάντηση στο RQ1 είναι επομένως διπλή. Ως προς το τελικό ονομαστικό επίπεδο, οι τρεις tabular/planning μέθοδοι δεν διαχωρίστηκαν. Ως προς την αποτελεσματικότητα μάθησης, η Dyna-Q+ είχε την υψηλότερη εκτιμώμενη time-average απόδοση. Το εύρημα παραμένει περιορισμένο στο συγκεκριμένο ελεγχόμενο περιβάλλον, στον διαθέσιμο προϋπολογισμό και στις προκαθορισμένες τελικές ρυθμίσεις.",
}

BIBLIOGRAPHY = {
    1: "[1] Richard S. Sutton and Andrew G. Barto, Reinforcement Learning: An Introduction, 2nd ed. MIT Press, 2018. ISBN 9780262039246.",
    2: "[2] Christian Steinparz, Thomas Schmied, Fabian Paischer, Marius-Constantin Dinu, Vihang Patil, Angela Bitto-Nemling, Hamid Eghbal-zadeh, and Sepp Hochreiter, “Reactive Exploration to Cope with Non-Stationarity in Lifelong Reinforcement Learning,” in Proceedings of the Conference on Lifelong Learning Agents, PMLR, vol. 199, 2022. [Online]. Available: https://proceedings.mlr.press/v199/steinparz22a.html",
    3: "[3] Shibhansh Dohare, J. Fernando Hernandez-Garcia, Qingfeng Lan, Parash Rahman, A. Rupam Mahmood, and Richard S. Sutton, “Loss of plasticity in deep continual learning,” Nature, vol. 632, pp. 768–774, 2024, doi: 10.1038/s41586-024-07711-7.",
    4: "[4] Evgenii Nikishin, Max Schwarzer, Pierluca D’Oro, Pierre-Luc Bacon, and Aaron Courville, “The Primacy Bias in Deep Reinforcement Learning,” in Proceedings of the 39th International Conference on Machine Learning, PMLR, vol. 162, pp. 16828–16847, 2022.",
    5: "[5] Andrew Patterson, Samuel Neumann, Martha White, and Adam White, “Empirical Design in Reinforcement Learning,” Journal of Machine Learning Research, vol. 25, no. 318, pp. 1–63, 2024.",
    6: "[6] Peter Henderson, Riashat Islam, Philip Bachman, Joelle Pineau, Doina Precup, and David Meger, “Deep Reinforcement Learning That Matters,” in Proceedings of AAAI-18, 2018.",
    7: "[7] Pouya Hamadanian, Arash Nasr-Esfahany, Malte Schwarzkopf, Siddhartha Sen, and Mohammad Alizadeh, “Online Reinforcement Learning in Non-Stationary Context-Driven Environments,” in Proceedings of the International Conference on Learning Representations (ICLR), 2025. [Online]. Available: https://proceedings.iclr.cc/paper_files/paper/2025/hash/fb21dae9e8710a272c0a0ca848f71553-Abstract-Conference.html",
    8: "[8] Safa Alver, Ali Rahimi-Kalahroudi, and Doina Precup, “Partial Models for Building Adaptive Model-Based Reinforcement Learning Agents,” in Proceedings of the 3rd Conference on Lifelong Learning Agents, PMLR, vol. 274, pp. 959–977, 2025. [Online]. Available: https://proceedings.mlr.press/v274/alver25a.html",
    9: "[9] Richard S. Sutton and Andrew G. Barto, “Q-learning: Off-policy TD Control,” in Reinforcement Learning: An Introduction, 2nd ed. MIT Press, 2018, pp. 131–135.",
    10: "[10] Volodymyr Mnih, Koray Kavukcuoglu, David Silver, Alex Graves, Ioannis Antonoglou, Daan Wierstra, and Martin Riedmiller, “Playing Atari with Deep Reinforcement Learning,” preprint, 2013.",
    11: "[11] William Fedus, Prajit Ramachandran, Rishabh Agarwal, Yoshua Bengio, Hugo Larochelle, Mark Rowland, and Will Dabney, “Revisiting Fundamentals of Experience Replay,” in Proceedings of the 37th International Conference on Machine Learning, PMLR, vol. 119, pp. 3061–3071, 2020.",
    12: "[12] John Schulman, Filip Wolski, Prafulla Dhariwal, Alec Radford, and Oleg Klimov, “Proximal Policy Optimization Algorithms,” arXiv:1707.06347, 2017.",
    13: "[13] Logan Engstrom, Andrew Ilyas, Shibani Santurkar, Dimitris Tsipras, Firdaus Janoos, Larry Rudolph, and Aleksander Madry, “Implementation Matters in Deep Policy Gradients: A Case Study on PPO and TRPO,” arXiv:2005.12729, 2020.",
    14: "[14] Chaofan Pan, Xin Yang, Yanhua Li, Wei Wei, Tianrui Li, Bo An, and Jiye Liang, “A Survey of Continual Reinforcement Learning,” arXiv:2506.21872, 2025.",
    15: "[15] Zihe Liu, Deep Reinforcement Learning in Non-stationary Environments, Ph.D. dissertation, University of Technology Sydney, 2024. [Online]. Available: https://opus.lib.uts.edu.au/handle/10453/186408",
    16: "[16] Rishabh Agarwal, Max Schwarzer, Pablo Samuel Castro, Aaron Courville, and Marc G. Bellemare, “Deep Reinforcement Learning at the Edge of the Statistical Precipice,” NeurIPS, 2021, arXiv:2108.13264.",
    17: "[17] Fabio Pardo, Arash Tavakoli, Vitaly Levdik, and Petar Kormushev, “Time Limits in Reinforcement Learning,” in Proceedings of the 35th International Conference on Machine Learning, PMLR, vol. 80, pp. 4045–4054, 2018.",
}

GLOSSARY_TERMS = {
    "Actual environment interaction", "Adaptation benefit", "Adaptive regime",
    "Censoring / Right-censoring", "Checkpoint", "Continual / continued learning",
    "Disturbance-associated loss", "Dyna-Q+", "Equal-layout reduction", "Frozen regime",
    "GridWorld", "Hidden change", "Independent root", "Non-stationarity",
    "Pointwise interval", "Recovery", "Recovery incidence / recovered proportion",
    "Recovery time", "Restricted recovery delay", "Resilience / Ανθεκτικότητα",
    "Stable recovery", "Study", "StudyRecipe", "Time-average return",
}


def _set_text(paragraph: Paragraph, value: str) -> None:
    paragraph.text = value
    for run in paragraph.runs:
        v13.v12.t711.builder.set_run_font(run, size=11)


def _find(doc: Document, value: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text == value:
            return paragraph
    raise RuntimeError(f"Required paragraph not found: {value!r}")


def _remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter() if node.tag == qn("w:t"))


def _move_block(doc: Document, start_text: str, end_text: str, target_text: str) -> None:
    body = doc.element.body
    children = list(body)
    def index_of(value: str) -> int:
        for idx, element in enumerate(children):
            if element.tag == qn("w:p") and _element_text(element) == value:
                return idx
        raise RuntimeError(f"Body marker missing: {value!r}")
    start = index_of(start_text)
    end = index_of(end_text)
    block = children[start:end]
    for element in block:
        body.remove(element)
    target_element = next(
        element for element in body
        if element.tag == qn("w:p") and _element_text(element) == target_text
    )
    pos = list(body).index(target_element)
    for offset, element in enumerate(block):
        body.insert(pos + offset, element)


def _field_result_lines(paragraph: Paragraph, lines: list[str]) -> None:
    field_run = None
    for run in paragraph._p.findall(qn("w:r")):
        if run.find(qn("w:instrText")) is not None:
            field_run = run
            break
    if field_run is None:
        raise RuntimeError(f"Field run missing in paragraph: {paragraph.text!r}")
    children = list(field_run)
    separator = None
    end = None
    for node in children:
        if node.tag != qn("w:fldChar"):
            continue
        field_type = node.get(qn("w:fldCharType"))
        if field_type == "separate":
            separator = node
        elif field_type == "end":
            end = node
    if separator is None or end is None:
        raise RuntimeError("Field result boundaries missing")
    start_index = list(field_run).index(separator) + 1
    end_index = list(field_run).index(end)
    for node in list(field_run)[start_index:end_index]:
        field_run.remove(node)
    insert_at = list(field_run).index(end)
    for idx, line in enumerate(lines):
        text = OxmlElement("w:t")
        text.set(qn("xml:space"), "preserve")
        text.text = line
        field_run.insert(insert_at, text)
        insert_at += 1
        if idx != len(lines) - 1:
            br = OxmlElement("w:br")
            field_run.insert(insert_at, br)
            insert_at += 1


def _add_seq_field(paragraph: Paragraph, label: str, display: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" SEQ {label} \\* ARABIC "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    v13.v12.t711.builder.set_run_font(run, size=10)


def _insert_config_table(doc: Document) -> None:
    body = doc.element.body
    a1 = _find(doc, "Α.1 Μέθοδοι και τελικές ρυθμίσεις")
    a2 = _find(doc, "Α.2 Τελικές διατάξεις και ανεξάρτητες επαναλήψεις")
    children = list(body)
    start = children.index(a1._p) + 1
    end = children.index(a2._p)
    for element in children[start:end]:
        body.remove(element)

    intro = doc.add_paragraph(style="Normal")
    _set_text(
        intro,
        "Οι τελικές ρυθμίσεις προέρχονται απευθείας από το παγωμένο protocol-v2.1 "
        "και συνοψίζονται εδώ ώστε η βασική παραμετροποίηση να είναι αναγνώσιμη "
        "χωρίς εξωτερική αναζήτηση στο repository.",
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Μέθοδος", "ID ρύθμισης", "Ρυθμός μάθησης", "Κύριες τελικές ρυθμίσεις"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    rows = [
        ("Q-Learning", "q-c06", "0,4", "γ=0,95 · ε=0,10 · Q₀=0"),
        ("SARSA", "sarsa-c06", "0,4", "γ=0,95 · ε=0,10 · Q₀=0"),
        ("DQN", "dqn-c05", "0,003", "γ=0,95 · buffer=16.384 · batch=32 · learning starts=128 · target update=128 · net 64×64 · ε: 1,00→0,05"),
        ("PPO", "ppo-c06", "0,001", "γ=0,95 · n_steps=128 · batch=64 · epochs=10 · clip=0,20 · GAE λ=0,95 · π/V 64×64"),
        ("Dyna-Q+", "dyna-c03", "0,2", "γ=0,95 · ε=0,10 · Q₀=0 · planning steps=10 · κ=0,0005"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    v13.v12.t711.builder.set_run_font(
                        run, size=8.5, bold=(row_index == 0)
                    )
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)

    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("Πίνακας ")
    v13.v12.t711.builder.set_run_font(run, size=10, bold=True)
    _add_seq_field(caption, "Πίνακας", "3")
    run = caption.add_run(" — Τελικές ρυθμίσεις των πέντε μεθόδων στο protocol-v2.1.")
    v13.v12.t711.builder.set_run_font(run, size=10)

    for element in (intro._p, table._tbl, caption._p):
        body.remove(element)
    pos = list(body).index(a2._p)
    for offset, element in enumerate((intro._p, table._tbl, caption._p)):
        body.insert(pos + offset, element)


def _build(output: Path, qa_output: Path) -> None:
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    media_before = v13.v12._media_hashes(output)
    doc = Document(output)

    heading_replacements = 0
    for paragraph in doc.paragraphs:
        replacement = HEADING_REPLACEMENTS.get(paragraph.text)
        if replacement is not None:
            paragraph.text = replacement
            heading_replacements += 1

    wording_replacements = 0
    for paragraph in doc.paragraphs:
        replacement = TEXT_REPLACEMENTS.get(paragraph.text)
        if replacement is not None:
            _set_text(paragraph, replacement)
            wording_replacements += 1

        if paragraph.text.startswith(
            "Το παρόν κεφάλαιο παρουσιάζει αποκλειστικά τα αποτελέσματα της προδηλωμένης ανάλυσης του protocol-v2.1 πάνω στο evidence set"
        ):
            _set_text(
                paragraph,
                "Το παρόν κεφάλαιο παρουσιάζει αποκλειστικά τα αποτελέσματα της "
                "προδηλωμένης ανάλυσης του protocol-v2.1. Ως ανεξάρτητη μονάδα "
                "χρησιμοποιείται η ανεξάρτητη επανάληψη (root): οι δύο τελικές "
                "διατάξεις συνδυάζονται με ίσο βάρος μέσα σε κάθε root και τα αριθμητικά "
                "μεγέθη συνοδεύονται από σημειακά διαστήματα Student-t 95% με βάση το "
                "πραγματικό πλήθος των ανεξάρτητων roots.",
            )
        elif paragraph.text.startswith(
            "Όλες οι προγραμματισμένες μονάδες ήταν διαθέσιμες για τα RQ1, RQ2 και RQ3. Η accepted replacement execution"
        ):
            _set_text(
                paragraph,
                "Όλες οι προγραμματισμένες μονάδες ήταν διαθέσιμες για τα RQ1, RQ2 και "
                "RQ3. Η τελική ανάλυση βασίζεται αποκλειστικά στην επικυρωμένη εκτέλεση "
                "αντικατάστασης· η αρχική αποτυχημένη εκτέλεση δεν συμμετέχει σε κανένα "
                "αποτέλεσμα. Οι λεπτομέρειες της γραμμής εκτέλεσης και οι αναγνωριστικές "
                "τιμές ακεραιότητας παρατίθενται στο Παράρτημα Γ.",
            )
        elif paragraph.text.startswith(
            "Οι άμεσες συγκρίσεις μεταξύ μεθόδων αναφέρονται ως root-paired διαφορές A−B μετά την equal-layout reduction."
        ):
            _set_text(
                paragraph,
                "Οι άμεσες συγκρίσεις μεταξύ μεθόδων αναφέρονται ως ζευγαρωμένες "
                "διαφορές A−B ανά ανεξάρτητη επανάληψη, αφού οι δύο διατάξεις "
                "συνδυαστούν με ίσο βάρος. Τα διαστήματα είναι σημειακά διαστήματα "
                "εκτίμησης και δεν αποτελούν ταυτόχρονη συμπερασματολογία για ολόκληρη "
                "οικογένεια συγκρίσεων. Δεν έχει οριστεί οικογένεια p-values ούτε "
                "σύνθετος δείκτης συνολικής κατάταξης· επομένως η παρουσίαση παραμένει "
                "περιγραφική ως προς τις εκτιμώμενες διαφορές και την αβεβαιότητά τους.",
            )

    # Result-table labels/captions remain real Word SEQ captions.
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text == "Condition":
            table.rows[0].cells[0].text = "Συνθήκη"
        if table.rows and len(table.columns) == 3 and table.rows[0].cells[0].text == "Μέθοδος":
            table.rows[0].cells[1].text = "Κυκλική"
            table.rows[0].cells[2].text = "Εναλλαγή"
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Caption" or len(paragraph.runs) < 3:
            continue
        if "5.3.1 Συνολική εικόνα ανά condition" in paragraph.text:
            paragraph.runs[2].text = " — Μέσες τιμές οφέλους προσαρμογής και σημειακά διαστήματα 95% ανά συνθήκη."
        elif "5.4.3 Ευαισθησία στην ανοχή ανάκαμψης" in paragraph.text:
            paragraph.runs[2].text = " — Αναλογία σταθερής ανάκαμψης στις ανοχές 0,05/0,10/0,20 για τις δύο συνθήκες ανααντιστοίχισης."

    # The acronym table is front-matter support and should not consume the numbered table sequence.
    glossary_caption_removed = False
    for paragraph in list(doc.paragraphs):
        if paragraph.style.name == "Caption" and "Σύνοψη για την ενότητα «Ακρωνύμια»" in paragraph.text:
            _remove_paragraph(paragraph)
            glossary_caption_removed = True
            break

    try:
        glossary_style = doc.styles["Glossary Term"]
    except KeyError:
        glossary_style = doc.styles.add_style("Glossary Term", WD_STYLE_TYPE.PARAGRAPH)
        glossary_style.base_style = doc.styles["Normal"]
        glossary_style.font.name = "Times New Roman"
        glossary_style.font.size = Pt(11)
        glossary_style.font.bold = True
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 3" and paragraph.text in GLOSSARY_TERMS:
            paragraph.style = glossary_style

    # Remove the review-build production note appendix; implementation screenshots remain optional T-538 work.
    old_d = _find(doc, "Παράρτημα Δ — Ερευνητική εφαρμογή και όριο παρουσίασης")
    old_e = _find(doc, "Παράρτημα Ε — Αναπαραγωγή και λογισμικό περιβάλλον")
    body = doc.element.body
    children = list(body)
    start = children.index(old_d._p)
    end = children.index(old_e._p)
    for element in children[start:end]:
        body.remove(element)
    old_e.text = "Παράρτημα Δ — Αναπαραγωγή και λογισμικό περιβάλλον"

    for paragraph in doc.paragraphs:
        if paragraph.text == "Το αναπαραγώγιμο λογισμικό περιβάλλον συνοψίζεται ως εξής:":
            _set_text(
                paragraph,
                "Η αναπαραγωγή της μελέτης βασίζεται στο ίδιο ελεγχόμενο λογισμικό και "
                "υπολογιστικό περιβάλλον που χρησιμοποιήθηκε στην τελική εκτέλεση. Η "
                "επιστημονική λογική παραμένει ανεξάρτητη από το γραφικό περιβάλλον και "
                "η ακεραιότητα των αποτελεσμάτων ελέγχεται μέσω των αποθηκευμένων πακέτων "
                "εκτέλεσης και των manifests.",
            )
        elif paragraph.text == "Δεν απαιτείται εκτενές code listing. Μικρά code/config snippets επιτρέπονται μόνο αν εξηγούν contract που δεν αποδίδεται καθαρότερα με ψευδοκώδικα, πίνακα ή διάγραμμα.":
            _set_text(
                paragraph,
                "Η πλήρης υλοποίηση παραμένει διαθέσιμη στο repository της εργασίας. "
                "Στο παρόν παράρτημα καταγράφονται μόνο τα στοιχεία που είναι αναγκαία "
                "για την αναπαραγωγή και τον έλεγχο της επιστημονικής διαδικασίας, χωρίς "
                "εκτενή αποσπάσματα κώδικα.",
            )
    bullet_replacements = {
        "Python 3.12 + locked uv environment,": "Python 3.12 σε κλειδωμένο περιβάλλον uv,",
        "project-owned package src/resilient_agents/,": "επιστημονικός πυρήνας στο project-owned package src/resilient_agents/,",
        "deterministic separated RNG streams,": "διαχωρισμένες και ντετερμινιστικές ροές τυχαιότητας (RNG),",
        "Study recipe/plan/store lifecycle,": "κύκλος ζωής Study recipe/plan/store με ακριβή ιχνηλασιμότητα,",
        "filesystem evidence as authority,": "τα αποθηκευμένα πακέτα τεκμηρίων στο filesystem ως πηγή αλήθειας,",
        "manifest/checksum validation,": "επικύρωση manifest και checksums,",
        "read-only thesis bibliography consumer.": "read-only κατανάλωση της συγχρονισμένης βιβλιογραφίας της διπλωματικής.",
    }
    for paragraph in doc.paragraphs:
        if paragraph.text in bullet_replacements:
            _set_text(paragraph, bullet_replacements[paragraph.text])

    _insert_config_table(doc)
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Οι δύο held-out final 7×7 layouts"):
            _set_text(
                paragraph,
                "Η τελική μελέτη χρησιμοποιεί δύο ανεξάρτητες διατάξεις GridWorld 7×7. "
                "Και στις δύο, η αρχική θέση είναι (0,0), ο στόχος (6,6), το μήκος της "
                "συντομότερης διαδρομής είναι 12 βήματα και το διοικητικό όριο επεισοδίου "
                "48 βήματα. Η στατιστική ανάλυση χρησιμοποιεί 12 ανεξάρτητες επαναλήψεις "
                "(roots)· οι δύο διατάξεις συνδυάζονται με ίσο βάρος μέσα σε κάθε root "
                "πριν από τη συμπερασματολογία.",
            )
        elif paragraph.text == "Οι τέσσερις frozen Phase-B conditions είναι:":
            _set_text(paragraph, "Οι τέσσερις προκαθορισμένες συνθήκες της Φάσης B είναι:")
        elif paragraph.text == "Οι ακριβείς mappings/probabilities/support rules πρέπει να προέρχονται από το frozen protocol/configuration.":
            _set_text(
                paragraph,
                "Στην εναλλαγή right–down, οι right και down ανταλλάσσουν συνέπειες ενώ "
                "up/left παραμένουν αμετάβλητες. Στην κυκλική ανααντιστοίχιση ισχύει "
                "up→right→down→left→up. Η αποτυχία ενέργειας εφαρμόζεται με πιθανότητα "
                "0,15. Η αλλοίωση παρατήρησης εφαρμόζεται με πιθανότητα 0,05 και "
                "δειγματοληψία ομοιόμορφα από έγκυρες μη εμποδισμένες θέσεις, εξαιρώντας "
                "την πραγματική κατάσταση.",
            )

    appendix_b = _find(doc, "Παράρτημα Β — Πλήρη αποτελέσματα και διαγνωστικά ανά ερευνητικό ερώτημα")
    a4 = doc.add_paragraph(style="Heading 3")
    a4.text = "Α.4 Προϋπολογισμοί και συμβόλαιο ανάκαμψης"
    a4_body = doc.add_paragraph(style="Normal")
    _set_text(
        a4_body,
        "Η Φάση A διαθέτει 8.192 πραγματικές αλληλεπιδράσεις εκπαίδευσης ανά μονάδα, "
        "με probes στα 0, 512, 1.024, 2.048, 4.096 και 8.192 interactions και 12 "
        "επεισόδια ανά probe. Η Φάση B έχει ορίζοντα 256 interactions. Για το RQ3 "
        "χρησιμοποιούνται παθητικά παράθυρα 32 interactions, κύρια ανοχή 0,10, "
        "προκαθορισμένες ανοχές ευαισθησίας 0,05 και 0,20 και απαίτηση δύο διαδοχικών "
        "παραθύρων εντός ανοχής για σταθερή ανάκαμψη. Η μη ανάκαμψη παραμένει δεξιά "
        "λογοκριμένη στο 256 με null recovery time.",
    )
    for element in (a4._p, a4_body._p):
        body.remove(element)
    pos = list(body).index(appendix_b._p)
    body.insert(pos, a4._p)
    body.insert(pos + 1, a4_body._p)

    # Official structure: glossary before main text, bibliography before appendices.
    _move_block(doc, "Γλωσσάριο και Ακρωνύμια", "Παραρτήματα", "Κεφάλαιο 1 — Εισαγωγή")
    bibliography = _find(doc, "Βιβλιογραφία")
    appendices = _find(doc, "Παραρτήματα")
    children = list(body)
    start = children.index(bibliography._p)
    block = []
    for element in children[start:]:
        if element.tag == qn("w:sectPr"):
            break
        block.append(element)
    for element in block:
        body.remove(element)
    pos = list(body).index(appendices._p)
    for offset, element in enumerate(block):
        body.insert(pos + offset, element)
    for title in ("Γλωσσάριο και Ακρωνύμια", "Κεφάλαιο 1 — Εισαγωγή", "Βιβλιογραφία", "Παραρτήματα"):
        _find(doc, title).paragraph_format.page_break_before = True

    bibliography_normalized = 0
    for paragraph in doc.paragraphs:
        match = re.match(r"^\[(\d+)\]\s", paragraph.text)
        if not match:
            continue
        number = int(match.group(1))
        if number in BIBLIOGRAPHY:
            _set_text(paragraph, BIBLIOGRAPHY[number])
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.first_line_indent = Cm(-0.6)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            bibliography_normalized += 1

    # Preserve automatic fields while giving headless review builds meaningful cached results.
    paragraphs = doc.paragraphs
    toc_start = next(i for i, p in enumerate(paragraphs) if p.text == "Γλωσσάριο και Ακρωνύμια")
    toc_lines: list[str] = []
    for paragraph in paragraphs[toc_start:]:
        if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            level = int(paragraph.style.name[-1])
            toc_lines.append("    " * (level - 1) + paragraph.text)
    _field_result_lines(Paragraph(_find(doc, "Πίνακας Περιεχομένων")._p.getnext(), doc), toc_lines)
    figure_lines = [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Caption" and paragraph.text.startswith("Σχήμα ")
    ]
    _field_result_lines(Paragraph(_find(doc, "Κατάλογος Σχημάτων")._p.getnext(), doc), figure_lines)
    table_lines = [
        "Πίνακας 1 — Μέσες τιμές οφέλους προσαρμογής και σημειακά διαστήματα 95% ανά συνθήκη.",
        "Πίνακας 2 — Αναλογία σταθερής ανάκαμψης στις ανοχές 0,05/0,10/0,20 για τις δύο συνθήκες ανααντιστοίχισης.",
        "Πίνακας 3 — Τελικές ρυθμίσεις των πέντε μεθόδων στο protocol-v2.1.",
    ]
    _field_result_lines(Paragraph(_find(doc, "Κατάλογος Πινάκων")._p.getnext(), doc), table_lines)

    doc.save(output)

    final_doc = Document(output)
    media_after = v13.v12._media_hashes(output)
    final_sha = v13.v12._sha256(output)
    final_paragraph_count = len(final_doc.paragraphs)
    final_table_count = len(final_doc.tables)
    placeholder_hits = v13.v12._placeholder_hits(final_doc)
    forbidden_wording = [
        paragraph.text
        for paragraph in final_doc.paragraphs
        if "σημαντικά υψηλότερη" in paragraph.text
        or "είναι η ισχυρότερη ως προς" in paragraph.text
        or "είχε το ισχυρότερο αποτέλεσμα" in paragraph.text
    ]
    h1_order = [p.text for p in final_doc.paragraphs if p.style.name == "Heading 1"]
    required_order = [
        "Γλωσσάριο και Ακρωνύμια",
        "Κεφάλαιο 1 — Εισαγωγή",
        "Κεφάλαιο 7 — Συμπεράσματα και Μελλοντική Εργασία",
        "Βιβλιογραφία",
        "Παραρτήματα",
    ]
    order_indices = [h1_order.index(value) for value in required_order]
    numbered_table_captions = [
        p.text for p in final_doc.paragraphs
        if p.style.name == "Caption" and p.text.startswith("Πίνακας ")
    ]

    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": final_paragraph_count,
            "post_synthesis_paragraph_count": final_paragraph_count,
            "table_count": final_table_count,
            "rendered_table_count": final_table_count,
            "cached_table_entry_count": 3,
            "t711a_hardening_version": 14,
            "academic_heading_replacement_count": heading_replacements,
            "neutral_wording_replacement_count": wording_replacements,
            "forbidden_superiority_wording": forbidden_wording,
            "glossary_table_unnumbered": glossary_caption_removed,
            "numbered_table_caption_count": len(numbered_table_captions),
            "appendix_config_table_present": final_table_count == 4,
            "removed_review_application_appendix": "Παράρτημα Δ — Ερευνητική εφαρμογή και όριο παρουσίασης" not in [p.text for p in final_doc.paragraphs],
            "bibliography_normalized_count": bibliography_normalized,
            "document_major_order": h1_order,
            "academic_order_valid": order_indices == sorted(order_indices),
            "embedded_media_bytes_preserved": media_before == media_after,
            "scientific_values_modified": False,
            "placeholder_hits": placeholder_hits,
            "placeholder_count": len(placeholder_hits),
            "final_visual_qa_required": True,
        }
    )

    hardening_ok = (
        report.get("status") == "pass"
        and report.get("citation_count") == report.get("verified_reference_identity_count") == 17
        and report.get("inserted_asset_count") == report.get("planned_figure_count") == 24
        and report.get("inline_shape_count_after_synthesis") == 25
        and media_before == media_after
        and heading_replacements >= 35
        and wording_replacements == 3
        and not forbidden_wording
        and glossary_caption_removed
        and final_table_count == 4
        and len(numbered_table_captions) == 3
        and bibliography_normalized == 17
        and order_indices == sorted(order_indices)
        and report.get("placeholder_count") == 3
        and report.get("registered_asset_bytes_modified") is False
    )
    if not hardening_ok:
        report["status"] = "fail"

    qa_output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if report["status"] != "pass":
        raise RuntimeError(f"T-711A v14 hardening failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
