#!/usr/bin/env python3
"""Build the T-711 review-ready Word thesis from the merged T-710 manuscript.

This is a composition tool, not a scientific analysis tool. It consumes only:
- docs/thesis/draft/ accepted manuscript files;
- research/bibliography/citation-ready/ canonical metadata;
- results/thesis-assets/protocol-v2.1-final/ registered T-613 assets.

It does not compute or reinterpret any scientific estimand.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DRAFT = ROOT / "docs" / "thesis" / "draft"
BIB = ROOT / "research" / "bibliography" / "citation-ready"
ASSET_ROOT = ROOT / "results" / "thesis-assets" / "protocol-v2.1-final"

MANUSCRIPT_FILES = [
    DRAFT / "CHAPTER_01_INTRODUCTION.md",
    DRAFT / "CHAPTER_02_BACKGROUND_RELATED_WORK.md",
    DRAFT / "CHAPTER_03_METHODOLOGY.md",
    DRAFT / "CHAPTER_04_SYSTEM_ARCHITECTURE.md",
    DRAFT / "CHAPTER_05_RESULTS.md",
    DRAFT / "CHAPTER_06_DISCUSSION.md",
    DRAFT / "CHAPTER_07_CONCLUSIONS.md",
]
GLOSSARY_FILE = DRAFT / "GLOSSARY_ACRONYMS.md"
APPENDIX_FILE = DRAFT / "APPENDIX_DRAFT.md"
FRONT_FILE = DRAFT / "FRONT_MATTER_SUMMARIES.md"

SRC_RE = re.compile(r"SRC-[0-9A-F]{10}")
CITATION_GROUP_RE = re.compile(r"\[(?:@SRC-[0-9A-F]{10})(?:;\s*@SRC-[0-9A-F]{10})*\]")
ASSET_RE = re.compile(r"FIG-[A-Z0-9-]+")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def set_run_font(run, name: str = "Times New Roman", size: float | None = None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    rfonts = run._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str, size: float, bold: bool | None = None):
    style.font.name = name
    style._element.get_or_add_rPr()
    rfonts = style._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        style._element.rPr.insert(0, rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(key), name)
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str, display: str = ""):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = display
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, txt, fld_end])
    set_run_font(run)
    return run


def set_update_fields(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, " PAGE ", "1")


def strip_markdown_inline(s: str) -> str:
    s = s.replace("`", "")
    s = s.replace("**", "")
    s = s.replace("*", "")
    return s


def load_bibliography() -> dict[str, dict[str, str]]:
    path = BIB / "catalog" / "sources.csv"
    with path.open(encoding="utf-8", newline="") as f:
        rows = list(csv.DictReader(f))
    by_id = {r["Κωδικός"]: r for r in rows}
    return by_id


def manuscript_order_text() -> str:
    parts = [text(FRONT_FILE)]
    parts.extend(text(p) for p in MANUSCRIPT_FILES)
    parts.extend([text(GLOSSARY_FILE), text(APPENDIX_FILE)])
    return "\n".join(parts)


def build_citation_map(bib: dict[str, dict[str, str]]) -> dict[str, int]:
    mapping: dict[str, int] = {}
    for match in SRC_RE.finditer(manuscript_order_text()):
        sid = match.group(0)
        if sid not in bib:
            raise ValueError(f"formal citation {sid} is absent from citation-ready catalog")
        if sid not in mapping:
            mapping[sid] = len(mapping) + 1
    return mapping


def replace_citations(s: str, cmap: dict[str, int]) -> str:
    def repl(m: re.Match) -> str:
        ids = SRC_RE.findall(m.group(0))
        return ", ".join(f"[{cmap[x]}]" for x in ids)
    return CITATION_GROUP_RE.sub(repl, s)


def add_inline(paragraph, value: str, cmap: dict[str, int]):
    value = replace_citations(value, cmap)
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    pos = 0
    for m in token_re.finditer(value):
        if m.start() > pos:
            r = paragraph.add_run(value[pos:m.start()])
            set_run_font(r, size=11)
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_run_font(r, size=11, bold=True)
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, name="Consolas", size=9.5)
        else:
            r = paragraph.add_run(token[1:-1])
            set_run_font(r, size=11, italic=True)
        pos = m.end()
    if pos < len(value):
        r = paragraph.add_run(value[pos:])
        set_run_font(r, size=11)


def add_body_paragraph(doc: Document, value: str, cmap: dict[str, int]):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(p, value, cmap)
    return p


def add_list_paragraph(doc: Document, value: str, level: int, ordered: bool, cmap: dict[str, int]):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.65 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    add_inline(p, value, cmap)
    return p


def add_equation(doc: Document, latex: str):
    replacements = {
        r"\le": "≤", r"\ge": "≥", r"\gamma": "γ", r"\alpha": "α", r"\pi": "π",
        r"\mu": "μ", r"\sigma": "σ", r"\Delta": "Δ", r"\times": "×", r"\pm": "±",
        r"\left": "", r"\right": "", r"\;": " ", r"\,": " ",
    }
    clean = latex.strip()
    for a, b in replacements.items():
        clean = clean.replace(a, b)
    clean = re.sub(r"\\text\{([^{}]+)\}", r"\1", clean)
    clean = clean.replace("\\", "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(clean)
    set_run_font(r, name="Cambria Math", size=11)
    return p


def load_assets() -> dict[str, dict]:
    manifest = json.loads(text(ASSET_ROOT / "asset-manifest.json"))
    return {a["asset_id"]: a for a in manifest["assets"]}


def choose_output(asset: dict, fmt: str) -> Path | None:
    for out in asset.get("outputs", []):
        if out.get("format") == fmt:
            path = ASSET_ROOT / out["relative_path"]
            if path.exists():
                if sha256(path) != out["sha256"]:
                    raise ValueError(f"asset hash mismatch: {asset['asset_id']} {path}")
                return path
    return None


def add_seq_caption(doc: Document, label: str, caption: str, prefix: str = ""):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if prefix:
        r = p.add_run(prefix)
        set_run_font(r, size=10, bold=True)
    r = p.add_run(f"{label} ")
    set_run_font(r, size=10, bold=True)
    add_field(p, f" SEQ {label} \\* ARABIC ", "1")
    r = p.add_run(f" — {caption}")
    set_run_font(r, size=10)
    return p


def add_figure(doc: Document, asset: dict, inserted: list[str]):
    png = choose_output(asset, "png")
    if png is None:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(png), width=Inches(6.1))
    add_seq_caption(doc, "Σχήμα", asset["caption"])
    inserted.append(asset["asset_id"])


def add_asset_table(doc: Document, asset: dict, inserted: list[str]):
    csv_path = choose_output(asset, "csv")
    if csv_path is None:
        return
    with csv_path.open(encoding="utf-8", newline="") as f:
        rows = list(csv.reader(f))
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i in range(cols):
        hdr[i].text = rows[0][i] if i < len(rows[0]) else ""
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(hdr[i])
        for run in hdr[i].paragraphs[0].runs:
            set_run_font(run, size=8.5, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(cols):
            cells[i].text = row[i] if i < len(row) else ""
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for run in cells[i].paragraphs[0].runs:
                set_run_font(run, size=8)
    add_seq_caption(doc, "Πίνακας", asset["caption"])
    inserted.append(asset["asset_id"])


def add_markdown_table(doc: Document, lines: list[str], section_title: str, cmap: dict[str, int], table_counter: list[int]):
    parsed = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in parsed[1]):
        parsed.pop(1)
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(cols):
        table.rows[0].cells[i].text = replace_citations(strip_markdown_inline(parsed[0][i] if i < len(parsed[0]) else ""), cmap)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_run_font(run, size=9, bold=True)
        set_cell_margins(table.rows[0].cells[i])
    for row in parsed[1:]:
        cells = table.add_row().cells
        for i in range(cols):
            cells[i].text = replace_citations(strip_markdown_inline(row[i] if i < len(row) else ""), cmap)
            for run in cells[i].paragraphs[0].runs:
                set_run_font(run, size=8.5)
            set_cell_margins(cells[i])
    table_counter[0] += 1
    add_seq_caption(doc, "Πίνακας", f"Σύνοψη για την ενότητα «{section_title}».")


def insert_asset_ids(doc: Document, ids: Iterable[str], assets: dict[str, dict], inserted: list[str], inserted_rq_tables: set[str]):
    for aid in ids:
        asset = assets.get(aid)
        if asset is None:
            continue
        if aid in inserted:
            continue
        if asset.get("kind") == "figure":
            add_figure(doc, asset, inserted)
            rq = asset.get("rq_scope")
            if rq and rq not in inserted_rq_tables:
                for other in assets.values():
                    if other.get("kind") == "table" and other.get("rq_scope") == rq and "main-thesis" in other.get("intended_use", []):
                        if other["asset_id"] not in inserted:
                            add_asset_table(doc, other, inserted)
                inserted_rq_tables.add(rq)


def render_markdown(doc: Document, md: str, cmap: dict[str, int], assets: dict[str, dict], inserted: list[str], inserted_rq_tables: set[str], start_new_page: bool = True):
    lines = md.splitlines()
    i = 0
    section_title = ""
    table_counter = [0]
    if start_new_page:
        doc.add_page_break()
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("> **Σημείωση drafting:**") or stripped.startswith("> Το T-710"):
            i += 1
            continue
        if "Προτεινόμενη T-711 τοποθέτηση" in stripped:
            insert_asset_ids(doc, ASSET_RE.findall(stripped), assets, inserted, inserted_rq_tables)
            i += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if level == 1:
                style = "Heading 1"
            elif level == 2:
                style = "Heading 2"
            else:
                style = "Heading 3"
            p = doc.add_paragraph(style=style)
            add_inline(p, title, cmap)
            section_title = strip_markdown_inline(replace_citations(title, cmap))
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_markdown_table(doc, block, section_title or "τρέχουσα ενότητα", cmap, table_counter)
            continue
        if stripped == r"\[":
            eq = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                eq.append(lines[i])
                i += 1
            add_equation(doc, " ".join(eq))
            i += 1
            continue
        m = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if m:
            level = 1 + len(m.group(1)) // 2
            add_list_paragraph(doc, m.group(2), level, False, cmap)
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if m:
            level = 1 + len(m.group(1)) // 2
            add_list_paragraph(doc, m.group(2), level, True, cmap)
            i += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip("> "), cmap)
            i += 1
            continue
        # Collect adjacent prose lines as one paragraph, stopping before Markdown structures.
        para = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith(("#", "|", ">", "- ", "* ", r"\[")) or re.match(r"\d+\.\s+", nxt):
                break
            para.append(nxt)
            i += 1
        add_body_paragraph(doc, " ".join(para), cmap)


def extract_front_sections() -> tuple[str, str, str, str]:
    md = text(FRONT_FILE)
    def section(name: str, next_name: str | None) -> str:
        start = md.index(f"## {name}") + len(f"## {name}")
        end = md.index(f"## {next_name}", start) if next_name else len(md)
        return md[start:end].strip()
    greek = section("Περίληψη", "Abstract")
    eng = section("Abstract", None)
    greek_kw = re.search(r"\*\*Λέξεις-κλειδιά:\*\*\s*(.+)", greek)
    eng_kw = re.search(r"\*\*Keywords:\*\*\s*(.+)", eng)
    greek = re.sub(r"\*\*Λέξεις-κλειδιά:\*\*.*", "", greek).strip()
    eng = re.sub(r"\*\*Keywords:\*\*.*", "", eng).strip()
    return greek, greek_kw.group(1) if greek_kw else "", eng, eng_kw.group(1) if eng_kw else ""


def add_title_page(doc: Document, title: str, subtitle: str, lang_note: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    r = p.add_run(title)
    set_run_font(r, size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=14, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(4)
    r = p.add_run("Ονοματεπώνυμο φοιτητή: [να συμπληρωθεί από το επίσημο έντυπο]")
    set_run_font(r, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(lang_note)
    set_run_font(r, size=10, italic=True)
    doc.add_page_break()


def add_front_matter(doc: Document, cmap: dict[str, int]):
    greek_title = "Σύγκριση και Αξιολόγηση Ανθεκτικών Πρακτόρων Τεχνητής Νοημοσύνης σε Περιβάλλοντα με Αβεβαιότητα"
    english_title = "Comparison and Evaluation of Resilient AI Agents in Uncertain Environments"
    add_title_page(doc, greek_title, "Διπλωματική Εργασία", "Ελληνική σελίδα τίτλου — στοιχεία φοιτητή/ιδρύματος συμπληρώνονται μόνο από authoritative metadata.")
    add_title_page(doc, english_title, "Diploma Thesis", "English title page — student/institution metadata remains unresolved until authoritative input is supplied.")

    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Δηλώσεις πνευματικών δικαιωμάτων και λογοκλοπής")
    p = doc.add_paragraph()
    r = p.add_run("[ΕΚΚΡΕΜΕΙ: να εισαχθεί η ακριβής επίσημη διατύπωση της δήλωσης από το ισχύον έντυπο/οδηγία. Δεν επινοείται κείμενο στο T-711.]")
    set_run_font(r, size=11, italic=True)
    doc.add_page_break()

    greek, greek_kw, eng, eng_kw = extract_front_sections()
    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Περίληψη")
    render_markdown(doc, greek, cmap, {}, [], set(), start_new_page=False)
    p = doc.add_paragraph()
    r = p.add_run("Λέξεις-κλειδιά: ")
    set_run_font(r, size=11, bold=True)
    r = p.add_run(greek_kw)
    set_run_font(r, size=11)
    doc.add_page_break()

    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Abstract")
    render_markdown(doc, eng, cmap, {}, [], set(), start_new_page=False)
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    set_run_font(r, size=11, bold=True)
    r = p.add_run(eng_kw)
    set_run_font(r, size=11)
    doc.add_page_break()

    for title, instruction, placeholder in (
        ("Πίνακας Περιεχομένων", ' TOC \\o "1-3" \\h \\z \\u ', "Ενημερώστε το πεδίο TOC στο Microsoft Word."),
        ("Κατάλογος Σχημάτων", ' TOC \\h \\z \\c "Σχήμα" ', "Ενημερώστε το πεδίο καταλόγου σχημάτων στο Microsoft Word."),
        ("Κατάλογος Πινάκων", ' TOC \\h \\z \\c "Πίνακας" ', "Ενημερώστε το πεδίο καταλόγου πινάκων στο Microsoft Word."),
    ):
        h = doc.add_paragraph(style="Heading 1")
        h.add_run(title)
        p = doc.add_paragraph()
        add_field(p, instruction, placeholder)
        doc.add_page_break()


def format_authors(authors: str) -> str:
    parts = [p.strip() for p in authors.split(";") if p.strip()]
    if len(parts) <= 1:
        return parts[0] if parts else ""
    if len(parts) == 2:
        return f"{parts[0]} and {parts[1]}"
    return ", ".join(parts[:-1]) + f", and {parts[-1]}"


def add_references(doc: Document, cmap: dict[str, int], bib: dict[str, dict[str, str]]):
    doc.add_page_break()
    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Βιβλιογραφία")
    for sid, n in sorted(cmap.items(), key=lambda x: x[1]):
        row = bib[sid]
        authors = format_authors(row.get("Συγγραφείς", ""))
        title = row.get("Τίτλος", "").strip()
        year = row.get("Έτος", "").strip()
        url = row.get("Σύνδεσμος", "").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.space_after = Pt(5)
        content = f"[{n}] {authors}, “{title},” {year}."
        if url:
            content += f" [Online]. Available: {url}"
        r = p.add_run(content)
        set_run_font(r, size=10)


def configure_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)

    styles = doc.styles
    set_style_font(styles["Normal"], "Times New Roman", 11)
    set_style_font(styles["Heading 1"], "Times New Roman", 14, True)
    set_style_font(styles["Heading 2"], "Times New Roman", 12, True)
    set_style_font(styles["Heading 3"], "Times New Roman", 11, True)
    set_style_font(styles["Caption"], "Times New Roman", 10)
    set_style_font(styles["List Bullet"], "Times New Roman", 11)
    set_style_font(styles["List Number"], "Times New Roman", 11)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = True
    styles["Normal"].paragraph_format.line_spacing = 1.5

    footer = sec.footer
    add_page_number(footer.paragraphs[0])
    set_update_fields(doc)


def build(output: Path, qa_output: Path):
    required = MANUSCRIPT_FILES + [FRONT_FILE, GLOSSARY_FILE, APPENDIX_FILE, BIB / "catalog" / "sources.csv", ASSET_ROOT / "asset-manifest.json"]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError("missing required T-711 inputs: " + ", ".join(missing))

    bib = load_bibliography()
    cmap = build_citation_map(bib)
    assets = load_assets()

    doc = Document()
    configure_document(doc)
    add_front_matter(doc, cmap)

    inserted: list[str] = []
    inserted_rq_tables: set[str] = set()
    for path in MANUSCRIPT_FILES:
        render_markdown(doc, text(path), cmap, assets, inserted, inserted_rq_tables, start_new_page=True)

    render_markdown(doc, text(GLOSSARY_FILE), cmap, assets, inserted, inserted_rq_tables, start_new_page=True)
    render_markdown(doc, text(APPENDIX_FILE), cmap, assets, inserted, inserted_rq_tables, start_new_page=True)
    add_references(doc, cmap, bib)

    # Traceable review note, deliberately not claiming final freeze.
    doc.add_page_break()
    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Σημείωση έκδοσης review-ready")
    notes = [
        "Η έκδοση αυτή παράχθηκε μηχανικά από το merged T-710 manuscript και δεν αποτελεί ακόμη T-713 final freeze.",
        "Εκκρεμούν authoritative προσωπικά/διοικητικά metadata και η ακριβής επίσημη διατύπωση δηλώσεων όπου δεν έχουν παρασχεθεί.",
        "Τα Word fields (TOC, κατάλογοι, caption numbering, page references) έχουν οριστεί ώστε να ενημερωθούν στο Microsoft Word πριν από τελική υποβολή.",
        "Κάθε quantitative figure/table προέρχεται από registered T-613 asset και επαληθεύεται με το manifest SHA-256 πριν από εισαγωγή.",
    ]
    for note in notes:
        add_list_paragraph(doc, note, 1, False, cmap)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    # Structural post-save checks.
    from docx import Document as ReloadedDocument
    check = ReloadedDocument(output)
    full_text = "\n".join(p.text for p in check.paragraphs)
    unresolved_src = sorted(set(SRC_RE.findall(full_text)))
    unresolved_drafting = "Σημείωση drafting" in full_text or "Προτεινόμενη T-711 τοποθέτηση" in full_text
    qa = {
        "schema_version": 1,
        "status": "pass" if not unresolved_src and not unresolved_drafting else "fail",
        "output": str(output.relative_to(ROOT) if output.is_relative_to(ROOT) else output),
        "output_sha256": sha256(output),
        "bibliography_source_commit": text(BIB / "SOURCE_COMMIT").strip(),
        "citation_count": len(cmap),
        "citation_map": {sid: n for sid, n in sorted(cmap.items(), key=lambda x: x[1])},
        "unresolved_src_ids": unresolved_src,
        "contains_drafting_notes": unresolved_drafting,
        "inserted_asset_ids": inserted,
        "inserted_asset_count": len(inserted),
        "paragraph_count": len(check.paragraphs),
        "table_count": len(check.tables),
        "section_count": len(check.sections),
        "input_sha256": {str(p.relative_to(ROOT)): sha256(p) for p in required},
        "scientific_boundary": "composition-only; no estimand/statistical recomputation",
    }
    qa_output.parent.mkdir(parents=True, exist_ok=True)
    qa_output.write_text(json.dumps(qa, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if qa["status"] != "pass":
        raise RuntimeError(f"T-711 structural QA failed: {qa}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=ROOT / "artifacts" / "t711" / "resilient-ai-agents-thesis-review-ready.docx")
    parser.add_argument("--qa-output", type=Path, default=ROOT / "artifacts" / "t711" / "qa-report.json")
    args = parser.parse_args()
    build(args.output.resolve(), args.qa_output.resolve())
    print(args.output)
    print(args.qa_output)


if __name__ == "__main__":
    main()
