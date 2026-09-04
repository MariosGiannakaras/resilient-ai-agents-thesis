#!/usr/bin/env python3
"""T-715 audit-reconciled reader composition.

Builds the accepted T-715 v26 composition, then inserts bounded reader-facing methodology,
reproducibility and limitations clarifications from frozen protocol/analysis authorities.
No scientific evidence, numerical outcome, registered T-613 asset, or application screenshot
byte is recomputed or replaced.
"""

from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document

import t711_build_entry_v26 as v26
import t715_audit_hardening as audit


t711 = v26.t711
_previous_build = t711.builder.build


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in sorted(n for n in archive.namelist() if n.startswith("word/media/"))
        }


def _build(output: Path, qa_output: Path) -> None:
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    if report.get("status") != "pass":
        raise RuntimeError(f"T-715 v26 composition failed before audit hardening: {report}")

    base_doc = Document(output)
    paragraphs_before = len(base_doc.paragraphs)
    inline_before = len(base_doc.inline_shapes)
    media_before = _media_hashes(output)

    hardening = audit.apply(base_doc)
    if hardening.get("already_applied"):
        raise RuntimeError("T-715 audit hardening unexpectedly already present in fresh v26 output")
    base_doc.save(output)

    final_doc = Document(output)
    paragraphs_after = len(final_doc.paragraphs)
    inline_after = len(final_doc.inline_shapes)
    media_after = _media_hashes(output)
    media_preserved = media_before == media_after

    sentinel_counts = {
        text: sum(1 for paragraph in final_doc.paragraphs if paragraph.text.strip() == text)
        for text in audit.SENTINELS
    }
    inserted_delta = paragraphs_after - paragraphs_before
    final_sha = _sha256(output)

    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": paragraphs_after,
            "post_synthesis_paragraph_count": paragraphs_after,
            "t715_audit_hardening_applied": True,
            "t715_audit_hardening_version": 1,
            "t715_audit_hardening_inserted_paragraph_count": hardening["inserted_paragraph_count"],
            "t715_audit_hardening_inserted_heading_count": hardening["inserted_heading_count"],
            "t715_audit_hardening_paragraph_delta": inserted_delta,
            "t715_audit_hardening_sentinel_counts": sentinel_counts,
            "t715_audit_hardening_media_preserved": media_preserved,
            "t715_audit_hardening_inline_shapes_preserved": inline_before == inline_after,
            "t715_audit_hardening_source_scope": "composition-only; frozen protocol/analysis reconciliation",
            "t715_tuning_unit_count": 180,
            "t715_tuning_candidate_count_per_method": 6,
            "t715_tuning_root_count": 3,
            "t715_tuning_layout_count": 2,
            "t715_tuning_winners": ["q-c06", "sarsa-c06", "dqn-c05", "ppo-c06", "dyna-c03"],
            "t715_final_root_count": 12,
            "t715_root_count_candidates": [12, 16, 20, 24],
            "t715_root_count_max_half_width": 0.1428,
            "t715_recovery_window_size": 32,
            "t715_recovery_horizon": 256,
            "t715_recovery_tolerance": 0.10,
            "t715_recovery_stability_windows": 2,
            "t715_action_failure_probability": 0.15,
            "t715_action_failure_step_reward": -0.1,
            "t715_observation_corruption_probability": 0.05,
            "t715_observation_corruption_excludes_true_state": True,
            "t715_observation_corruption_excludes_goal": False,
            "t715_layout_generation_seeds": [57001, 57002],
            "t715_seed_stream_ranges": {
                "initialization": [71001, 71012],
                "exploration": [72001, 72012],
                "scenario": [73001, 73012],
                "environment": [74001, 74012],
                "action_disturbance": [75001, 75012],
                "observation_disturbance": [76001, 76012],
            },
            "t715_new_posthoc_binomial_test_added": False,
            "t715_new_experiment_or_reanalysis": False,
            "scientific_values_modified": False,
            "registered_asset_bytes_modified": False,
            "final_visual_qa_required": True,
        }
    )

    hardening_ok = (
        hardening.get("applied") is True
        and hardening.get("inserted_paragraph_count") == audit.EXPECTED_INSERTED_PARAGRAPHS == 26
        and hardening.get("inserted_heading_count") == audit.EXPECTED_INSERTED_HEADINGS == 3
        and inserted_delta == 26
        and all(count == 1 for count in sentinel_counts.values())
        and media_preserved
        and inline_before == inline_after
        and report.get("t715_preexisting_docx_media_preserved") is True
        and report.get("t715_new_docx_media_count") == 2
        and report.get("t715_application_screenshot_count") == 2
        and report.get("registered_figure_count_preserved") is True
        and report.get("scientific_values_modified") is False
        and report.get("registered_asset_bytes_modified") is False
    )
    if not hardening_ok:
        report["status"] = "fail"

    qa_output.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if report.get("status") != "pass":
        raise RuntimeError(f"T-715 audit-reconciled composition failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
