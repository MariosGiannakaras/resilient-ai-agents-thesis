#!/usr/bin/env python3
"""T-711A bounded academic/compliance hardening over the accepted T-711 v11 build.

This layer is composition/editorial only. It does not recalculate or reinterpret any
protocol-v2.1 result, estimand, interval, denominator, censoring decision or registered
T-613 scientific asset. It performs only verified reader-facing corrections identified
by the final document audits.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

import t711_build_entry_v11 as v11


t711 = v11.t711
_previous_build = t711.builder.build


CAPTIONS: dict[int, str] = {
    1: "Ροή του πειραματικού πρωτοκόλλου· τα ακριβή checkpoints προηγούνται των αντιστοιχισμένων κλάδων FN/FD/AN/AD.",
    2: "Αντιστοίχιση κάθε ερευνητικού ερωτήματος στα προκαθορισμένα μεγέθη εκτίμησης και στα καταγεγραμμένα αποτελέσματα.",
    3: "Τελική ονομαστική απόδοση ανά μέθοδο με σημειακά διαστήματα Student-t 95% σε 12 ανεξάρτητες επαναλήψεις.",
    4: "Μέση απόδοση στον άξονα αλληλεπιδράσεων, ξεχωριστά από την τελική επίδοση.",
    5: "Όφελος προσαρμογής ανά μέθοδο και συνθήκη· θετική τιμή σημαίνει μικρότερη απώλεια στην Adaptive από ό,τι στην Frozen εκτέλεση.",
    6: "Η απώλεια στην Frozen και στην Adaptive εκτέλεση παρουσιάζεται ως δύο διακριτά μεγέθη.",
    7: "Όφελος προσαρμογής ανά συνθήκη σε μικρά πολλαπλά γραφήματα με κοινές κλίμακες.",
    8: "Καταγεγραμμένες τροχιές κατευθυνόμενης απόκλισης ανά root για την κύρια οικογένεια action-remap· η διακεκομμένη γραμμή αντιστοιχεί στην ανοχή 0,10.",
    9: "Αναλογία roots που πέτυχαν σταθερή ανάκαμψη στην κύρια ανοχή 0,10· οι υπόλοιπες παραμένουν δεξιά λογοκριμένες (right-censored).",
    10: "Περιορισμένος χρόνος καθυστέρησης ανάκαμψης έως το όριο των 256 αλληλεπιδράσεων, με ρητή διατήρηση της λογοκρισίας.",
    11: "Παρατηρούμενος χρόνος ανάκαμψης μόνο μεταξύ των roots που ανέκαμψαν· το εμφανιζόμενο n αποκλείει τις right-censored roots.",
    12: "Προκαθορισμένη ανάλυση ευαισθησίας για ανοχές 0,05/0,10/0,20· κύρια ανοχή παραμένει το 0,10.",
    13: "Κατανομή τελικής ονομαστικής απόδοσης ανά root μετά από ισότιμη μείωση ως προς τα layouts.",
    14: "Κατανομή της μέσης απόδοσης στον άξονα αλληλεπιδράσεων ανά root μετά από ισότιμη μείωση ως προς τα layouts.",
    15: "Προκαθορισμένες paired συγκρίσεις A−B ανά root· τα διαστήματα είναι περιγραφικά σημειακά t intervals.",
    16: "Κατανομές του οφέλους προσαρμογής ανά root, μέθοδο και συνθήκη.",
    17: "Αριθμητικός θερμικός χάρτης του μέσου καταγεγραμμένου οφέλους προσαρμογής· υψηλότερη τιμή σημαίνει μεγαλύτερο όφελος.",
    18: "Όλες οι προκαθορισμένες paired συγκρίσεις A−B του RQ2 ανά συνθήκη και estimand.",
    19: "Paired διαγνωστικά Frozen-to-Adaptive ανά root· οι συνδετικές γραμμές δεν ορίζουν πρόσθετο estimand.",
    20: "Αναλυτικές καταγεγραμμένες τροχιές ανά root, μέθοδο και κύρια συνθήκη· δεν εφαρμόζεται averaging μεταξύ τροχιών.",
    21: "Σύνθεση των roots που ανέκαμψαν και των right-censored roots στην ανοχή 0,10.",
    22: "Προκαθορισμένες A−B συγκρίσεις για recovery status και restricted delay στον κύριο άξονα ανάκαμψης.",
    23: "Καταγεγραμμένοι χρόνοι ανάκαμψης και επιβεβαίωσης· οι right-censored roots σημειώνονται στο horizon χωρίς κατασκευή ψευδούς χρόνου ανάκαμψης.",
    24: "Γραμμή ιχνηλασιμότητας του αποδεκτού evidence· η αποτυχημένη εκτέλεση 216 jobs παραμένει αποκλεισμένη.",
}

REPLACEMENTS: dict[str, str] = {
    "Η Deep Q-Network (DQN) επεκτείνει τη value-based λογική σε νευρωνική προσέγγιση της action-value function και αποτελεί βασικό σημείο αναφοράς του deep RL [10]. Η DQN χρησιμοποιεί experience replay ώστε οι μεταβάσεις να αποθηκεύονται και να επαναχρησιμοποιούνται για updates, και target network ώστε οι bootstrapped targets να μεταβάλλονται πιο ελεγχόμενα.":
        "Η Deep Q-Network (DQN) επεκτείνει τη value-based λογική σε νευρωνική προσέγγιση της action-value function και αποτελεί βασικό σημείο αναφοράς του deep RL [10]. Η DQN χρησιμοποιεί experience replay ώστε οι μεταβάσεις να αποθηκεύονται και να επαναχρησιμοποιούνται για ενημερώσεις [10].",
    "Οι paired τελικές συγκρίσεις μεταξύ Q-Learning, SARSA και Dyna-Q+ ήταν ακριβώς μηδενικές. Κάθε μία από τις τρεις υπερείχε της DQN κατά 1,762 return units όταν η διαφορά εκφράζεται ως method-minus-DQN, με paired interval [0,637, 2,888]. Η αντίστοιχη διαφορά έναντι PPO ήταν επίσης 1,762, με interval [0,469, 3,056]. Η DQN έναντι PPO είχε μέση paired διαφορά 0,000 και ευρύ interval [-2,112, 2,112].":
        "Οι paired τελικές συγκρίσεις μεταξύ Q-Learning, SARSA και Dyna-Q+ ήταν ακριβώς μηδενικές. Για καθεμία από τις τρεις, η εκτιμώμενη method-minus-DQN διαφορά ήταν 1,762 return units, με paired interval [0,637, 2,888]. Η αντίστοιχη εκτιμώμενη διαφορά έναντι PPO ήταν επίσης 1,762, με interval [0,469, 3,056]. Η DQN έναντι PPO είχε μέση paired διαφορά 0,000 και ευρύ interval [-2,112, 2,112].",
    "Στις root-paired συγκρίσεις time-average, η Dyna-Q+ υπερείχε της Q-Learning κατά 1,143 [0,936, 1,350], της SARSA κατά 1,126 [0,949, 1,303], της DQN κατά 2,377 [1,441, 3,314] και της PPO κατά 2,419 [1,352, 3,486]. Η διαφορά Q-Learning έναντι SARSA ήταν -0,017 [-0,294, 0,260], ενώ η DQN έναντι PPO ήταν 0,042 [-1,763, 1,847].":
        "Στις root-paired συγκρίσεις time-average, οι εκτιμώμενες διαφορές Dyna-Q+−Q-Learning, Dyna-Q+−SARSA, Dyna-Q+−DQN και Dyna-Q+−PPO ήταν αντίστοιχα 1,143 [0,936, 1,350], 1,126 [0,949, 1,303], 2,377 [1,441, 3,314] και 2,419 [1,352, 3,486]. Η διαφορά Q-Learning έναντι SARSA ήταν -0,017 [-0,294, 0,260], ενώ η DQN έναντι PPO ήταν 0,042 [-1,763, 1,847].",
    "Η Q-Learning υπερείχε της SARSA κατά 8,879 [3,432, 14,327]. Η SARSA υπερείχε της PPO κατά 14,300 [10,146, 18,454]. Η διαφορά Dyna-Q+−SARSA ήταν -4,073 [-9,528, 1,382] και παρέμεινε αβέβαιη. Η DQN−PPO ήταν 2,237 [-3,639, 8,114], επίσης με interval που περιλαμβάνει το μηδέν.":
        "Η εκτιμώμενη paired διαφορά Q-Learning−SARSA ήταν 8,879 [3,432, 14,327], ενώ η SARSA−PPO ήταν 14,300 [10,146, 18,454]. Η διαφορά Dyna-Q+−SARSA ήταν -4,073 [-9,528, 1,382] και παρέμεινε αβέβαιη. Η DQN−PPO ήταν 2,237 [-3,639, 8,114], επίσης με interval που περιλαμβάνει το μηδέν.",
}

EXPECTED_REVIEW_PLACEHOLDERS = {
    "Ονοματεπώνυμο φοιτητή: [να συμπληρωθεί από το επίσημο έντυπο]",
    "Student: [to be completed from the official form]",
    "[Θέση για την ακριβή επίσημη δήλωση πριν από την τελική υποβολή]",
}

PLACEHOLDER_PATTERNS = (
    re.compile(r"\bTODO\b", re.I),
    re.compile(r"\bTBD\b", re.I),
    re.compile(r"\[να συμπληρωθεί[^\]]*\]", re.I),
    re.compile(r"\[to be completed[^\]]*\]", re.I),
    re.compile(r"\[Θέση[^\]]*\]", re.I),
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }


def _caption_number(paragraph: Paragraph) -> int | None:
    if paragraph.style.name != "Caption":
        return None
    match = re.match(r"^Σχήμα\s+(\d+)\s+—\s+", paragraph.text.strip())
    return int(match.group(1)) if match else None


def _insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_xml = OxmlElement("w:p")
    paragraph._p.addnext(new_xml)
    new_paragraph = Paragraph(new_xml, paragraph._parent)
    new_paragraph.style = "Normal"
    run = new_paragraph.add_run(text)
    run.italic = True
    return new_paragraph


def _localize_cached_figure_list(doc: Document) -> int:
    for paragraph in doc.paragraphs:
        instructions = [node.text or "" for node in paragraph._p.xpath(".//w:instrText")]
        if not any('\\c "Σχήμα"' in item for item in instructions):
            continue
        text_nodes = paragraph._p.xpath(".//w:t")
        if len(text_nodes) != len(CAPTIONS):
            raise RuntimeError(
                f"Unexpected cached List-of-Figures entry count: {len(text_nodes)}"
            )
        for number, node in enumerate(text_nodes, start=1):
            node.text = f"Σχήμα {number} — {CAPTIONS[number]}"
        return len(text_nodes)
    raise RuntimeError("Cached List of Figures field was not found")


def _assign_alt_text(doc: Document) -> int:
    paragraphs = doc.paragraphs
    assigned = 0
    synthesis_used = False
    for index, paragraph in enumerate(paragraphs):
        if not paragraph._p.xpath(".//w:drawing"):
            continue
        description: str | None = None
        for lookahead in range(index + 1, min(index + 4, len(paragraphs))):
            number = _caption_number(paragraphs[lookahead])
            if number in CAPTIONS:
                description = f"Σχήμα {number}: {CAPTIONS[number]}"
                break
        if description is None:
            if not synthesis_used:
                description = (
                    "Αποτελέσματα με μία ματιά: σύνθεση των κύριων ευρημάτων των "
                    "RQ1, RQ2 και RQ3 από τα αποδεκτά αποτελέσματα."
                )
                synthesis_used = True
            else:
                description = (
                    "Επιστημονικό διάγραμμα της διπλωματικής· η πλήρης περιγραφή "
                    "δίνεται στην αμέσως επόμενη λεζάντα."
                )
        for docpr in paragraph._p.xpath(".//wp:docPr"):
            docpr.set("descr", description)
            docpr.set("title", description.split(":", 1)[0])
            assigned += 1
    return assigned


def _placeholder_hits(doc: Document) -> list[str]:
    hits: list[str] = []
    for paragraph in doc.paragraphs:
        if any(pattern.search(paragraph.text) for pattern in PLACEHOLDER_PATTERNS):
            hits.append(paragraph.text)
    return hits


def _generator_metadata_residue(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        core = archive.read("docProps/core.xml").decode("utf-8")
    residue: list[str] = []
    for phrase in ("python-docx", "generated by python-docx", "2013-12-23"):
        if phrase in core:
            residue.append(phrase)
    return residue


def _build(output: Path, qa_output: Path) -> None:
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    media_before = _media_hashes(output)
    doc = Document(output)

    replacement_count = 0
    for paragraph in doc.paragraphs:
        replacement = REPLACEMENTS.get(paragraph.text)
        if replacement is not None:
            paragraph.text = replacement
            replacement_count += 1

    caption_count = 0
    rq3_caption: Paragraph | None = None
    for paragraph in doc.paragraphs:
        number = _caption_number(paragraph)
        if number not in CAPTIONS:
            continue
        if len(paragraph.runs) < 3:
            raise RuntimeError(
                f"Unexpected caption run structure for figure {number}: {paragraph.text!r}"
            )
        paragraph.runs[2].text = " — " + CAPTIONS[number]
        for extra_run in paragraph.runs[3:]:
            extra_run.text = ""
        caption_count += 1
        if number == 11:
            rq3_caption = paragraph

    if rq3_caption is None:
        raise RuntimeError("Figure 11 conditional-recovery caption was not found")
    _insert_after(
        rq3_caption,
        "Σημείωση ερμηνείας: ο χρόνος αυτός είναι υπό συνθήκη ανάκαμψης "
        "(conditional on recovery). Όταν το recovered n είναι πολύ μικρό, ιδίως "
        "n=1 ή n=2, η εκτίμηση είναι ασταθής και πρέπει να διαβάζεται μαζί με "
        "το recovery proportion και το restricted delay.",
    )

    cached_figure_count = _localize_cached_figure_list(doc)
    alt_text_count = _assign_alt_text(doc)

    props = doc.core_properties
    props.title = (
        "Σύγκριση και Αξιολόγηση Ανθεκτικών Πρακτόρων Τεχνητής Νοημοσύνης "
        "σε Περιβάλλοντα με Αβεβαιότητα"
    )
    props.subject = (
        "Διπλωματική εργασία — ενισχυτική μάθηση, ανθεκτικότητα και "
        "προσαρμογή υπό αβεβαιότητα"
    )
    props.author = ""
    props.last_modified_by = ""
    props.keywords = (
        "reinforcement learning; resilient AI agents; non-stationarity; adaptation; recovery"
    )
    props.comments = ""
    props.category = "Diploma Thesis"
    props.created = datetime(2026, 9, 3, 0, 0, tzinfo=timezone.utc)
    props.modified = datetime(2026, 9, 3, 0, 0, tzinfo=timezone.utc)

    doc.save(output)

    final_doc = Document(output)
    media_after = _media_hashes(output)
    placeholder_hits = _placeholder_hits(final_doc)
    placeholder_set = set(placeholder_hits)
    final_mode = os.environ.get("T711_FINAL_MODE", "0").strip().lower() in {
        "1",
        "true",
        "yes",
    }
    superiority_residue = [
        paragraph.text
        for paragraph in final_doc.paragraphs
        if "υπερείχε" in paragraph.text
    ]
    metadata_residue = _generator_metadata_residue(output)
    final_sha = _sha256(output)
    final_paragraph_count = len(final_doc.paragraphs)

    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": final_paragraph_count,
            "post_synthesis_paragraph_count": final_paragraph_count,
            "t711a_hardening_version": 12,
            "t711a_reader_text_rescanned_after_save": True,
            "dqn_foundation_citation_wording_corrected": replacement_count == 4,
            "superiority_wording_neutralized_count": 3,
            "superiority_wording_residue": superiority_residue,
            "localized_figure_caption_count": caption_count,
            "localized_cached_figure_entry_count": cached_figure_count,
            "rq3_conditional_recovery_note_added": True,
            "image_alt_text_count": alt_text_count,
            "embedded_media_bytes_preserved": media_before == media_after,
            "docx_core_metadata_scrubbed": not metadata_residue,
            "generator_metadata_residue": metadata_residue,
            "placeholder_hits": placeholder_hits,
            "placeholder_count": len(placeholder_hits),
            "review_placeholders_expected": placeholder_set == EXPECTED_REVIEW_PLACEHOLDERS,
            "final_mode": final_mode,
            "final_submission_ready": final_mode and not placeholder_hits,
            "scientific_values_modified": False,
        }
    )

    hardening_ok = (
        report.get("status") == "pass"
        and replacement_count == 4
        and caption_count == cached_figure_count == 24
        and alt_text_count == 25
        and media_before == media_after
        and not superiority_residue
        and not metadata_residue
        and final_paragraph_count == 708
        and (
            (not final_mode and placeholder_set == EXPECTED_REVIEW_PLACEHOLDERS)
            or (final_mode and not placeholder_hits)
        )
        and report.get("registered_asset_bytes_modified") is False
    )
    if not hardening_ok:
        report["status"] = "fail"

    qa_output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if report["status"] != "pass":
        raise RuntimeError(f"T-711A v12 hardening failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
