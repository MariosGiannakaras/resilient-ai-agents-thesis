#!/usr/bin/env python3
"""T-714 glossary pagination hardening.

Keeps each glossary-term heading with the paragraph that begins its definition so a
term cannot be stranded at the bottom of a page. Scientific content and media are
untouched.
"""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document

import t711_build_entry_v16 as v16


t711 = v16.t711
_previous_build = v16._build

GLOSSARY_STYLE = "Glossary Term"
EXPECTED_GLOSSARY_TERM_COUNT = 24


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _build(output: Path, qa_output: Path) -> None:
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))

    doc = Document(output)
    glossary_terms = [p for p in doc.paragraphs if p.style.name == GLOSSARY_STYLE]
    if len(glossary_terms) != EXPECTED_GLOSSARY_TERM_COUNT:
        raise RuntimeError(
            f"Expected {EXPECTED_GLOSSARY_TERM_COUNT} glossary terms, found {len(glossary_terms)}"
        )

    for paragraph in glossary_terms:
        paragraph.paragraph_format.keep_with_next = True

    doc.save(output)

    final_doc = Document(output)
    final_glossary_terms = [
        p for p in final_doc.paragraphs if p.style.name == GLOSSARY_STYLE
    ]
    glossary_terms_without_keep = [
        p.text for p in final_glossary_terms if p.paragraph_format.keep_with_next is not True
    ]

    final_sha = _sha256(output)
    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": len(final_doc.paragraphs),
            "post_synthesis_paragraph_count": len(final_doc.paragraphs),
            "t711a_hardening_version": 17,
            "glossary_term_count": len(final_glossary_terms),
            "glossary_terms_keep_with_next": not glossary_terms_without_keep,
            "glossary_terms_without_keep_with_next": glossary_terms_without_keep,
            "scientific_values_modified": False,
            "registered_asset_bytes_modified": False,
            "final_visual_qa_required": True,
        }
    )
    if (
        len(final_glossary_terms) != EXPECTED_GLOSSARY_TERM_COUNT
        or glossary_terms_without_keep
    ):
        report["status"] = "fail"

    qa_output.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    if report.get("status") != "pass":
        raise RuntimeError(f"T-714 v17 glossary pagination hardening failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
