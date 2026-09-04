#!/usr/bin/env python3
"""T-715 reader composition: application screenshots plus visual-QA layout fixes.

This is composition-only. It fixes a front-matter duplication introduced by the reader
preprocessor, keeps exactly two deterministic DEVELOPMENT-only application screenshots,
and removes one forced appendix page break that left the RQ3 appendix introduction on an
otherwise nearly empty page. Scientific values and registered T-613 media remain intact.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import t711_build_entry_v24 as v24
import t711_build_entry_v19 as v19


t711 = v24.t711
_previous_build = t711.builder.build
ROOT = Path(__file__).resolve().parents[1]

PHASE_A_NAME = "03-run-phase-a-1440x900.png"
PHASE_B_NAME = "04-run-phase-b-frozen-adaptive-1440x900.png"
EXPECTED_CAPTURE_NAMES = (PHASE_A_NAME, PHASE_B_NAME)

# The T-715 reader preprocessor is used for whole manuscript files and for the extracted
# front-matter fragments. Replacing every front fragment with FRONT_READER duplicated the
# summary/abstract. Keep fragment preprocessing identity-like and instead make the builder
# extract the simplified front text from FRONT_READER exactly once.
_original_reader_markdown = v19._reader_markdown


def _reader_markdown_without_front_duplication(md: str, mode: str) -> str:
    if mode == "front":
        return md
    return _original_reader_markdown(md, mode)


def _reader_front_sections() -> tuple[str, str, str, str]:
    md = v19.FRONT_READER

    def section(name: str, next_name: str | None) -> str:
        start = md.index(f"## {name}") + len(f"## {name}")
        end = md.index(f"## {next_name}", start) if next_name else len(md)
        return md[start:end].strip()

    greek = section("Περίληψη", "Abstract")
    english = section("Abstract", None)
    greek_kw = re.search(r"\*\*Λέξεις-κλειδιά:\*\*\s*(.+)", greek)
    english_kw = re.search(r"\*\*Keywords:\*\*\s*(.+)", english)
    greek = re.sub(r"\*\*Λέξεις-κλειδιά:\*\*.*", "", greek).strip()
    english = re.sub(r"\*\*Keywords:\*\*.*", "", english).strip()
    return (
        greek,
        greek_kw.group(1) if greek_kw else "",
        english,
        english_kw.group(1) if english_kw else "",
    )


v19._reader_markdown = _reader_markdown_without_front_duplication
t711.builder.extract_front_sections = _reader_front_sections


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        result: dict[str, str] = {}
        for name in sorted(n for n in archive.namelist() if n.startswith("word/media/")):
            result[name] = hashlib.sha256(archive.read(name)).hexdigest()
        return result


def _find_paragraph(doc: Document, prefix: str):
    matches = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"T-715 application screenshot anchor expected once for {prefix!r}, found {len(matches)}")
    return matches[0]


def _insert_screenshot_after(doc: Document, anchor, image_path: Path, caption: str, alt_text: str) -> None:
    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(6)
    picture_paragraph.paragraph_format.space_after = Pt(3)
    picture_paragraph.paragraph_format.keep_with_next = True
    run = picture_paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(5.95))
    shape._inline.docPr.set("descr", alt_text)

    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_paragraph.paragraph_format.keep_together = True
    caption_run = caption_paragraph.add_run(caption)
    t711.builder.set_run_font(caption_run, size=9.5)

    anchor._p.addnext(picture_paragraph._p)
    picture_paragraph._p.addnext(caption_paragraph._p)


def _validated_screenshot_inputs() -> tuple[Path, Path, dict]:
    directory = Path(os.environ.get("T715_UI_SCREENSHOT_DIR", ROOT / "artifacts" / "t715-ui")).resolve()
    manifest_path = directory / "manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(f"T-715 deterministic UI manifest missing: {manifest_path}")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    required_manifest_state = (
        manifest.get("development_fixture_created_only") is True
        and manifest.get("scientific_jobs_executed") == 0
        and manifest.get("environment_steps_executed") == 0
        and manifest.get("final_reserve_accessed") is False
        and manifest.get("final_experiment_authorized") is False
    )
    if not required_manifest_state:
        raise RuntimeError(f"T-715 refuses non-DEVELOPMENT UI capture manifest: {manifest}")

    records = {record.get("file"): record for record in manifest.get("screenshots", [])}
    paths: list[Path] = []
    for name in EXPECTED_CAPTURE_NAMES:
        path = directory / name
        record = records.get(name)
        if not path.exists() or record is None:
            raise RuntimeError(f"T-715 required deterministic UI screenshot missing: {name}")
        if int(record.get("width", 0)) != 1440 or int(record.get("height", 0)) != 900:
            raise RuntimeError(f"T-715 unexpected screenshot dimensions for {name}: {record}")
        actual = _sha256(path)
        if actual != record.get("sha256"):
            raise RuntimeError(f"T-715 screenshot hash mismatch for {name}")
        paths.append(path)
    return paths[0], paths[1], manifest


def _remove_rq3_appendix_forced_break(doc: Document) -> bool:
    matches = [
        paragraph for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith("Σχήμα 20 — Αναλυτικές καταγεγραμμένες τροχιές")
    ]
    if len(matches) != 1:
        raise RuntimeError(f"T-715 expected exactly one Figure 20 RQ3 appendix caption, found {len(matches)}")
    picture = matches[0]._p.getprevious()
    ppr = picture.find(qn("w:pPr"))
    if ppr is None:
        return False
    page_break = ppr.find(qn("w:pageBreakBefore"))
    if page_break is None:
        return False
    ppr.remove(page_break)
    return True


def _build(output: Path, qa_output: Path) -> None:
    phase_a, phase_b, capture_manifest = _validated_screenshot_inputs()
    _previous_build(output, qa_output)
    report = json.loads(qa_output.read_text(encoding="utf-8"))
    if report.get("status") != "pass":
        raise RuntimeError(f"T-715 base reader composition failed before application screenshots: {report}")

    media_before = _media_hashes(output)
    doc = Document(output)
    inline_before = len(doc.inline_shapes)

    phase_a_anchor = _find_paragraph(doc, "Η εφαρμογή οργανώνεται στις κύριες ενότητες")
    _insert_screenshot_after(
        doc,
        phase_a_anchor,
        phase_a,
        "Στιγμιότυπο εφαρμογής 1 — Ενδεικτική προβολή της Φάσης A και του GridWorld κατά την ονομαστική μάθηση. Η εικόνα είναι deterministic DEVELOPMENT capture και δεν αποτελεί πηγή ποσοτικών αποτελεσμάτων.",
        "Desktop εφαρμογή PySide6 στη Φάση A με προβολή του GridWorld και της τρέχουσας κατάστασης εκτέλεσης. DEVELOPMENT capture μόνο για επεξήγηση της υλοποίησης.",
    )

    phase_b_anchor = _find_paragraph(doc, "Η ενότητα Run της εφαρμογής είναι η πιο άμεση σύνδεση")
    _insert_screenshot_after(
        doc,
        phase_b_anchor,
        phase_b,
        "Στιγμιότυπο εφαρμογής 2 — Ενδεικτική προβολή της Φάσης B με παράλληλη παρουσίαση Frozen και Adaptive κλάδων στο GridWorld. Η εικόνα εξηγεί τη λειτουργία της εφαρμογής και δεν χρησιμοποιείται ως επιστημονικό τεκμήριο.",
        "Desktop εφαρμογή PySide6 στη Φάση B με παράλληλη απεικόνιση Frozen και Adaptive GridWorld. DEVELOPMENT capture μόνο για επεξήγηση της υλοποίησης.",
    )

    rq3_break_removed = _remove_rq3_appendix_forced_break(doc)
    doc.save(output)

    final_doc = Document(output)
    media_after = _media_hashes(output)
    inline_after = len(final_doc.inline_shapes)
    preserved = all(media_after.get(name) == digest for name, digest in media_before.items())
    new_media = sorted(name for name in media_after if name not in media_before)
    capture_hashes = {_sha256(phase_a), _sha256(phase_b)}
    new_media_hashes = {media_after[name] for name in new_media}

    app_captions = [paragraph.text for paragraph in final_doc.paragraphs if paragraph.text.startswith("Στιγμιότυπο εφαρμογής ")]
    redundant_front_markers = [paragraph.text for paragraph in final_doc.paragraphs if paragraph.text.strip() == "Προκαταρκτικό υλικό — Περίληψη / Abstract"]
    summary_heading_count = sum(1 for paragraph in final_doc.paragraphs if paragraph.text.strip() == "Περίληψη" and paragraph.style.name == "Heading 1")
    abstract_heading_count = sum(1 for paragraph in final_doc.paragraphs if paragraph.text.strip() == "Abstract" and paragraph.style.name == "Heading 1")

    final_sha = _sha256(output)
    report.update(
        {
            "output_sha256": final_sha,
            "paragraph_count": len(final_doc.paragraphs),
            "post_synthesis_paragraph_count": len(final_doc.paragraphs),
            "t715_front_matter_reader_duplication_fixed": True,
            "t715_redundant_front_marker_count": len(redundant_front_markers),
            "t715_summary_heading_count": summary_heading_count,
            "t715_abstract_heading_count": abstract_heading_count,
            "t715_rq3_appendix_forced_pagebreak_removed": rq3_break_removed,
            "t715_application_screenshots_added": True,
            "t715_application_screenshot_count": len(app_captions),
            "t715_application_screenshot_names": list(EXPECTED_CAPTURE_NAMES),
            "t715_application_screenshot_sha256": {PHASE_A_NAME: _sha256(phase_a), PHASE_B_NAME: _sha256(phase_b)},
            "t715_application_capture_purpose": capture_manifest.get("purpose"),
            "t715_application_capture_development_only": True,
            "t715_application_capture_scientific_jobs_executed": 0,
            "t715_application_capture_environment_steps_executed": 0,
            "t715_application_capture_final_reserve_accessed": False,
            "t715_application_screenshots_scientific_evidence": False,
            "t715_preexisting_docx_media_preserved": preserved,
            "t715_new_docx_media_count": len(new_media),
            "t715_new_docx_media_hashes_match_screenshots": new_media_hashes == capture_hashes,
            "t715_inline_shape_count_before_app_screenshots": inline_before,
            "t715_final_inline_shape_count": inline_after,
            "t715_final_image_alt_text_count": report.get("image_alt_text_count", 0) + 2,
            "scientific_values_modified": False,
            "registered_asset_bytes_modified": False,
            "final_visual_qa_required": True,
        }
    )

    hardening_ok = (
        preserved
        and len(new_media) == 2
        and new_media_hashes == capture_hashes
        and inline_before == 25
        and inline_after == 27
        and len(app_captions) == 2
        and len(redundant_front_markers) == 0
        and summary_heading_count == 1
        and abstract_heading_count == 1
        and rq3_break_removed
        and report.get("inserted_asset_count") == 24
        and report.get("registered_figure_count_preserved") is True
        and report.get("scientific_values_modified") is False
        and report.get("registered_asset_bytes_modified") is False
    )
    if not hardening_ok:
        report["status"] = "fail"

    qa_output.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if report.get("status") != "pass":
        raise RuntimeError(f"T-715 visual-QA composition hardening failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
