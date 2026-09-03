#!/usr/bin/env python3
"""T-711 reader-aid synthesis infographic.

Adds one unnumbered, composition-only "results at a glance" graphic immediately after
Chapter 5.1. The graphic summarizes already accepted RQ1/RQ2/RQ3 primary values and
recovery incidence from the frozen T-612/T-613 result set. It is deliberately *not* a
registered T-613 figure, does not enter the automatic List of Figures, does not change
any scientific value, and does not introduce a composite score/ranking or new estimand.

The source manuscript hash is pinned so the summary cannot silently drift if Chapter 5
changes later. The existing 24 registered figures and their bytes remain untouched.
"""

from __future__ import annotations

import hashlib
import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont

import t711_build_entry_v9 as v9


t711 = v9.t711
_previous_build = t711.builder.build

REPO_ROOT = Path(__file__).resolve().parents[1]
CHAPTER_5 = REPO_ROOT / "docs/thesis/draft/CHAPTER_05_RESULTS.md"
EXPECTED_CHAPTER_5_SHA256 = "08cec7f0c7dba6ced80bb01c749e1fb76f20e52b830ad3ff1c3f24581ed694a4"
ANCHOR = "5.2 RQ1 — Ονομαστική μάθηση"
ALT_TEXT = (
    "Συνοπτική οπτική σύνθεση των αποδεκτών RQ1, RQ2 και RQ3 αποτελεσμάτων. "
    "Παρουσιάζει final και time-average RQ1 summaries, adaptation benefit για τα δύο "
    "persistent action remaps και recovered roots στο primary recovery tolerance 0,10. "
    "Δεν αποτελεί νέο estimand ή composite ranking."
)

METHOD_COLORS = {
    "Q-Learning": (0, 114, 178),
    "SARSA": (230, 159, 0),
    "Dyna-Q+": (0, 158, 115),
    "DQN": (204, 121, 167),
    "PPO": (213, 94, 0),
}


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _font_path(*, bold: bool) -> Path:
    candidates = (
        [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
        ]
        if bold
        else [
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
        ]
    )
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise RuntimeError("T-711 results synthesis requires a Unicode-capable DejaVu/Liberation font")


def _font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(_font_path(bold=bold)), size=size)


def _render_results_synthesis(path: Path) -> None:
    width, height = 1800, 860
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = _font(44, bold=True)
    panel_font = _font(29, bold=True)
    body_font = _font(23)
    small_font = _font(19)
    number_font = _font(25, bold=True)

    title = "Αποτελέσματα με μία ματιά"
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(
        ((width - (title_box[2] - title_box[0])) / 2, 26),
        title,
        font=title_font,
        fill=(20, 20, 20),
    )

    margin, gap, top, bottom = 55, 28, 100, 740
    panel_width = (width - 2 * margin - 2 * gap) // 3
    panels: list[tuple[int, int, int, int]] = []
    for index in range(3):
        x0 = margin + index * (panel_width + gap)
        x1 = x0 + panel_width
        panels.append((x0, top, x1, bottom))
        draw.rounded_rectangle(
            (x0, top, x1, bottom),
            radius=24,
            fill=(247, 249, 251),
            outline=(190, 200, 210),
            width=3,
        )

    # RQ1: accepted final and interaction-axis time-average summaries.
    x0, y0, x1, y1 = panels[0]
    draw.text((x0 + 25, y0 + 20), "RQ1 · Ονομαστική μάθηση", font=panel_font, fill=(20, 20, 20))
    draw.text((x0 + 25, y0 + 73), "Τελικό probe (8.192 interactions)", font=small_font, fill=(70, 70, 70))
    draw.text((x0 + 25, y0 + 111), "Q-Learning · SARSA · Dyna-Q+", font=body_font, fill=(20, 20, 20))
    value = "−0,100"
    box = draw.textbbox((0, 0), value, font=number_font)
    draw.text((x1 - 30 - (box[2] - box[0]), y0 + 108), value, font=number_font, fill=METHOD_COLORS["Q-Learning"])
    draw.text((x0 + 25, y0 + 150), "DQN · PPO", font=body_font, fill=(20, 20, 20))
    value = "−1,862"
    box = draw.textbbox((0, 0), value, font=number_font)
    draw.text((x1 - 30 - (box[2] - box[0]), y0 + 147), value, font=number_font, fill=METHOD_COLORS["DQN"])
    draw.line((x0 + 25, y0 + 197, x1 - 25, y0 + 197), fill=(210, 215, 220), width=2)
    draw.text((x0 + 25, y0 + 218), "Time-average return", font=small_font, fill=(70, 70, 70))

    rq1_time = [
        ("Dyna-Q+", "−0,485"),
        ("SARSA", "−1,611"),
        ("Q-Learning", "−1,628"),
        ("DQN", "−2,862"),
        ("PPO", "−2,904"),
    ]
    y = y0 + 258
    for method, value in rq1_time:
        draw.ellipse((x0 + 28, y + 4, x0 + 42, y + 18), fill=METHOD_COLORS[method])
        draw.text((x0 + 55, y - 1), method, font=body_font, fill=(30, 30, 30))
        value_box = draw.textbbox((0, 0), value, font=body_font)
        draw.text((x1 - 30 - (value_box[2] - value_box[0]), y - 1), value, font=body_font, fill=(30, 30, 30))
        y += 48
    draw.text((x0 + 25, y1 - 44), "μεγαλύτερο = καλύτερο", font=small_font, fill=(90, 90, 90))

    # RQ2: accepted adaptation-benefit means for the two persistent remap conditions.
    x0, y0, x1, y1 = panels[1]
    draw.text((x0 + 25, y0 + 20), "RQ2 · Όφελος προσαρμογής", font=panel_font, fill=(20, 20, 20))
    draw.text((x0 + 25, y0 + 73), "Persistent action remaps", font=small_font, fill=(70, 70, 70))

    cycle = [
        ("Q-Learning", 32.269),
        ("SARSA", 31.127),
        ("Dyna-Q+", 26.102),
        ("DQN", 6.623),
        ("PPO", 0.060),
    ]
    swap = [
        ("Q-Learning", 22.665),
        ("SARSA", 13.785),
        ("Dyna-Q+", 9.712),
        ("DQN", 1.723),
        ("PPO", -0.515),
    ]

    def mini_group(y: int, label: str, data: list[tuple[str, float]]) -> int:
        draw.text((x0 + 25, y), label, font=body_font, fill=(30, 30, 30))
        y += 35
        label_x, zero, end = x0 + 25, x0 + 180, x1 - 48
        draw.line((zero, y - 3, zero, y + 5 * 34 - 9), fill=(125, 125, 125), width=2)
        scale = (end - zero) / 35.0
        for index, (method, numeric) in enumerate(data):
            row_y = y + index * 34
            draw.text((label_x, row_y - 2), method, font=small_font, fill=METHOD_COLORS[method])
            value_x = zero + numeric * scale
            if numeric >= 0:
                draw.rectangle((zero, row_y + 7, value_x, row_y + 20), fill=METHOD_COLORS[method])
            else:
                draw.rectangle((value_x, row_y + 7, zero, row_y + 20), fill=METHOD_COLORS[method])
            label_value = f"{numeric:+.2f}".replace("-", "−").replace(".", ",")
            label_box = draw.textbbox((0, 0), label_value, font=small_font)
            text_x = value_x + 6 if numeric >= 0 else value_x - (label_box[2] - label_box[0]) - 6
            draw.text((text_x, row_y - 3), label_value, font=small_font, fill=(35, 35, 35))
        return y + 5 * 34

    next_y = mini_group(y0 + 112, "Action-remap cycle", cycle)
    mini_group(next_y + 24, "Action-remap swap", swap)
    draw.text((x0 + 25, y1 - 44), "θετικό = Adaptive loss < Frozen loss", font=small_font, fill=(90, 90, 90))

    # RQ3: primary recovery incidence only; timing/censoring semantics remain explicit.
    x0, y0, x1, y1 = panels[2]
    draw.text((x0 + 25, y0 + 20), "RQ3 · Ανάκαμψη @ τ=0,10", font=panel_font, fill=(20, 20, 20))
    draw.text((x0 + 25, y0 + 73), "Recovered roots / 12", font=small_font, fill=(70, 70, 70))
    draw.text((x0 + 270, y0 + 112), "cycle", font=small_font, fill=(70, 70, 70))
    draw.text((x0 + 405, y0 + 112), "swap", font=small_font, fill=(70, 70, 70))

    rq3 = [
        ("Q-Learning", 12, 12),
        ("SARSA", 12, 12),
        ("Dyna-Q+", 12, 8),
        ("DQN", 2, 8),
        ("PPO", 1, 4),
    ]
    y = y0 + 150
    for method, cycle_n, swap_n in rq3:
        draw.ellipse((x0 + 28, y + 5, x0 + 42, y + 19), fill=METHOD_COLORS[method])
        draw.text((x0 + 55, y - 1), method, font=body_font, fill=(30, 30, 30))
        draw.text((x0 + 288, y - 3), f"{cycle_n}/12", font=number_font, fill=METHOD_COLORS[method])
        draw.text((x0 + 423, y - 3), f"{swap_n}/12", font=number_font, fill=METHOD_COLORS[method])
        y += 55

    draw.line((x0 + 25, y + 12, x1 - 25, y + 12), fill=(210, 215, 220), width=2)
    draw.text((x0 + 25, y + 34), "Censoring:", font=small_font, fill=(70, 70, 70))
    draw.text((x0 + 135, y + 34), "unrecovered roots → recovery_time = null", font=small_font, fill=(30, 30, 30))
    draw.text((x0 + 25, y + 69), "Το 256 δεν εμφανίζεται ως fake recovery time.", font=small_font, fill=(30, 30, 30))
    draw.text((x0 + 25, y1 - 44), "recovery incidence ≠ recovery speed", font=small_font, fill=(90, 90, 90))

    footer = "Σύνοψη ήδη προδηλωμένων estimands · όχι composite ranking · καμία νέα ανάλυση ή επαναϋπολογισμός"
    footer_box = draw.textbbox((0, 0), footer, font=small_font)
    draw.text(
        ((width - (footer_box[2] - footer_box[0])) / 2, 800),
        footer,
        font=small_font,
        fill=(90, 90, 90),
    )

    image.save(path, format="PNG", optimize=False)


def _insert_results_synthesis(output: Path, png_path: Path) -> int:
    doc = Document(output)
    anchors = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == ANCHOR]
    if len(anchors) != 1:
        raise RuntimeError(f"T-711 results synthesis anchor expected once, found {len(anchors)}")

    before_count = len(doc.inline_shapes)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(png_path), width=Inches(6.1))
    shape._inline.docPr.set("descr", ALT_TEXT)

    anchors[0]._p.addprevious(paragraph._p)
    doc.save(output)

    reopened = Document(output)
    after_count = len(reopened.inline_shapes)
    if after_count != before_count + 1:
        raise RuntimeError(
            f"T-711 results synthesis expected one additional inline shape: before={before_count}, after={after_count}"
        )
    return after_count


def _build(output: Path, qa_output: Path) -> None:
    chapter_sha = _sha256(CHAPTER_5)
    if chapter_sha != EXPECTED_CHAPTER_5_SHA256:
        raise RuntimeError(
            "T-711 results synthesis source drift: Chapter 5 changed after the reviewed v9 manuscript; "
            f"expected {EXPECTED_CHAPTER_5_SHA256}, got {chapter_sha}"
        )

    _previous_build(output, qa_output)

    with tempfile.TemporaryDirectory(prefix="t711-results-synthesis-") as temp_dir:
        png_path = Path(temp_dir) / "results-at-a-glance.png"
        _render_results_synthesis(png_path)
        png_sha = _sha256(png_path)
        inline_shape_count = _insert_results_synthesis(output, png_path)

    report = json.loads(qa_output.read_text(encoding="utf-8"))
    report.update(
        {
            "results_synthesis_infographic": True,
            "results_synthesis_numbered_figure": False,
            "results_synthesis_anchor": ANCHOR,
            "results_synthesis_basis": "selected accepted T-612/T-613 RQ1/RQ2/RQ3 summaries only; no recomputation or composite ranking",
            "results_synthesis_source_file": str(CHAPTER_5.relative_to(REPO_ROOT)),
            "results_synthesis_source_sha256": chapter_sha,
            "results_synthesis_png_sha256": png_sha,
            "registered_figure_count_preserved": report.get("inserted_asset_count") == 24,
            "inline_shape_count_after_synthesis": inline_shape_count,
            "registered_asset_bytes_modified": False,
            "scientific_values_modified": False,
            "final_visual_qa_required": True,
        }
    )

    if (
        report.get("status") != "pass"
        or report.get("inserted_asset_count") != 24
        or report.get("planned_figure_count") != 24
        or inline_shape_count != 25
        or report.get("reader_visible_asset_id_residue")
    ):
        report["status"] = "fail"

    qa_output.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if report["status"] != "pass":
        raise RuntimeError(f"T-711 v10 results-synthesis QA failed: {report}")


t711.builder.build = _build

if __name__ == "__main__":
    t711.builder.main()
