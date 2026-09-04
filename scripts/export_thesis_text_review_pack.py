from pathlib import Path
import argparse
import re, json, csv, hashlib, zipfile, shutil
from collections import defaultdict
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

parser = argparse.ArgumentParser(description='Export the complete reader-facing thesis into a provenance/rewrite review pack.')
parser.add_argument('--docx', type=Path, required=True)
parser.add_argument('--qa', type=Path, required=True)
parser.add_argument('--output', type=Path, required=True)
args = parser.parse_args()
DOCX = args.docx.resolve()
QA = args.qa.resolve()
OUT = args.output.resolve()
if OUT.exists(): shutil.rmtree(OUT)
(OUT/'chapters').mkdir(parents=True)
(OUT/'categories').mkdir()
(OUT/'sources').mkdir()
(OUT/'data').mkdir()

qa=json.loads(QA.read_text(encoding='utf-8'))
citation_map=qa['citation_map']
num_to_sid={n:sid for sid,n in citation_map.items()}
sha=hashlib.sha256(DOCX.read_bytes()).hexdigest()
doc=Document(DOCX)

# Bibliography entries from final Word, which are the exact reader-visible references.
bib={}
in_bib=False
for p in doc.paragraphs:
    t=p.text.strip()
    if t=='Βιβλιογραφία': in_bib=True; continue
    if t=='Παραρτήματα': in_bib=False
    if in_bib:
        m=re.match(r'^\[(\d+)\]\s+(.*)$',t)
        if m: bib[int(m.group(1))]=m.group(2)

# Ordered block iterator.
def iter_blocks(document):
    parent=document.element.body
    p_by_el={p._p:p for p in document.paragraphs}
    table_by_el={t._tbl:t for t in document.tables}
    for child in parent.iterchildren():
        if child.tag==qn('w:p'):
            yield ('p',p_by_el[child])
        elif child.tag==qn('w:tbl'):
            yield ('t',table_by_el[child])

file_key='00_front_matter'
major='Front matter'
subsection=''
seq=defaultdict(int)
records=[]
chapter_text=defaultdict(list)
category_text=defaultdict(list)
source_usage=defaultdict(list)

major_to_file={
 'Κεφάλαιο 1 — Εισαγωγή':'01_chapter_01_introduction',
 'Κεφάλαιο 2 — Θεωρητικό Υπόβαθρο και Σχετική Βιβλιογραφία':'02_chapter_02_background',
 'Κεφάλαιο 3 — Μεθοδολογία και Πειραματικός Σχεδιασμός':'03_chapter_03_methodology',
 'Κεφάλαιο 4 — Αρχιτεκτονική και Υλοποίηση του Συστήματος':'04_chapter_04_implementation',
 'Κεφάλαιο 5 — Πειραματικά Αποτελέσματα':'05_chapter_05_results',
 'Κεφάλαιο 6 — Συζήτηση':'06_chapter_06_discussion',
 'Κεφάλαιο 7 — Συμπεράσματα και Μελλοντική Εργασία':'07_chapter_07_conclusions',
 'Βιβλιογραφία':'08_bibliography',
 'Παραρτήματα':'09_appendices',
}

def refs_in(text):
    nums=[]
    for m in re.finditer(r'\[(\d+)\]',text):
        n=int(m.group(1))
        if n in num_to_sid: nums.append(n)
    return sorted(set(nums))

def labels_for(text, major, subsection, is_table=False):
    labels=[]
    refs=refs_in(text)
    if refs:
        labels += ['AI-assisted synthesis', 'externally supported']
    elif major.startswith('Κεφάλαιο 2'):
        labels += ['AI-assisted synthesis', 'theory/related-work synthesis']
    elif major.startswith('Κεφάλαιο 3'):
        labels += ['AI-assisted synthesis', 'project methodology/protocol fact']
    elif major.startswith('Κεφάλαιο 4'):
        labels += ['AI-assisted synthesis', 'project implementation fact']
    elif major.startswith('Κεφάλαιο 5'):
        labels += ['AI-assisted result narration', 'frozen experiment result']
    elif major.startswith('Κεφάλαιο 6'):
        labels += ['AI-assisted interpretation/synthesis']
        if any(k in text for k in ['Q-Learning','SARSA','Dyna-Q+','DQN','PPO','ανάκαμ','προσαρμογ']): labels += ['derived from frozen results']
    elif major.startswith('Κεφάλαιο 7'):
        labels += ['AI-assisted synthesis']
        if any(k in text for k in ['RQ1','RQ2','RQ3','Q-Learning','SARSA','Dyna-Q+','DQN','PPO']): labels += ['derived from frozen results']
    elif major=='Front matter':
        labels += ['front matter / generated composition']
        if any(k in text for k in ['αποτελέσματα','results','Q-Learning','SARSA','Dyna-Q+','DQN','PPO']): labels += ['AI-assisted synthesis','derived from frozen results']
    elif major=='Βιβλιογραφία': labels += ['bibliography entry']
    elif major=='Παραρτήματα':
        labels += ['appendix material']
        if subsection.startswith('Παράρτημα Α') or subsection.startswith('Α.'):
            labels += ['project methodology/protocol fact']
        elif subsection.startswith('Παράρτημα Β') or subsection.startswith('Β.'):
            labels += ['frozen experiment result/diagnostic']
        elif subsection.startswith('Παράρτημα Γ') or subsection.startswith('Γ.'):
            labels += ['project provenance/traceability fact']
        elif subsection.startswith('Παράρτημα Δ') or subsection.startswith('Δ.'):
            labels += ['project reproducibility/software fact']
    else: labels += ['generated thesis text']
    if is_table and 'bibliography entry' not in labels: labels.append('table')
    return list(dict.fromkeys(labels))

def add_record(kind,text,style=None):
    global major,file_key,subsection
    clean=' '.join(text.split()) if kind=='p' else text.strip()
    if not clean: return
    if kind=='p' and clean in ['Ενημερώστε το πεδίο TOC στο Microsoft Word.','Ενημερώστε το πεδίο καταλόγου σχημάτων στο Microsoft Word.','Ενημερώστε το πεδίο καταλόγου πινάκων στο Microsoft Word.']:
        return
    refs=refs_in(clean)
    labels=labels_for(clean,major,subsection,kind=='table')
    seq[file_key]+=1
    rid=f"{file_key.upper()}-{seq[file_key]:03d}"
    sids=[num_to_sid[n] for n in refs]
    rec={'id':rid,'file':file_key,'major':major,'section':subsection,'kind':kind,'style':style or '', 'text':clean,'citation_numbers':';'.join(map(str,refs)),'source_ids':';'.join(sids),'labels':';'.join(labels)}
    records.append(rec)
    if kind=='p' and style and style.startswith('Heading'):
        level=int(style.split()[-1]) if style.split()[-1].isdigit() else 2
        chapter_text[file_key].append('#'*min(level,3)+' '+clean+'\n')
        return
    meta=[]
    if labels: meta.append('**Κατηγορία:** '+', '.join(labels))
    if refs:
        srcs=[]
        for n in refs:
            sid=num_to_sid[n]
            title=bib.get(n,'')
            srcs.append(f"[{n}] `{sid}`" + (f" — {title}" if title else ''))
        meta.append('**Στήριξη:** '+'; '.join(srcs))
    else:
        meta.append('**Στήριξη:** καμία εξωτερική παραπομπή στο συγκεκριμένο block')
    block=f"### {rid}\n\n"+'  \n'.join(meta)+"\n\n**Κείμενο διπλωματικής:**\n\n"+clean+"\n\n**Δική μου επαναδιατύπωση / σημειώσεις:**\n\n> \n\n---\n"
    chapter_text[file_key].append(block)
    for label in labels:
        category_text[label].append(block)
    for n in refs:
        source_usage[n].append(rid)

for typ,obj in iter_blocks(doc):
    if typ=='p':
        text=obj.text.strip(); style=obj.style.name if obj.style else ''
        if not text: continue
        if style=='Heading 1' and text in major_to_file:
            major=text; file_key=major_to_file[text]; subsection=text
        elif style and style.startswith('Heading'):
            subsection=text
        add_record('p',text,style)
    else:
        rows=[]
        for row in obj.rows:
            cells=[' '.join(c.text.split()).replace('|','\\|') for c in row.cells]
            rows.append(cells)
        if not rows: continue
        md='| '+' | '.join(rows[0])+' |\n| '+' | '.join(['---']*len(rows[0]))+' |\n'
        for r in rows[1:]: md+='| '+' | '.join(r)+' |\n'
        add_record('table',md,'Table')

intro=(f"# Thesis text review pack\n\nGenerated from the current reader-facing Word build.\n\n- DOCX SHA-256: `{sha}`\n- Automated QA status: `{qa.get('status')}`\n- Visible bibliography entries: {qa.get('citation_count')}\n- Purpose: manual authorship/source review and genuine rewrite in the student's own words.\n\nThe blocks below contain the **complete reader-visible thesis text** for this section (except Word-generated TOC/list caches). Each block is tagged by provenance.\n\n")
for key,parts in chapter_text.items():
    (OUT/'chapters'/f'{key}.md').write_text(intro+'\n'.join(parts),encoding='utf-8')

for label,parts in sorted(category_text.items()):
    safe=re.sub(r'[^a-z0-9]+','_',label.lower()).strip('_')
    (OUT/'categories'/f'{safe}.md').write_text(f"# Category: {label}\n\nThese are complete thesis blocks carrying this label. Blocks can appear in more than one category.\n\n"+'\n'.join(parts),encoding='utf-8')

fields=['id','file','major','section','kind','style','labels','citation_numbers','source_ids','text']
with (OUT/'data'/'paragraph_ledger.csv').open('w',encoding='utf-8-sig',newline='') as f:
    w=csv.DictWriter(f,fieldnames=fields); w.writeheader(); w.writerows(records)

source_index=['# Source / citation index','',f'This Word build uses **{len(citation_map)} formal references**. Numeric in-text citations such as `[5]` are IEEE-style **παραπομπές/citations**; the full entries appear in the **Βιβλιογραφία/references** section.','', 'The pack does **not** duplicate full copyrighted papers. For the original-language source, open the canonical `citation-ready` evidence/analysis record or its verified URL/original pointer.','']
for n in sorted(num_to_sid):
    sid=num_to_sid[n]; entry=bib.get(n,'')
    source_index += [f'## [{n}] `{sid}`', '', entry, '', f'- Thesis blocks using this source: {", ".join(source_usage.get(n,[])) or "none detected in visible text"}', f'- Canonical evidence: `research/bibliography/citation-ready/evidence/{sid}.md`', f'- Canonical analysis: `research/bibliography/citation-ready/analyses/{sid}.md`', '- Original-language document / verified URL: follow the “Πρωτότυπο που ελέγχθηκε” and DOI/arXiv/URL fields in the canonical analysis/evidence record.', '']
    (OUT/'sources'/f'{n:02d}_{sid}.md').write_text(f'# [{n}] {sid}\n\n{entry}\n\n## Where it is used in the thesis\n\n'+('\n'.join(f'- `{x}`' for x in source_usage.get(n,[])) or '- No visible block detected.')+f'\n\n## Canonical source records\n\n- `research/bibliography/citation-ready/evidence/{sid}.md`\n- `research/bibliography/citation-ready/analyses/{sid}.md`\n\n## Original-language source\n\nOpen the verified original/URL named inside the canonical analysis record. This pack intentionally does not reproduce the full copyrighted publication.\n',encoding='utf-8')
(OUT/'sources'/'SOURCE_INDEX.md').write_text('\n'.join(source_index),encoding='utf-8')

(OUT/'sources'/'ONLINE_AND_OFFICIAL_WEB_GUIDANCE.md').write_text('''# Online / official-web material\n\nThe final thesis scientific prose is not built from uncited ad-hoc web snippets. External scientific claims are tied to the verified `SRC-*` bibliography layer.\n\nOfficial University/Department web material was used as **format/compliance guidance**, not as a scientific source for RL claims. The canonical project record is:\n\n- `docs/thesis/OFFICIAL_GUIDANCE_SNAPSHOT_2026-09-03.md`\n\nExample completed theses were used only as structure/style context and are **not** scientific sources.\n\nIf a paragraph has no formal citation, that does not mean it came from a website: it is normally project-specific methodology/implementation/result narration or AI-assisted synthesis based on the project's own evidence.\n''',encoding='utf-8')

readme=f'''# Thesis Text Review / Provenance Pack\n\nThis pack is for manually checking what the current thesis says, where each statement comes from, and rewriting it in your own genuine wording after you understand it.\n\n**Build identity**\n- Word SHA-256: `{sha}`\n- QA: `{qa.get('status')}`\n- Formal references: {qa.get('citation_count')}\n- Main-body approx. words: {qa.get('t715_main_word_count_approx')}\n\n## What is inside\n\n- `chapters/`: the complete reader-visible text, separated by chapter/front matter/bibliography/appendices.\n- `categories/`: the same complete blocks regrouped by provenance category.\n- `sources/`: one file per formal source, showing its `SRC-*` identity and exactly which thesis blocks cite it.\n- `data/paragraph_ledger.csv`: machine-readable paragraph/block ledger.\n\n## Provenance labels\n\n- **AI-assisted synthesis**: prose composed/synthesized for the thesis; not a verbatim translation from one source.\n- **externally supported**: the block contains one or more formal IEEE citations.\n- **project methodology/protocol fact**: taken from the frozen protocol/project records.\n- **project implementation fact**: describes the implemented GridWorld/backend/PySide6 system.\n- **frozen experiment result**: values or result narration from T-611/T-612/T-613 accepted evidence.\n- **derived from frozen results**: interpretation/conclusion based on those results.\n- **project provenance/traceability fact**: repository/evidence lineage, not literature.\n\n## Important distinction\n\nThe thesis generally does **not** contain copied source passages. Literature-backed paragraphs are synthesized prose with citations. Therefore you should not treat a thesis paragraph as a translation of one paper; use its listed `SRC-*` records to read the original source, then write the idea in your own words while keeping the citation.\n\n`ONLINE_AND_OFFICIAL_WEB_GUIDANCE.md` explains what came from official web guidance versus formal scientific bibliography.\n'''
(OUT/'README.md').write_text(readme,encoding='utf-8')

all_text=[]
for key in sorted(chapter_text):
    all_text.append(f'# {key}\n')
    for r in records:
        if r['file']==key:
            all_text.append(r['text']+'\n')
(OUT/'FULL_THESIS_TEXT_FOR_REVIEW.md').write_text('\n'.join(all_text),encoding='utf-8')

print('OUT', OUT)
print('records',len(records),'chapter files',len(chapter_text),'category files',len(category_text),'sources',len(citation_map))
